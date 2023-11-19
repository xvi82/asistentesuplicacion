from tkcalendar import DateEntry
import datetime
import PyPDF2
import itertools
import os
import tkinter as tk
from tkinter import filedialog, messagebox, PhotoImage, ttk
from tkinter.ttk import Progressbar
from PyPDF2 import PdfReader
import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import threading
import subprocess
import pyperclip
import re


# Historial del portapapeles y configuración
historial_portapapeles = []
MAX_ANCHO_MENU = 50  # Máximo número de caracteres por elemento en el menú

# FUNCIONES PARA EXTRAER TEXTO CONCRETO
def extract_text_between_phrases(text, start_phrase, end_phrase):
    """
    Extrae y devuelve el texto que se encuentra entre dos frases específicas dentro de una cadena de texto.
    Si alguna de las frases no se encuentra o el orden es incorrecto, devuelve un mensaje indicándolo.
    """
    try:
        start_index = text.find(start_phrase)
        end_index = text.find(end_phrase, start_index)

        if start_index != -1 and end_index != -1 and start_index < end_index:
            return text[start_index + len(start_phrase):end_index].strip()
        else:
            return "No se encontró el texto especificado o el orden de las frases es incorrecto."
    except Exception as e:
        return f"Error al procesar el texto: {e}"

#crear hilos
def ejecutar_en_hilo(funcion):
    """Ejecuta una función dada en un hilo separado."""
    hilo = threading.Thread(target=funcion)
    hilo.start()
    return hilo

#Centrar la ventana
def center_window(root):
    try:
        root.update_idletasks()
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        window_width = root.winfo_width()
        window_height = root.winfo_height()
        x = (screen_width // 2) - (window_width // 2)
        y = (screen_height // 2) - (window_height // 2)
        y -= 34
        root.geometry(f"+{x}+{y}")
    except Exception as e:
        print(f"Error al centrar la ventana: {e}")


# Credenciales de OpenAI API
openai.api_key = os.getenv('OPENAI_API_KEY')

#Preparar texto
def borrar_saltos_linea_vacios(texto):
    return texto.replace("\n\n", "\n")


def tabular_parrafos(texto):
    parrafos = texto.split("\n")
    parrafos_tabular = []
    for parrafo in parrafos:
        parrafos_tabular.append("\t" + parrafo)
    return "\n".join(parrafos_tabular)


def texto_preparado_para_procesador(texto):
    preparado1 = borrar_saltos_linea_vacios(texto)
    preparado2 = tabular_parrafos(preparado1)
    return preparado2


# Definir una función para generar rangos de números como cadenas
def rango_numerico_como_cadenas(inicio, fin):
    return [str(i) for i in range(inicio, fin + 1)]

# Partido judiciales
partido_juzgado_map = {
    "Galdar": ['1'],
    "Arrecife": ['1', '2'],
    "Puerto del Rosario": ['3', '4'],
    "Las Palmas de Gran Canaria": rango_numerico_como_cadenas(1, 11)
}

def word_to_number(word):
    word_number_mapping = {
        "uno": 1, "dos": 2, "tres": 3, "cuatro": 4, "cinco": 5,
        "seis": 6, "siete": 7, "ocho": 8, "nueve": 9,
        "diez": 10, "once": 11, "doce": 12
    }
    return word_number_mapping.get(word.lower(), word)

def get_response(prompt, model):
    response = openai.ChatCompletion.create(
        model=model,
        messages=[
            {"role": "system", "content": prompt},
        ]
    )
    return response['choices'][0]['message']['content']


def extract_data_from_response(response):
    # Uso de expresiones regulares para encontrar todos los textos entre corchetes
    data_matches = re.findall(r'\[(.*?)\]', response)

    if len(data_matches) == 6:  # Verificamos que hay 5 datos como esperado
        numero_procedimiento = data_matches[0]
        fecha_sentencia = data_matches[1]
        numero_juzgado = data_matches[2]
        nombre_demandante = data_matches[3]
        nombre_demandado = data_matches[4]
        partido_judicial = data_matches[5]

        return numero_procedimiento, fecha_sentencia, numero_juzgado, nombre_demandante, nombre_demandado, partido_judicial
    else:
        print("Error: El formato de la respuesta no es el esperado.")
        return None

def extract_full_text_from_pdf(reader):
    return ''.join(page.extract_text() for page in reader.pages)

def extract_first_page_text(reader):
    return reader.pages[0].extract_text()

def extract_info_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        full_text = extract_full_text_from_pdf(reader)
        extracted_text = extract_text_between_phrases(full_text, "DERECHO", "FALLO")

        prompt_for_summary = create_summary_prompt(extracted_text)
        summary_response = get_response(prompt_for_summary, "gpt-4-1106-preview")
        print(summary_response)

        first_page_text = extract_first_page_text(reader)
        combined_prompt = create_data_extraction_prompt(first_page_text)
        response = get_response(combined_prompt, "gpt-3.5-turbo")

        response = convert_response_to_numbers(response)
        extracted_data = extract_data_from_response(response)

        if extracted_data:
            display_extracted_data(extracted_data)
            update_ui_components(extracted_data, summary_response)

def create_summary_prompt(extracted_text):
    return (
        f"Eres un Juez experto en Derecho Laboral y de la Seguridad Social, actua como tal, tienes que hacer un resumen de 370 palabras del sentido de una sentencia y las razones para estimar o desestimar. Habla siempre en pasado. Nunca hables de la posibilidad de recurso. No digas si ha posibilidad de recurso. Utiliza a lo largo del texto expresiones como 'la sentencia de instancia/la resolución combatida/el pronunciamiento impugnado   apreció/consideró probado/resolvió/entendió'. Empieza el resumen con la expresión 'La sentencia de instancia estimaba/desestimaba la demanda'. La sentencia es la siguiente:\n: {extracted_text}"
    )

def create_data_extraction_prompt(first_page_text):
    return (
        f"A partir del siguiente texto: '{first_page_text}':\n"
        "- Indica el número de procedimiento sin ceros al principio.\n"
        "- Proporciona la fecha de la sentencia, con este formato '%d/%m/%Y'.\n"
        "- Menciona el número del juzgado en formato número, 1...12.\n"
        "- Lista el o los nombres del demandante (si hay más de uno, sepáralos con 'y').\n"
        "- Especifica el nombre del demandado (excluye 'el FOGASA' o 'Fondo de Garantía Salarial').\n"
        "- Especifíca el nombre del partido judicial, puede ser: Las Palmas de Gran Canaria, Arrecife, Galdar o Puerto del Rosario.\n"
        "- Sólo quiero que pongas los datos con el siguiente formato [número del procedimiento sin ceros al principio] [fecha de la sentencia] [número del juzgado] [nombre del demandante] [nombre del demandado] [nombre del partido judicial].\n"
        "- Pon cada dato entre corchetes."
    )

def convert_response_to_numbers(response):
    return " ".join(word_to_number(part) for part in response.split(" "))

def display_extracted_data(extracted_data):
    # Imprime los datos extraídos
    numero_procedimiento, fecha_sentencia, numero_juzgado, nombre_demandante, nombre_demandado, partido_judicial = extracted_data
    print(f"Número de Procedimiento: {numero_procedimiento}")
    print(f"Fecha de Sentencia: {fecha_sentencia}")
    print(f"Nombre del Partido Judicial: {partido_judicial}")
    print(f"Número de Juzgado: {numero_juzgado}")
    print(f"Nombre del Demandante: {nombre_demandante.title()}")
    print(f"Nombre del Demandado: {nombre_demandado}")

def update_ui_components(extracted_data, summary_response):
    # Actualizar componentes de la UI
    numero_procedimiento, fecha_sentencia, numero_juzgado, _, _, partido_judicial = extracted_data
    partido_var.set(partido_judicial)
    update_juzgado()
    juzgado_var.set(numero_juzgado)
    fecha_cal.set_date(datetime.datetime.strptime(fecha_sentencia, '%d/%m/%Y').date())
    autos_var.set(numero_procedimiento.split('/')[0])
    autos_year_var.set(numero_procedimiento.split('/')[1])
    sintesis_var.set(texto_preparado_para_procesador(summary_response))
    update_text()


def cambiar_estado_boton(boton, estado):
    boton['state'] = estado


pdf_path = None  # Esto se define en la parte superior del script


def cargar_pdf():
    global pdf_path  # Indicar que vamos a usar la variable global
    progress_bar.grid(row=0, column=2, padx=2, pady=2, sticky="ew")  # Mostras la barra de progreso
    progress_bar.start(5)  # Iniciar la animación de la barra
    # Abre la ventana de diálogo para seleccionar un archivo PDF
    filepath = filedialog.askopenfilename(
        title="Seleccionar archivo PDF",
        filetypes=[("Archivos PDF", "*.pdf")]
    )

    if filepath:
        # Aquí puedes hacer algo con el archivo seleccionado
        pdf_path = filepath  # Aquí asignas el valor seleccionado a la variable global
        extract_info_from_pdf(filepath)
        print(f"Archivo seleccionado: {filepath}")
        # Por ejemplo, puedes almacenar la ruta del archivo en una variable,
        # abrir y leer el archivo, etc.
    cambiar_estado_boton(recurso_button, 'normal')
    progress_bar.stop()  # Detener la animación
    progress_bar.grid_remove()  # Ocultar la barra de progreso



def update_juzgado(*args):
    partido = partido_var.get()
    juzgados = partido_juzgado_map.get(partido, [])
    juzgado_var.set(juzgados[0] if juzgados else '')
    juzgado_options['values'] = juzgados

def fecha_a_texto(fecha):
    meses = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
        7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    dia = fecha.day
    mes = meses[fecha.month]
    año = fecha.year
    return f"{dia} de {mes} de {año}"


# Pre-compilar expresiones regulares
motivo_ordinal_pattern = re.compile(r"(único|primer|segundo|tercer|cuarto|quinto|sexto|séptimo|octavo|noveno|décimo) motivo")
comillas_pattern = re.compile(r'"([^"]*?)"|“(.*?)”|\'([^\']*?)\'|«(.*?)»')

def save_to_doc():
    # Verificación de los campos de recurso
    numero_recurso = recurso_var.get()
    año_recurso = recurso_year_var.get()

    if not numero_recurso:
        tk.messagebox.showwarning("Advertencia", "Por favor, ingrese el número de recurso.")
        recurso_entry.focus()
        return
    elif not año_recurso:
        tk.messagebox.showwarning("Advertencia", "Por favor, seleccione el año del recurso.")
        recurso_year_options.focus()
        return
    resumen = on_button_click()
    doc = configurar_documento()

    text_lines = right_text.get("1.0", tk.END).split('\n')
    procesar_lineas(text_lines, doc)

    # Guardar y abrir el documento
    filename = guardar_documento(doc, resumen)
    abrir_documento(filename)

def configurar_documento():
    doc = Document()
    # Configuración del estilo...
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return doc

def procesar_lineas(text_lines, doc):
    for line in text_lines:
        # Procesamiento de cada línea
        if line.startswith(("PRIMERO.-", "SEGUNDO.-", "TERCERO.-", "CUARTO.-", "QUINTO.-", "SEXTO.-", "SÉPTIMO.-", "OCTAVO.-")):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = p.add_run(line)
            run.bold = True
        elif "FALLAMOS" in line:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(line)
            run.bold = True
        elif '"' in line:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            parts = re.split(r'(".*?")', line)
            for part in parts:
                if part.startswith('"') and part.endswith('"'):
                    run = p.add_run(part)
                    run.italic = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph()
            start = 0
            for match in sorted(itertools.chain(comillas_pattern.finditer(line), motivo_ordinal_pattern.finditer(line)), key=lambda m: m.start()):
                p.add_run(line[start:match.start()])
                run = p.add_run(match.group(0))
                if comillas_pattern.match(match.group(0)):
                    run.italic = True
                elif motivo_ordinal_pattern.match(match.group(0)):
                    run.bold = True
                start = match.end()
            p.add_run(line[start:])

def guardar_documento(doc, resumen):
    meses = {
        "January": "enero", "February": "febrero", "March": "marzo", "April": "abril",
        "May": "mayo", "June": "junio", "July": "julio", "August": "agosto",
        "September": "septiembre", "October": "octubre", "November": "noviembre", "December": "diciembre"
    }
    now = datetime.datetime.now()
    nombre_mes = meses[now.strftime("%B")]
    numero_recurso = recurso_var.get()
    año_recurso = recurso_year_var.get()
    directory_path = os.path.dirname(os.path.realpath(__file__))
    filename = os.path.join(directory_path, f"Rec. {numero_recurso}-{año_recurso} - {resumen}.docx")
    doc.save(filename)
    print(f"Documento guardado en: {filename}")
    return filename

def abrir_documento(filename):
    try:
        if os.name == 'nt':
            os.startfile(filename)
        elif os.name == 'posix':
            subprocess.call(('open', filename))
        else:
            subprocess.call(('xdg-open', filename))
    except Exception as e:
        print(f"Error abriendo el archivo: {e}")

def on_button_click():
    prompt = "Haz un resumen de 10 palabras del siguiente texto. Pon las palabras más significativas del caso. El resumen debe ser telegráfico:\n" + sintesis_var.get() # Obtener texto del widget de entrada
    response = get_response(prompt, "gpt-3.5-turbo")
    # Eliminar caracteres específicos de la respuesta
    response = response.replace('.', '').replace(',', '').replace(';', '').replace(':', '')
    return response


def clear_all():
    reset_text_variables([partido_var, recurrente_var, autos_var, autos_year_var, sintesis_var,
                          articulos_var, jurisprudencia_var, revocacion_var, revocacion_parcial_var,
                          razon_var, juzgado_var, revisiones_var, censuras_var, recurso_year_var, recurso_var,
                          impugnante_var])

    reset_checkboxes([cuestiones_var, inadmisibilidad_var, infraccion_var,
                      revision_var, sustantivas_var])

    reset_radiobuttons()

    clear_text_fields([right_text])

    # Actualizar interfaz y elementos dependientes
    fecha_cal.set_date(datetime.date.today())
    update_ui_elements()


def reset_text_variables(variables):
    for var in variables:
        var.set('')


def reset_checkboxes(checkbox_vars):
    for chk_var in checkbox_vars:
        chk_var.set(False)


def reset_radiobuttons():
    costas_var.set(None)
    trabajador_empresa_var.set("Trabajador")
    for tipo_var in tipos_vars:
        tipo_var.set(False)


def clear_text_fields(text_widgets):
    for text_widget in text_widgets:
        text_widget.delete("1.0", tk.END)


def update_ui_elements():
    update_juzgado()
    toggle_infraccion()
    toggle_costas_frames()
    toggle_censuras()


# Función que inserta la respuesta en la sección adecuada
def insertar_respuesta_revision(respuesta_texto):
    # Encuentra la posición en el texto donde queremos insertar la respuesta
    busqueda_frase = "Sentado lo anterior, pasamos a analizar las pretensiones concretas."
    content = right_text.get("1.0", tk.END)

    # Encuentra el final de la frase donde agregaremos la respuesta
    inicio_index = content.find(busqueda_frase)
    if inicio_index != -1:
        fin = inicio_index + len(busqueda_frase)
        pos = right_text.search(busqueda_frase, "1.0", tk.END)
        if pos:
            # Obtiene la posición de la línea y la columna para insertar texto
            linea, columna = pos.split('.')
            posicion = f"{linea}.{int(columna) + len(busqueda_frase)}"
            # Agrega un salto de línea antes y después del texto insertado
            right_text.insert(posicion, "\n\n" + respuesta_texto)
            right_text.see(posicion)  # Desplaza el texto para que la inserción sea visible
    else:
        print("No se encontró la frase para insertar la respuesta.")

def update_text(*args):
    fecha_seleccionada = fecha_a_texto(fecha_cal.get_date())

    ordinals = ["SEGUNDO.-", "TERCERO.-", "CUARTO.-", "QUINTO.-", "SEXTO.-", "SÉPTIMO.-", "OCTAVO.-", "NOVENO.-"]


    content_sections = [f"PRIMERO.- Síntesis de la litis\n{sintesis_var.get()}"]

    index = 0

    if cuestiones_var.get():
        content_sections.append(f"""{ordinals[index]} Cuestiones de orden público
	Con carácter prioritario, debemos recordar en los términos de la Sentencia del Tribunal Supremo de 11 de mayo de 2018 (Pleno, rec. 1800/2016), reiterada por otras posteriores, que la cuestión de la recurribilidad de la Sentencia de instancia ha de examinarse de oficio por afectar al orden público procesal y a la propia competencia funcional, incluso “antes de llevar a cabo cualquier pronunciamiento sobre la contradicción o sobre el fondo del asunto”.
	El artículo 190 de la Ley Reguladora de la Jurisdicción Social (en adelante LRJS) establece con carácter general el ámbito de la suplicación, precisando de forma positiva y negativa el artículo 191 LRJS las resoluciones susceptibles de ser combatidas a través del citado recurso extraordinario.
	En el presente caso...
""")
        index += 1

    if inadmisibilidad_var.get():
        content_sections.append(f"""{ordinals[index]} Inadmisibilidad del recurso 
	Por razones de sistemática hemos de comenzar con el análisis y resolución del motivo de inadmisibilidad alegado por la parte recurrida, cuya estimación haría innecesaria la resolución del resto de motivos impugnatorios. 
	El artículo 190 de la Ley Reguladora de la Jurisdicción Social (en adelante LRJS) establece con carácter general el ámbito de la suplicación, precisando de forma positiva y negativa el artículo 191 LRJS las resoluciones susceptibles de ser combatidas a través del citado recurso extraordinario. 
	En el presente caso…
""")
        index += 1

    if infraccion_var.get():
        infraccion_content = [f"""{ordinals[index]} Infracción de normas procesales
	La nulidad de actuaciones sólo puede articularse por la vía del art. 193.a) LRJS, que exige como requisitos para que pueda prosperar el motivo aducido, los siguientes: 
		1) Identificar el precepto procesal que se entienda infringido o doctrina emanada del TS, TC o los órganos jurisdiccionales instituidos en los Tratados y Acuerdos internacionales en materia de derechos humanos y libertades fundamentales ratificados por España, o la del TJUE (art. 219.2 LRJS y art. 1.6 CC) 
		2) La infracción debe haber provocado un perjuicio real sobre los derechos de defensa del interesado, irrogándole indefensión (STC 168/2002). 
		3) El defecto procesal no puede alegarse por la parte que lo provocó. 
		4) Es preciso que la parte perjudicada haya formulado protesta en tiempo y forma.
	Ha de recordarse que es criterio reiterado de las diversas salas de suplicación que la declaración de nulidad de actuaciones es un remedio excepcional, que ha de aplicarse con criterio restrictivo, pues una interpretación amplia de la posibilidad de anulación podría incluso vulnerar el derecho a la tutela judicial efectiva, proclamado en el artículo 24 CE, en su vertiente del derecho a un proceso público sin dilaciones indebidas. A este respecto, no solamente es la celeridad uno de los principios orientadores de la interpretación y aplicación de las normas reguladoras del proceso laboral ordinario (artículo 74.1 LJS) sino que la Ley Reguladora de la Jurisdicción Social aumenta considerablemente las posibilidades de la Sala de suplicación de examinar el fondo del asunto, tanto por vulneración de las normas reguladoras de la sentencia (artículo 202.2) como por no haber entrado la resolución de instancia en el fondo al haber apreciado alguna circunstancia obstativa, como puede ser una excepción procesal (artículo 202.3), siempre que el relato de hechos probados de la sentencia y demás antecedentes no cuestionados sean suficientes para realizar tal pronunciamiento sobre el fondo."""]
        for i, var in enumerate(tipos_vars):
            if var.get():
                infraccion_content.append(tipos_options[i])
        content_sections.append("\n".join(infraccion_content))
        content_sections.append("	En el presente caso...\n")
        index += 1
    else:
        for var in tipos_vars:
            var.set(False)  # If "Infracción de normas procesales" is unchecked, uncheck all related checkboxes

    if revision_var.get():
        content_sections.append(f"""{ordinals[index]} Revisión de hechos probados 
	La parte recurrente interesa la revisión fáctica de la sentencia, al amparo del art. 193.b) LRJS. 
	En primer lugar, debe razonarse que, con carácter general, el órgano que conoce del recurso extraordinario de suplicación no puede efectuar una nueva ponderación de la prueba, pues es a quien ha presidido el acto del juicio en la instancia a quien corresponde apreciar los elementos de convicción para establecer la verdad procesal intentando que la misma se acerque lo más posible a la verdad material. Ahora bien, tal principio debe ser matizado en el sentido de que el tribunal ad quem está autorizado para revisar las conclusiones fácticas cuando los documentos o pericias citados por la parte recurrente pongan de manifiesto de manera patente e incuestionable el error en el que ha incurrido la sentencia recurrida, o la irracionalidad o arbitrariedad de sus conclusiones. De otra forma, carecería de sentido la previsión del artículo 193.b) LRJS. 
	Además, debe señalarse que la jurisprudencia viene exigiendo con reiteración, hasta el punto de constituir doctrina pacifica, que para estimar este motivo es necesario que concurran los siguientes requisitos: 
		1.- Que se señale con precisión y claridad cuál es el hecho afirmado, negado u omitido, que la parte recurrente considera equivocado, contrario a lo acreditado o que consta con evidencia y no ha sido incorporado al relato fáctico. 
		2.- Que se ofrezca un texto alternativo concreto para figurar en la narración fáctica calificada de errónea, bien sustituyendo a alguno de sus puntos, bien complementándolos. 
		3.- Que se citen pormenorizadamente los documentos o pericias de los que se considera se desprende la equivocación de la sentencia, sin que sea dable admitir su invocación genérica, ni plantearse la revisión de cuestiones fácticas no discutidas a lo largo del proceso; la ley señala que el error debe ponerse de manifiesto precisamente merced a las pruebas documentales o periciales practicadas en la instancia. 
		4.- Que esos documentos o pericias pongan de manifiesto el error de manera clara, evidente, directa y patente; sin necesidad de acudir a conjeturas, suposiciones o argumentaciones más o menos lógicas, naturales y razonables, de modo que sólo son admisibles para poner de manifiesto el error de hecho, los documentos que ostenten un decisivo valor probatorio, tengan concluyente poder de convicción por su eficacia, suficiencia, fehaciencia o idoneidad. 
		5.- Que la revisión pretendida sea trascendente a la parte dispositiva de la sentencia, con efectos modificadores de ésta, pues el principio de economía procesal impide incorporar hechos cuya inclusión a nada práctico conduciría, si bien cabrá admitir la modificación fáctica cuando no siendo trascendente en esta instancia pudiera resultarlo en otras superiores. 
		6.- Que no se trate de una nueva valoración global de la prueba incorporada al proceso. 
	Estos criterios han sido reafirmados, entre otras, por la reciente Sentencia 90/2022 del Tribunal Supremo, de 1 de febrero (rec. 2429/2019). 
	Sentado lo anterior, pasamos a analizar las pretensiones concretas.""")
        if revisiones_var.get():
            content_sections.append(tabular_parrafos(revisiones_var.get()))
        else:
            content_sections.append("")
        index += 1

    if sustantivas_var.get():
        content_sections.append(f"""{ordinals[index]} Infracción de normas sustantivas y de la jurisprudencia """)
        articulos = articulos_var.get().strip()
        jurisprudencia = jurisprudencia_var.get().strip()

        if articulos and jurisprudencia:
            content_sections.append(
                f"""	La parte recurrente interesa la revocación de la sentencia, al amparo del art. 193.c) LRJS, alegando la infracción de los arts. {articulos}, así como la infracción de la jurisprudencia sentada por {jurisprudencia}. 
	Con carácter previo debemos indicar – por lo que respecta a las normas citadas por la parte recurrente en su recurso como infringidas – que el recurso de suplicación no es una apelación o segunda instancia, sino un recurso extraordinario sujeto a motivos tasados en cuya formulación se han de respetar los requisitos legales.
	Los motivos basados en el apartado c) del art. 193 LRJS se destinan a la impugnación del fallo por error in iudicando, y el recurrente tiene la carga de:
		a)	Citar debidamente el precepto o preceptos sustantivos y en su caso la jurisprudencia que, a su juicio, han sido vulnerados por el fallo de la sentencia, articulando motivos separados para cada precepto o grupo de preceptos que guarden unidad temática.
		b)	Razonar la pertinencia y fundamentación de los motivos (art. 196.2 LRJS) lo cual exige argumentar la conexión entre el contenido normativo de las normas o jurisprudencia citadas y el litigio, mostrando cómo su correcta aplicación debería haber llevado a dar distinta solución al debate.
	Dicha doctrina jurisprudencial, así mismo señala que no basta que el recurso cite la disposición legal conculcada si contiene diversos artículos, sino que es preciso que se señale el específico precepto que se entiende vulnerado, y si el precepto contiene varios apartados resulta igualmente indispensable señalar expresamente cuál de ellos se reputa infringido.""")
            if censuras_var.get():
                content_sections.append(tabular_parrafos(censuras_var.get()))
        elif articulos:
            content_sections.append(f"""	La parte recurrente interesa la revocación de la sentencia, al amparo del art. 193.c) LRJS, alegando la infracción de los arts. {articulos}.
	Con carácter previo debemos indicar – por lo que respecta a las normas citadas por la parte recurrente en su recurso como infringidas – que el recurso de suplicación no es una apelación o segunda instancia, sino un recurso extraordinario sujeto a motivos tasados en cuya formulación se han de respetar los requisitos legales.
	Los motivos basados en el apartado c) del art. 193 LRJS se destinan a la impugnación del fallo por error in iudicando, y el recurrente tiene la carga de:
		a)	Citar debidamente el precepto o preceptos sustantivos y en su caso la jurisprudencia que, a su juicio, han sido vulnerados por el fallo de la sentencia, articulando motivos separados para cada precepto o grupo de preceptos que guarden unidad temática.
		b)	Razonar la pertinencia y fundamentación de los motivos (art. 196.2 LRJS) lo cual exige argumentar la conexión entre el contenido normativo de las normas o jurisprudencia citadas y el litigio, mostrando cómo su correcta aplicación debería haber llevado a dar distinta solución al debate.
	Dicha doctrina jurisprudencial, así mismo señala que no basta que el recurso cite la disposición legal conculcada si contiene diversos artículos, sino que es preciso que se señale el específico precepto que se entiende vulnerado, y si el precepto contiene varios apartados resulta igualmente indispensable señalar expresamente cuál de ellos se reputa infringido.""")
            if censuras_var.get():
                content_sections.append(tabular_parrafos(censuras_var.get()))
        index += 1

    # Añadir el texto de "Costas, depósitos y consignaciones"
    if costas_var.get() == "Estimación":
        content_sections.append(f"""{ordinals[index]} Costas, depósitos y consignaciones""")
        if trabajador_empresa_var.get() == "Trabajador":
            content_sections.append(f"""	La estimación del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, implica que no haya lugar a hacer pronunciamiento alguno sobre las costas del recurso. 


	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	ESTIMAR el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, revocando la misma en el sentido de que:
""")
        else:
            content_sections.append(f"""	La estimación del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, implica que no haya lugar a hacer pronunciamiento alguno sobre las costas del recurso.
	Se decreta también, como preceptúa el artículo 203.1 LRJS, la devolución de la totalidad del depósito constituido para recurrir y dar el destino legal a los aseguramientos si los hubiere, todo ello firme que sea la presente resolución. 


	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	ESTIMAR el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, revocando la misma en el sentido de que:
""")
        content_sections.append(f'"{revocacion_entry.get()}"')
        content_sections.append(f"""\n	Sin costas.""")
    elif costas_var.get() == "Estimación parcial":
        content_sections.append(f"""{ordinals[index]} Costas, depósitos y consignaciones""")
        if trabajador_empresa_var.get() == "Trabajador":
            content_sections.append(f"""	La estimación parcial del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, implica que no haya lugar a hacer pronunciamiento alguno sobre las costas del recurso. 


	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	ESTIMAR PARCIALMENTE el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, revocando la misma en el sentido de que:
""")
        else:
            content_sections.append(f"""	La estimación parcial del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, implica que no haya lugar a hacer pronunciamiento alguno sobre las costas del recurso.
	Se decreta también, como preceptúa el artículo 203.1 LRJS, la devolución de la totalidad del depósito constituido para recurrir y dar el destino legal a los aseguramientos si los hubiere, todo ello firme que sea la presente resolución. 


	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	ESTIMAR PARCIALMENTE el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, revocando la misma en el sentido de que:
""")
        content_sections.append(f'"{revocacion_parcial_entry.get()}"')
        content_sections.append(f"""\n	Sin costas.""")
    elif costas_var.get() == "Desestimación con costas":
        content_sections.append(f"""{ordinals[index]} Costas, depósitos y consignaciones
	La desestimación del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, implica la expresa imposición de condena en las costas del recurso a la parte vencida en el mismo, condena que incluye el pago de la Minuta de Honorarios del Letrado de la parte impugnante, en la cuantía que esta Sala, y dentro de los límites legales, señala en {cuantia_costas_var.get()} euros.
	Se decreta también, como preceptúa el artículo 204.4 LRJS, la pérdida del depósito constituido para recurrir que, una vez sea firme la sentencia, se ingresará al Tesoro público; y por fin, se condena a la misma parte vencida a la pérdida de la cantidad consignada a la que se dará el destino legal.

	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	DESESTIMAR el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, confirmando la misma en su integridad.

	Se imponen las costas a la parte recurrente, en la cuantía de {cuantia_costas_var.get()} euros. Se decreta la pérdida del depósito y las consignaciones efectuadas, en su caso, para recurrir.""")
    elif costas_var.get() == "Desestimación sin costas":
        content_sections.append(f"""{ordinals[index]} Costas, depósitos y consignaciones
	La desestimación del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, no implica en el presente caso la expresa imposición de condena en las costas del recurso a la parte vencida en el mismo, al ser {razon_var.get()}.

	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	DESESTIMAR el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, confirmando la misma en su integridad.

	Sin costas""")

    final_text = "\n".join(content_sections)
    right_text.delete("1.0", tk.END)
    right_text.insert(tk.END, final_text)

# FUNCIONES PARA LA CARGA DEL RECURSO

pdf_path_for_text_extraction = None  # Definición de la nueva variable global
def load_pdf_and_get_text():
    global pdf_path_for_text_extraction  # Usar la nueva variable global

    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if file_path:
        pdf_path_for_text_extraction = file_path  # Almacenar la ruta del archivo PDF en la nueva variable global
        reader = PdfReader(file_path)
        text = ''
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    else:
        return None

def numero_a_palabra(n):
    numeros = ["cero", "un", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve", "diez"]
    return numeros[n] if n < len(numeros) else str(n)

def get_response3(pdf_text):
    contenido_pdf = f"ESTE ES EL RECURSO:\n{pdf_text}\n###########\n"
    response = openai.ChatCompletion.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": contenido_pdf},
            {"role": "user", "content": """Eres un experto analizador de recursos de suplicación. Quiero que analices el Recurso. Los recursos de suplicación pueden interponerse por el art. 193.a) LRJS, art. 193.b) LRJS o art. 193.c) LRJS.
            - El apartado a) del art. 193 LRJS regula un motivo de recurso que pretende reponer los autos al estado en el que se encontraban en el momento de cometerse una infracción de normas o garantías del procedimiento que haya producido indefensión.
            - El apartado b) del art. 193 LRJS regula un motivo de recurso que pretende revisar los hechos declarados probados, a la vista de las pruebas documentales y periciales practicadas.
            - El apartado c) del art. 193 LRJS regula un motivo de recurso que pretende examinar las infracciones de normas sustantivas o de la jurisprudencia.
            - Tienes que decirme, cuántos motivos suplicacionales hay de la letra a), cuántos motivos suplicacionales hay de la letra b) y cuántos motivos suplicacionales hay de la letra c).
            - Extrae también el nombre del recurrente, NO PONGAS EL NOMBRE DEL ABOGADO, asegurate.
            - Si hay algún motivo del apartado c), examinalo detenidamente, y haz una lista de los artículos que dice que se han infringido. Haz una relación de ellos. Por ejemplo [art. 1 ET, art. 24 RD 1457/2023, art. 9 LGSS].
            - El formato de la respuesta será el siguiente: [nombre recurrente] [numero a)] [numero b)] [numero c)] [artículos infringidos], por ejemplo [Pedro Rodríguez] [0] [4] [2] [art. 1 ET, art. 24 RD 1457/2023, art. 9 LGSS].
            - Respeta el formato, tienes prohibido cambiar el formato.
            - El formato de la respuesta será el siguiente: [nombre recurrente] [numero a)] [numero b)] [numero c)] [artículos infringidos], por ejemplo [Pedro Rodríguez] [0] [4] [2] [art. 1 ET, art. 24 RD 1457/2023, art. 9 LGSS].
            - Respeta el formato, tienes prohibido cambiar el formato.
            - Take a deep breath and work on this problem step by step."""}
        ]
    )
    return response['choices'][0]['message']['content']

motivos_b = 0  # Definiendo la variable a nivel global
motivos_c = 0  # Definiendo la variable a nivel global
def interpret_response(response_str):
    global motivos_b  # Declarando que vamos a usar la variable global
    global motivos_c  # Declarando que vamos a usar la variable global
    elements = response_str.strip('[]').split('] [')
    if len(elements) == 5:
        nombre_recurrente = elements[0].title()
        recurrente_var.set(nombre_recurrente)
        motivos_a = int(elements[1])
        infraccion_var.set(motivos_a > 0)
        motivos_b = int(elements[2])
        revision_var.set(motivos_b > 0)
        motivos_c = int(elements[3])
        sustantivas_var.set(motivos_c > 0)
        articulos_infringidos = elements[4]
        if articulos_infringidos:
            articulos_var.set(f'{articulos_infringidos}')
            toggle_censuras()
        else:
            articulos_var.set('NO SE HAN PODIDO IDENTIFICAR LOS ARTÍCULOS, INTRODÚCELOS')
            toggle_censuras()

        partes_frase = []
        if motivos_a > 0:
            partes_frase.append(
                f"{numero_a_palabra(motivos_a)} motivo{'s' if motivos_a > 1 else ''} de infracción de normas o garantías procesales")
        if motivos_b > 0:
            partes_frase.append(
                f"{numero_a_palabra(motivos_b)} motivo{'s' if motivos_b > 1 else ''} de revisión fáctica")
        if motivos_c > 0:
            partes_frase.append(
                f"{numero_a_palabra(motivos_c)} motivo{'s' if motivos_c > 1 else ''} de censura jurídica")

        # Construir la frase base
        frase_base = f"\nDisconforme la parte actuante, {nombre_recurrente}, interpone el presente recurso de suplicación articulando "

        # Unir las frases de motivos
        if len(partes_frase) > 1:
            frase_motivos = ', '.join(partes_frase[:-1]) + f" y {partes_frase[-1]}"
        elif len(partes_frase) == 1:
            frase_motivos = partes_frase[0]
        else:
            frase_motivos = ""

        # Completar la frase final
        if frase_motivos:
            frase_final = frase_base + frase_motivos + ", pretendiendo la revocación de la Sentencia con la condena de la parte demandada. "
        else:
            frase_final = frase_base + "pretendiendo la revocación de la Sentencia con la condena de la parte demandada. "

        if impugnante_var.get():
            frase_final += "El recurso fue impugnado por la representación letrada de " + impugnante_var.get() + "\n"
        else:
            frase_final += "\n"
        sintesis_entry.insert(tk.END, tabular_parrafos(frase_final))
        update_text()
        pass
    else:
        return "El formato de respuesta no es correcto."


def analyze_pdf():
    progress_bar.grid(row=0, column=2, padx=2, pady=2, sticky="ew")  # Mostras la barra de progreso
    progress_bar.start(10)  # Iniciar la animación de la barra
    pdf_text = load_pdf_and_get_text()
    if pdf_text:
        gpt_response = get_response3(pdf_text)
        formatted_response = interpret_response(gpt_response)  # Usa la función para formatear la respuesta

    else:
        messagebox.showinfo("Información", "No se ha cargado ningún archivo PDF.")
    cambiar_estado_boton(analizar_revisiones_button, 'normal')
    cambiar_estado_boton(analizar_censuras_button, 'normal')
    progress_bar.stop()  # Detener la animación
    progress_bar.grid_remove()  # Ocultar la barra de progreso

#FIN DE LAS FUNCIONES PARA LA CARGA DEL RECURSO

#ANALISIS DE LAS REVISIONES FÁCTICAS

# Función para unir los PDFs y añadir los textos necesarios
def unir_pdfs_y_anadir_texto(pdf1, pdf2):
    texto_unido = "ESTA ES LA SENTENCIA:\n\n"
    texto_unido += leer_pdf(pdf1)
    texto_unido += "\n\nESTE ES EL RECURSO:\n\n"
    texto_unido += leer_pdf(pdf2)
    return texto_unido

# Función para leer el contenido de un archivo PDF
def leer_pdf(archivo):
    with open(archivo, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        texto = ""
        for page in reader.pages:
            texto += page.extract_text() + "\n"
    return texto

# Función para obtener la respuesta de ChatGPT
def get_response4(prompt, texto_unido):
    response = openai.ChatCompletion.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": texto_unido},
            {"role": "user", "content": prompt},
        ]
    )
    return response['choices'][0]['message']['content']

# FUNCIÓN PARA CONSTRUIR EL PROMPT

def reemplazar_motivos(texto, numero_de_repeticiones):
    # Mapeo de los ordinales
    ordinales = {
        "primero": "primer",
        "segundo": "segundo",
        "tercero": "tercer",
        "cuarto": "cuarto",
        "quinto": "quinto",
        "sexto": "sexto",
        "séptimo": "séptimo",
        "octavo": "octavo",
        # Añade más si es necesario
    }

    # Reemplazar "motivo primero" por "único motivo" si numero_de_repeticiones es 1
    if numero_de_repeticiones == 1:
        texto = texto.replace("motivo primero", "único motivo")
    else:
        # Reemplazar cada "motivo [ordinal]" por "[ordinal] motivo"
        for ordinal in ordinales:
            texto = texto.replace(f"motivo {ordinal}", f"{ordinales[ordinal]} motivo")

    return texto

def generar_respuestas(prompt_base, numero_de_repeticiones, pdf2):
    ordinales = ["primero", "segundo", "tercero", "cuarto", "quinto", "sexto", "séptimo", "octavo"]
    respuesta_final = ""

    for i in range(numero_de_repeticiones):
        # Reemplazar [numero] por el número de repeticiones y ajustar singular o plural
        if numero_de_repeticiones == 1:
            prompt_con_numero = prompt_base.replace("[numero]", "1 motivo")
        else:
            prompt_con_numero = prompt_base.replace("[numero]", f"{numero_de_repeticiones} motivos")

        # Reemplazar todas las instancias de 'primero' por el ordinal correspondiente
        prompt_modificado = prompt_con_numero.replace("primero", ordinales[i])

        # Generar la frase con los ordinales para la frase 'No analices ni el motivo...'
        frase_ordinales = ", ni el ".join(ordinales[:i])
        if frase_ordinales:
            frase_ordinales = "No analices el motivo " + frase_ordinales + ", ya los analizaste antes. "

        # Insertar la frase en el lugar correcto del prompt modificado
        partes_prompt = prompt_modificado.split("El motivo puede ser una modificación de un hecho probado o la adición de un nuevo hecho probado o de un nuevo párrafo a un hecho probado ya existente.")
        prompt_modificado = partes_prompt[0] + frase_ordinales + "El motivo puede ser una modificación de un hecho probado o la adición de un nuevo hecho probado o de un nuevo párrafo a un hecho probado ya existente." + partes_prompt[1]
        print(prompt_modificado)
        respuesta = get_response4(prompt_modificado, pdf2)
        pyperclip.copy(respuesta)
        respuesta_final += f"{respuesta}\n"

    return respuesta_final

# Función para manejar
def analizar_revisiones_facticas():
    progress_bar.grid(row=0, column=2, padx=2, pady=2, sticky="ew")  # Mostras la barra de progreso
    progress_bar.start(5)  # Iniciar la animación de la barra
    with open(pdf_path, 'rb') as pdf_file:
        reader = PyPDF2.PdfReader(pdf_path)
        full_text = ''
        for page in reader.pages:
            full_text += page.extract_text()

        start_phrase = "PROBADOS"
        end_phrase = "FUNDAMENTOS"
        extracted_text = extract_text_between_phrases(full_text, start_phrase, end_phrase)

        # Añadir texto al inicio y al final de extracted_text
        extracted_text = 'ESTOS SON LOS HECHOS PROBADOS\n\n' + extracted_text + '\n\nESTE ES EL RECURSO\n\n'

    # Leer el segundo PDF y añadir su contenido a extracted_text
    pdf2_content = leer_pdf(pdf_path_for_text_extraction)
    extracted_text += pdf2_content
    print(extracted_text)
    numero_de_repeticiones = motivos_b  # Puedes cambiar este número según lo necesites
    respuesta_final = generar_respuestas(promptrevision, numero_de_repeticiones, extracted_text)
    respuesta_final = reemplazar_motivos(respuesta_final, numero_de_repeticiones)
    progress_bar.stop()  # Detener la animación
    progress_bar.grid_remove()  # Ocultar la barra de progreso

# Asegúrate de definir la variable 'prompt' fuera de la función, como una variable global o directamente en la función.
    revisiones_var.set(respuesta_final)  # Devuelve la respuesta completa para su uso posterior)


promptrevision = """Eres un experto analizador de recursos de suplicación. Quiero que analices el Recurso. Los recursos de suplicación pueden interponerse por el art. 193.a) LRJS, art. 193.b) LRJS o art. 193.c) LRJS.
            - El apartado a) del art. 193 LRJS regula un motivo de recurso que pretende reponer los autos al estado en el que se encontraban en el momento de cometerse una infracción de normas o garantías del procedimiento que haya producido indefensión.
            - El apartado b) del art. 193 LRJS regula un motivo de recurso que pretende revisar los hechos declarados probados, a la vista de las pruebas documentales y periciales practicadas.
            - El apartado c) del art. 193 LRJS regula un motivo de recurso que pretende examinar las infracciones de normas sustantivas o de la jurisprudencia.
            - Tienes que buscar los motivos suplicacionales de la letra b).
            - En el documento hay [numero] motivos de la letra b).
            - Una vez que los encuentres centrate sólo en el motivo primero. Eres un experto contanto motivos. Sólo quiero que analices el motivo primero, no te confundas. El motivo puede ser una modificación de un hecho probado o la adición de un nuevo hecho probado o de un nuevo párrafo a un hecho probado ya existente. 
            - Si el primero de los motivos se refiere a la modificación de un hecho probado de la sentencia tienes que escribir lo siguiente:

"Como motivo primero de revisión fáctica, la parte recurrente, interesa la modificación del Hecho Probado [decir qué número de hecho probado es], cuya redacción original es:

[ve a los HECHOS PROBADOS y copia exactamente el hecho probado que se pretende modificar]

La redacción que se propone sería la siguiente:

[ve al RECURSO y copia exactamente la modificación que se pretende]

Para ello, el recurrente se apoya en [hacer una relación de las pruebas en las que el recurrente se basa para proponer la modificación]."

            - Tienes que seguir esta estructura. Tienes prohibido utilizar otra estructura, me enfadaría si utilizaras otra estructura.
            - Si el primero de los motivos se refiere a la adición de un nuevo hecho probado o de un nuevo párrafo a un hecho probado ya existente, tienes que escribir lo siguiente:

"Como motivo primero de revisión fáctica, la parte recurrente, interesa la adición de un nuevo Hecho Probado [decir qué número de hecho probado es]/de un nuevo párrafo al Hecho Probado [decir qué número de hecho probado es], cuya redacción sería la siguiente:

[ve al RECURSO y copia exactamente la adición que se pretende]

Para ello, el recurrente se apoya en [hacer una relación de las pruebas en las que el recurrente se basa para proponer la adición]."

            - Tienes que seguir esta estructura. Tienes prohibido utilizar otra estructura, me enfadaría si utilizaras otra estructura.
            - Take a deep breath and work on this problem step by step."""

# FIN DE ANALISIS DE REVISIONES FÁCTICAS
# ANALISIS DE LAS CENSURAS JURÍDICAS

def generar_respuestas2(prompt_base, numero_de_repeticiones, pdf2):
    ordinales = ["primero", "segundo", "tercero", "cuarto", "quinto", "sexto", "séptimo", "octavo"]
    respuesta_final = ""

    for i in range(numero_de_repeticiones):
        # Reemplazar [numero] por el número de repeticiones y ajustar singular o plural
        if numero_de_repeticiones == 1:
            prompt_con_numero = prompt_base.replace("[numero]", "1 motivo")
        else:
            prompt_con_numero = prompt_base.replace("[numero]", f"{numero_de_repeticiones} motivos")

        # Reemplazar todas las instancias de 'primero' por el ordinal correspondiente
        prompt_modificado = prompt_con_numero.replace("primero", ordinales[i])

        # Generar la frase con los ordinales para la frase 'No analices ni el motivo...'
        frase_ordinales = ", ni el ".join(ordinales[:i])
        if frase_ordinales:
            frase_ordinales = "No analices el motivo " + frase_ordinales + ", ya los analizaste antes. "
        # Insertar la frase en el lugar correcto del prompt modificado

        partes_prompt = prompt_modificado.split("Think step by step.")
        prompt_modificado = partes_prompt[0] + frase_ordinales + "Think step by step." + partes_prompt[1]

        print(prompt_modificado)
        respuesta = get_response4(prompt_modificado, pdf2)
        pyperclip.copy(respuesta)

        # Agregar el prompt modificado a la respuesta final
        respuesta_final += f"{respuesta}\n"
        respuesta_final = reemplazar_motivos(respuesta_final, numero_de_repeticiones)

    return respuesta_final


def analizar_censuras_juridicas():
    progress_bar.grid(row=0, column=2, padx=2, pady=2, sticky="ew")  # Mostras la barra de progreso
    progress_bar.start(5)  # Iniciar la animación de la barra
    pdf2 = leer_pdf(pdf_path_for_text_extraction)
    numero_de_repeticiones = motivos_c  # Puedes cambiar este número según lo necesites
    respuesta_final = generar_respuestas2(promptcensura, numero_de_repeticiones, pdf2)
    progress_bar.stop()  # Detener la animación
    progress_bar.grid_remove()  # Ocultar la barra de progreso

    # Asegúrate de definir la variable 'prompt' fuera de la función, como una variable global o directamente en la función.
    censuras_var.set(respuesta_final)  # Devuelve la respuesta completa para su uso posterior)

promptcensura = """Eres un experto analizador de recursos de suplicación. Quiero que analices el Recurso. Los recursos de suplicación pueden interponerse por el art. 193.a) LRJS, art. 193.b) LRJS o art. 193.c) LRJS.
            - El apartado a) del art. 193 LRJS regula un motivo de recurso que pretende reponer los autos al estado en el que se encontraban en el momento de cometerse una infracción de normas o garantías del procedimiento que haya producido indefensión.
            - El apartado b) del art. 193 LRJS regula un motivo de recurso que pretende revisar los hechos declarados probados, a la vista de las pruebas documentales y periciales practicadas.
            - El apartado c) del art. 193 LRJS regula un motivo de recurso que pretende examinar las infracciones de normas sustantivas o de la jurisprudencia.
            - Tienes que buscar los motivos suplicacionales de la letra c).
            - En el documento hay [numero] de la letra c).
            - Una vez que los encuentres centrate sólo en el motivo primero. Eres un experto contanto motivos. Sólo quiero que analices el motivo primero.
            - Think step by step. 
            - Tienes que hacer un resumen de 200 palabras del motivo primero, de las razones por las cuales el recurrente recurre la sentencia.
            - Sigue la siguiente estructura. Sólo del motivo primero de la letra c):

"Como motivo primero de censura jurídica, la parte recurrente interesa la revocación de la sentencia de instancia al considerar que se ha producido la infracción [decir los artículos que el recurrente dice que se han infringido en el primero motivo de la letra c)], a saber [aquí empieza el resumen]."

            - Tienes que seguir esta estructura. Tienes prohibido utilizar otra estructura, me enfadaría si utilizaras otra estructura.
            - Take a deep breath and work on this problem step by step."""

# FIN DEL ANALISIS DE LA CENSURA JURÍDICA

def on_mousewheel(event):
    left_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

def bind_to_mousewheel(widget):
    widget.bind("<MouseWheel>", on_mousewheel)

# Funciones relacionadas con el portapapeles
def verificar_portapapeles():
    try:
        ultimo_copiado = root.clipboard_get()
        if ultimo_copiado and (ultimo_copiado not in historial_portapapeles):
            historial_portapapeles.append(ultimo_copiado)
    except tk.TclError:
        pass
    root.after(1000, verificar_portapapeles)

def pegar_en_texto(elemento, lugar):
    lugar.insert(tk.END, elemento + '\n')

def truncar_texto(texto, max_ancho):
    return (texto[:max_ancho] + '...') if len(texto) > max_ancho else texto

def mostrar_menu(event, lugar):
    menu = tk.Menu(root, tearoff=0)
    for elemento in historial_portapapeles:
        texto_menu = truncar_texto(elemento, MAX_ANCHO_MENU)
        menu.add_command(label=texto_menu, command=lambda e=elemento: pegar_en_texto(e, lugar))
    try:
        menu.tk_popup(event.x_root, event.y_root)
    finally:
        menu.grab_release()

# Función de validación
def solo_numeros(P):
    # Permite solo si el texto es un número o está vacío (para permitir borrado)
    return P.isdigit() or P == ""

root = tk.Tk()
root.title('Asistente a la Suplicación')
root.geometry("1300x800")  # Establece el tamaño de la ventana a 800x600

# Left frame for inputs
left_frame = ttk.Frame(root, width=200)
left_frame.grid(row=0, column=0, sticky="ns")

# Right frame for text display
right_frame = ttk.Frame(root)
right_frame.grid(row=0, column=1, sticky="nsew")


# Left frame content
left_canvas = tk.Canvas(left_frame, highlightthickness=0)
left_scrollbar = ttk.Scrollbar(left_frame, orient="vertical", command=left_canvas.yview)
left_scrollable_frame = ttk.Frame(left_canvas)
# En Windows, el evento MouseWheel puede ser usado directamente.
left_canvas.bind_all("<MouseWheel>", lambda e: left_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))


# Añadir un padding interno al canvas para centrar los widgets
left_scrollable_frame.bind(
    "<Configure>",
    lambda e: left_canvas.configure(
        scrollregion=left_canvas.bbox("all"),
        width=e.width + 10  # Agregar un padding interno de 10 píxeles a cada lado
    )
)

left_canvas.create_window((0, 0), window=left_scrollable_frame, anchor="nw")
left_canvas.configure(yscrollcommand=left_scrollbar.set)

# Frame que contiene los botones
opciones_frame = ttk.Frame(left_scrollable_frame)
opciones_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=2)

# Configuramos el weight de las columnas del frame para que se expandan adecuadamente
opciones_frame.grid_columnconfigure(0, weight=1)
opciones_frame.grid_columnconfigure(1, weight=1)

# Botón "cargar" dentro del frame
cargar_button = ttk.Button(opciones_frame, text="Cargar Sentencia", command=lambda: ejecutar_en_hilo(cargar_pdf))
cargar_button.grid(row=0, column=0, sticky="ew", padx=2, pady=2)

# Botón "resolver" dentro del frame
recurso_button = ttk.Button(opciones_frame, text="Cargar Recurso", state='disabled', command=lambda: ejecutar_en_hilo(analyze_pdf))
recurso_button.grid(row=0, column=1, sticky="ew", padx=2, pady=2)


# Partido de procedencia checkboxes
partido_var = tk.StringVar()
partido_frame = ttk.LabelFrame(left_scrollable_frame, text="Partido de procedencia")
partido_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=5)

for partido in ["Las Palmas de Gran Canaria", "Arrecife", "Puerto del Rosario", "Galdar"]:
    check = ttk.Checkbutton(partido_frame, text=partido, variable=partido_var, onvalue=partido, offvalue='', command=update_juzgado)
    check.pack(side="top", anchor="w", padx=5, pady=2)

# Juzgado de procedencia dropdown
juzgado_var = tk.StringVar()
juzgado_label = ttk.Label(left_scrollable_frame, text="Juzgado de procedencia")
juzgado_label.grid(row=2, column=0, sticky="ew", padx=10, pady=(5,0))
juzgado_options = ttk.Combobox(left_scrollable_frame, textvariable=juzgado_var)
juzgado_options.grid(row=3, column=0, sticky="ew", padx=10, pady=(0,5))

# Fecha calendar
fecha_label = ttk.Label(left_scrollable_frame, text="Fecha de la sentencia recurrida")
fecha_label.grid(row=4, column=0, sticky="ew", padx=10, pady=(5,0))
fecha_cal = DateEntry(left_scrollable_frame)
fecha_cal.grid(row=5, column=0, sticky="ew", padx=10, pady=(0,5))

# Autos de procedencia
autos_var = tk.StringVar()
autos_year_var = tk.StringVar()
autos_frame = ttk.Frame(left_scrollable_frame)
autos_frame.grid(row=6, column=0, sticky="ew", padx=10, pady=5)

autos_label = ttk.Label(autos_frame, text="Autos")
autos_label.grid(row=0, column=0, columnspan=2, sticky="w")

# Configuración del validador
vcmd = (root.register(solo_numeros), '%P')

autos_entry = ttk.Entry(autos_frame, textvariable=autos_var, validate="key", validatecommand=vcmd)
autos_entry.grid(row=1, column=0, sticky="ew", padx=(0,10))
autos_entry.bind("<Button-3>", lambda event: mostrar_menu(event, autos_entry))
autos_year_options = ttk.Combobox(autos_frame, textvariable=autos_year_var, values=['2021', '2022', '2023'])
autos_year_options.grid(row=1, column=1, sticky="ew")

autos_frame.columnconfigure(0, weight=3)
autos_frame.columnconfigure(1, weight=1)

# Numero de recurso
recurso_var = tk.StringVar()
recurso_year_var = tk.StringVar()
recurso_frame = ttk.Frame(left_scrollable_frame)
recurso_frame.grid(row=7, column=0, sticky="ew", padx=10, pady=5)

recurso_label = ttk.Label(recurso_frame, text="Recurso")
recurso_label.grid(row=0, column=0, columnspan=2, sticky="w")

recurso_entry = ttk.Entry(recurso_frame, textvariable=recurso_var, validate="key", validatecommand=vcmd)
recurso_entry.grid(row=1, column=0, sticky="ew", padx=(0,10))
recurso_entry.bind("<Button-3>", lambda event: mostrar_menu(event, recurso_entry))
recurso_year_options = ttk.Combobox(recurso_frame, textvariable=recurso_year_var, values=['2021', '2022', '2023'])
recurso_year_options.grid(row=1, column=1, sticky="ew")

recurso_frame.columnconfigure(0, weight=3)
recurso_frame.columnconfigure(1, weight=1)

# Recurrente entry

recurrente_var = tk.StringVar()
recurrente_label = ttk.Label(left_scrollable_frame, text="Recurrente")
recurrente_label.grid(row=8, column=0, sticky="ew", padx=10, pady=(5,0))
recurrente_entry = ttk.Entry(left_scrollable_frame, textvariable=recurrente_var)
recurrente_entry.grid(row=9, column=0, sticky="ew", padx=10, pady=(0,5))
recurrente_entry.bind("<Button-3>", lambda event: mostrar_menu(event, recurrente_entry))

# Impugnante

impugnante_var = tk.StringVar()
impugnante_label = ttk.Label(left_scrollable_frame, text="Impugnante")
impugnante_label.grid(row=10, column=0, sticky="ew", padx=10, pady=(5,0))
impugnante_entry = ttk.Entry(left_scrollable_frame, textvariable=impugnante_var)
impugnante_entry.grid(row=11, column=0, sticky="ew", padx=10, pady=(0,5))
impugnante_entry.bind("<Button-3>", lambda event: mostrar_menu(event, impugnante_entry))


def update_sintesis_with_impugnante(*args):
    # Base de la frase para la actualización.
    frase_base_impugnacion = " El recurso fue impugnado por la representación letrada de "

    # Extraer el contenido actual de la síntesis.
    sintesis_text = sintesis_var.get().strip()

    # Obtener el nombre del impugnante del Entry asociado.
    impugnante_name = impugnante_var.get().strip()

    # Si hay un nombre del impugnante, verificamos su presencia y actualizamos el nombre.
    if impugnante_name:
        # Si ya contiene la frase base, actualizamos solo el nombre del impugnante.
        if frase_base_impugnacion.strip() in sintesis_text:
            # Encontrar el comienzo de la frase y actualizar solo esa parte.
            start = sintesis_text.find(frase_base_impugnacion.strip())
            end = sintesis_text.find("\n", start + len(frase_base_impugnacion)) if "\n" in sintesis_text[
                                                                                           start:] else None
            sintesis_text = (sintesis_text[
                             :start].strip() + frase_base_impugnacion + impugnante_name + "." + "\n") if end is None else (
                        sintesis_text[:start] + frase_base_impugnacion + impugnante_name + "." + "\n" + sintesis_text[end:])
        else:
            # Si no, añadimos la frase al final.
            sintesis_text += frase_base_impugnacion + impugnante_name
    else:
        # Si no hay impugnante, eliminamos la frase si existe.
        if frase_base_impugnacion.strip() in sintesis_text:
            # Encontrar y eliminar la línea entera que contiene la frase.
            start = sintesis_text.find(frase_base_impugnacion.strip())
            end = sintesis_text.find("\n", start + len(frase_base_impugnacion)) if "\n" in sintesis_text[
                                                                                           start:] else None
            sintesis_text = (sintesis_text[:start].rstrip()) if end is None else (
                        sintesis_text[:start] + sintesis_text[end + 1:])

    # Actualizamos el Entry de la síntesis con el nuevo o ajustado contenido.
    sintesis_var.set(sintesis_text)

# Asociar la nueva función al cambio de la variable impugnante_var
impugnante_var.trace("w", update_sintesis_with_impugnante)


# Revisiones entry
revisiones_var = tk.StringVar()
revisiones_entry = ttk.Entry(left_scrollable_frame, textvariable=revisiones_var)
revisiones_entry.grid_forget()

# Censuras entry
censuras_var = tk.StringVar()
censuras_entry = ttk.Entry(left_scrollable_frame, textvariable=censuras_var)
censuras_entry.grid_forget()

# Add a new entry widget for 'Sintesis de la litis' after 'Recurrente'
sintesis_var = tk.StringVar()
sintesis_label = ttk.Label(left_scrollable_frame, text="Síntesis de la litis")
sintesis_label.grid(row=12, column=0, sticky="ew", padx=10, pady=(5,0))
sintesis_entry = ttk.Entry(left_scrollable_frame, textvariable=sintesis_var, width=40)
sintesis_entry.grid(row=13, column=0, sticky="ew", padx=10, pady=(0,5))
sintesis_entry.bind("<Button-3>", lambda event: mostrar_menu(event, sintesis_entry))

# Variable del resumen de la sentencia
resumen_caso_var = tk.StringVar()

# Checkbox 'Cuestiones de orden público'
cuestiones_var = tk.BooleanVar(value=False)
cuestiones_check = ttk.Checkbutton(left_scrollable_frame, text="Cuestiones de orden público", variable=cuestiones_var)
cuestiones_check.grid(row=14, column=0, sticky="w", padx=8, pady=2)

# Checkbox 'Inadmisibilidad del recurso'
inadmisibilidad_var = tk.BooleanVar(value=False)
inadmisibilidad_check = ttk.Checkbutton(left_scrollable_frame, text="Inadmisibilidad del recurso", variable=inadmisibilidad_var)
inadmisibilidad_check.grid(row=15, column=0, sticky="w", padx=8, pady=2)

# Frame para los checkboxes que dependen de 'Infracción de las normas procesales'
dynamic_frame = ttk.Frame(left_scrollable_frame)

# Checkbox 'Infracción de las normas procesales'
infraccion_var = tk.BooleanVar(value=False)
infraccion_check = ttk.Checkbutton(left_scrollable_frame, text="Infracción de las normas procesales", variable=infraccion_var)
infraccion_check.grid(row=16, column=0, sticky="w", padx=8, pady=2)


# Se define la función toggle_infraccion más arriba en tu código
def toggle_infraccion():
    if infraccion_var.get():
        # Mostrar el frame con los checkboxes
        dynamic_frame.grid(row=17, column=0, sticky="ew", padx=15, pady=2)
    else:
        # Ocultar el frame con los checkboxes
        dynamic_frame.grid_forget()

infraccion_var.trace("w", lambda *_: toggle_infraccion())

# Definición de los checkboxes dentro del frame 'tipos_frame'
tipos_vars = [tk.BooleanVar(value=False) for _ in range(6)]
tipos_options = [
    "Denegación de prueba",
    "Falta de motivación pidiendo nulidad",
    "Falta de motivación sin pedir nulidad",
    "Incongruencia Extrapetitum",
    "Incongruencia Infrapetitum",
    "Defectuosa grabación"
]

for i, option in enumerate(tipos_options):
    chk = ttk.Checkbutton(dynamic_frame, text=option, variable=tipos_vars[i])
    chk.pack(anchor="w", pady=2)

tipos_options[0] = """	La parte recurrente interesa la nulidad de la Sentencia, al amparo del art. 193.a) LRJS, porque en el acto del juicio oral se le denegó la práctica de prueba propuesta; la parte recurrente formuló protesta ante la denegación de dicha práctica y justificó que era necesaria.
	La parte recurrente plantea que la denegación de dichas pruebas le produce indefensión.
	El derecho a un proceso con todas las garantías, previsto en el artículo 24.2 de la Constitución Española, ha sido interpretado por el Tribunal Constitucional entre otras en su Sentencia 199/1992, de 19 de noviembre, en el sentido de que «la indefensión que prohíbe el artículo 24.1 de la Constitución no nace de la sola y simple infracción de las reglas procesales por parte de los órganos competentes, sino que se produce cuando la vulneración de las normas procesales lleva consigo la prohibición del derecho a la defensa, con el consiguiente perjuicio real y efectivo para los intereses del afectado».
	En cuanto al derecho a la prueba, la Sala Cuarta del Tribunal Supremo en Sentencia de 25 de enero de 2018, (rec. 1648/2016) señala que “la tutela judicial efectiva presupone la garantía de las personas de que pueden acudir al Juez, ser oídas, proponer y practicar pruebas, a una resolución fundada en derecho y a los recursos […]”."""
tipos_options[1] = """	La parte recurrente interesa la nulidad de la Sentencia, al amparo del art. 193.a) LRJS, por falta de motivación.
	El derecho a un proceso con todas las garantías, previsto en el artículo 24.2 de la Constitución Española, ha sido interpretado por el Tribunal Constitucional entre otras en su Sentencia 199/1992, de 19 de noviembre, en el sentido de que «la indefensión que prohíbe el artículo 24.1 de la Constitución no nace de la sola y simple infracción de las reglas procesales por parte de los órganos competentes, sino que se produce cuando la vulneración de las normas procesales lleva consigo la prohibición del derecho a la defensa, con el consiguiente perjuicio real y efectivo para los intereses del afectado».
	La doctrina del TC respecto a la necesidad de motivación de las sentencias como se razona en la STC 80/2000, de 27 de marzo, establece que: «el requisito de motivación de las sentencias, como parte integrante del derecho a la tutela judicial efectiva, no impone que la resolución ofrezca una exhaustiva descripción del proceso intelectual llevado a cabo por el juzgador para resolver, ni una pormenorizada respuesta a todas las alegaciones de las partes, ni un determinado alcance o entidad en el razonamiento empleado, ni siquiera la corrección jurídica interna de la fundamentación empleada; basta que la argumentación vertida, con independencia de su parquedad o concentración, cumpla la doble finalidad de exteriorizar el motivo de la decisión, su ‘ratio decidendi’ (STC 5/1995, de 10 de enero), excluyente de un mero voluntarismo selectivo o de la pura arbitrariedad, y que permita su eventual revisión jurisdiccional a través del efectivo ejercicio de los recursos establecidos, como aquí sucede con este recurso de amparo».
	La doctrina constitucional y también la jurisprudencia del Tribunal Supremo vienen siendo constantes en señalar que la motivación debe cumplir con dos requisitos: el primero es que debe ser suficiente, lo que no quiere decir exhaustiva, y en segundo lugar fundamentada o argumentada en derecho y no en otros criterios del juez/a o tribunal."""
tipos_options[2] = """	La parte recurrente interesa invoca el art. 193.a) LRJS, por falta de motivación. Sin embargo, no plantea la nulidad de la sentencia.
	El planteamiento resulta interesante, pero dándose la circunstancia de que no se solicita la nulidad de la sentencia, entendemos que debería haberse articulado al amparo de la letra c) del artículo 193 LRJS y por tanto en este punto carece de sentido entrar a analizar más profundamente lo planteado, sin perjuicio de su análisis más adelante cuando tratemos los motivos impugnatorios del art. 193.c) LRJS."""
tipos_options[3] = """	La parte recurrente interesa la nulidad de la Sentencia, al amparo del art. 193.a) LRJS, por incongruencia "extra petitum".
	Sobre la figura de la congruencia, se ha ocupado el Tribunal Constitucional en numerosas ocasiones. Así, desde la STC 20/1982, de 5 de mayo, viene señalando que “la misma se mide por el ajuste o adecuación entre lo resuelto y los términos en que las partes han formulado sus pretensiones y peticiones, de manera tal que no puede la sentencia otorgar más de lo que se hubiera pedido en la demanda, ni menos de lo que hubiera sido admitido por el demandado, ni otorgar otra cosa diferente, que no hubiera sido pretendida”. 
	El Tribunal Supremo, en ST de 22 de diciembre de 2016 señala que la congruencia constituye “un ajuste sustancial entre lo pedido y lo resuelto que, por tanto, no exige del fallo una conformidad literal y rígida con las peticiones de las partes, sino racional y flexible”.
	La incongruencia "extra petitum", es la que se origina cuando se resuelve sobre cuestiones distintas y ajenas a lo solicitado por las partes, lo que implica una invasión frontal del derecho de defensa contradictorio, privando a los litigantes de la facultad de alegar lo que a su derecho proceda o lo que estimen conveniente a sus legítimos intereses."""
tipos_options[4] = """	La parte recurrente interesa la nulidad de la Sentencia, al amparo del art. 193.a) LRJS, por incongruencia omisiva.
	Sobre la figura de la congruencia, se ha ocupado el Tribunal Constitucional en numerosas ocasiones. Así, desde la STC 20/1982, de 5 de mayo, viene señalando que “la misma se mide por el ajuste o adecuación entre lo resuelto y los términos en que las partes han formulado sus pretensiones y peticiones, de manera tal que no puede la sentencia otorgar más de lo que se hubiera pedido en la demanda, ni menos de lo que hubiera sido admitido por el demandado, ni otorgar otra cosa diferente, que no hubiera sido pretendida”.
	El Tribunal Supremo, en ST de 22 de diciembre de 2016 señala que la congruencia constituye “un ajuste sustancial entre lo pedido y lo resuelto que, por tanto, no exige del fallo una conformidad literal y rígida con las peticiones de las partes, sino racional y flexible”.
	La incongruencia "infra petitum", es la que se origina por falta de respuesta judicial razonada a alguna de las cuestiones o elementos esenciales de la pretensión sometidos por las partes a la consideración del Juzgador/a, y cuyo conocimiento y decisión sea trascendente para fijar el fallo.
	No todos los casos de ausencia de respuesta judicial expresa producen una indefensión constitucionalmente relevante, debiendo ponderarse las circunstancias concurrentes en cada caso para determinar:
		a)	Si el silencio de la resolución judicial constituye una auténtica lesión del art. 24.1 CE o si, por el contrario, puede razonablemente interpretarse como una desestimación tácita.
		b)	Si efectivamente se ha planteado la cuestión cuyo conocimiento se afirma eludido por el Tribunal.
		c)	Si la incongruencia omisiva apreciada causó un efectivo perjuicio de los derechos de defensa de quien se queja en amparo."""
tipos_options[5] = """	La parte recurrente interesa la nulidad de la Sentencia, al amparo del art. 193.a) LRJS, por defectos en la grabación de la vista.
	La defectuosa grabación del juicio oral no conlleva necesariamente la anulación de las actuaciones de instancia cuando no se ha causado indefensión a la parte recurrente. En este sentido, el Tribunal Constitucional entre otras en su Sentencia 199/1992 señala que «la indefensión que prohíbe el artículo 24.1 de la Constitución no nace de la sola y simple infracción de las reglas procesales por parte de los órganos competentes, sino que se produce cuando la vulneración de las normas procesales lleva consigo la prohibición del derecho a la defensa, con el consiguiente perjuicio real y efectivo para los intereses del afectado».
	Igualmente, la STC n.º 4/2004, de 14 de enero, explica que la destrucción del acta de juicio no es en sí una causa de nulidad del juicio mismo si ese juicio materialmente se celebró con todas las garantías y en él las partes pudieron ejercer sus derechos de alegación y prueba sin limitación."""

# Añadir nuevos checkboxes debajo de 'Infracción de normas procesales'
revision_var = tk.BooleanVar(value=False)
revision_check = ttk.Checkbutton(left_scrollable_frame, text="Revisión de hechos probados", variable=revision_var)
revision_check.grid(row=18, column=0, sticky="w", padx=8, pady=2)

# Botón para 'Analizar revisiones fácticas'
analizar_revisiones_button = ttk.Button(left_scrollable_frame, text="Analizar revisiones fácticas", state='disabled', command=lambda: ejecutar_en_hilo(analizar_revisiones_facticas))
analizar_revisiones_button.grid(row=19, column=0, padx=10, pady=10, sticky="ew")
analizar_revisiones_button.grid_remove()  # Ocultar por defecto

def toggle_revision_button():
    if revision_var.get():
        analizar_revisiones_button.grid()  # Mostrar el botón
    else:
        analizar_revisiones_button.grid_remove()  # Ocultar el botón

# Vincular la función toggle_revision_button al cambio de estado del checkbox 'Revisión de hechos probados'
revision_var.trace("w", lambda *args: toggle_revision_button())

# Llamar a la función una vez para configurar el estado inicial del botón
toggle_revision_button()

sustantivas_var = tk.BooleanVar(value=False)
sustantivas_check = ttk.Checkbutton(left_scrollable_frame, text="Infracción de normas sustantivas", variable=sustantivas_var)
sustantivas_check.grid(row=20, column=0, sticky="w", padx=8, pady=2)

# Cuadros de texto para 'Infracción de normas sustantivas'
articulos_var = tk.StringVar()
jurisprudencia_var = tk.StringVar()

# Artículos infringidos
articulos_label = ttk.Label(left_scrollable_frame, text="Artículos infringidos")
articulos_label.grid(row=22, column=0, sticky="ew", padx=10, pady=(5,0))
articulos_entry = ttk.Entry(left_scrollable_frame, textvariable=articulos_var, width=40)
articulos_entry.grid(row=23, column=0, sticky="ew", padx=10, pady=(0,5))
articulos_entry.bind("<Button-3>", lambda event: mostrar_menu(event, articulos_entry))

# Jurisprudencia infringida
jurisprudencia_label = ttk.Label(left_scrollable_frame, text="Jurisprudencia infringida")
jurisprudencia_label.grid(row=24, column=0, sticky="ew", padx=10, pady=(5,0))
jurisprudencia_entry = ttk.Entry(left_scrollable_frame, textvariable=jurisprudencia_var, width=40)
jurisprudencia_entry.grid(row=25, column=0, sticky="ew", padx=10, pady=(0,5))
jurisprudencia_entry.bind("<Button-3>", lambda event: mostrar_menu(event, jurisprudencia_entry))

# Botón para 'Analizar revisiones fácticas'
analizar_censuras_button = ttk.Button(left_scrollable_frame, text="Analizar censuras jurídicas", state='disabled', command=lambda: ejecutar_en_hilo(analizar_censuras_juridicas))
analizar_censuras_button.grid(row=26, column=0, padx=10, pady=10, sticky="ew")
analizar_censuras_button.grid_remove()  # Ocultar el botón

def toggle_censuras():
    if sustantivas_var.get():
        articulos_label.grid()
        articulos_entry.grid()
        jurisprudencia_label.grid()
        jurisprudencia_entry.grid()
        analizar_censuras_button.grid()  # Mostrar el botón
    else:
        articulos_label.grid_remove()
        articulos_entry.grid_remove()
        jurisprudencia_label.grid_remove()
        jurisprudencia_entry.grid_remove()
        analizar_censuras_button.grid_remove()  # Ocultar el botón

# Vincular la función toggle_revision_button al cambio de estado del checkbox 'Revisión de hechos probados'
sustantivas_var.trace("w", lambda *args: toggle_censuras())

# Llamar a la función una vez para configurar el estado inicial del botón
toggle_censuras()

# Frame "Costas, depósitos y consignaciones"
costas_frame = ttk.LabelFrame(left_scrollable_frame, text="Costas, depósitos y consignaciones", borderwidth=1, relief="groove")
costas_frame.grid(row=27, column=0, sticky="ew", padx=10, pady=5)

# Variables para los radiobuttons
costas_var = tk.StringVar(value=None)

textos_costas = {
    "Estimación": "La estimación del recurso caca...",
    "Estimación parcial": "La estimación parcial del recurso...",
    "Desestimación con costas": "La desestimación implica la imposición de [XXX] euros de costas",
    "Desestimación sin costas": "La desestimación no implica la imposición de costas por ser [razón]"
}

# Usamos los valores cortos para los Radiobuttons
for option in textos_costas.keys():
    ttk.Radiobutton(costas_frame, text=option, value=option, variable=costas_var, command=update_text).pack(anchor="w", padx=5, pady=2)


# Subframe y widgets para "Estimación"
estimacion_frame = ttk.LabelFrame(left_scrollable_frame, text="Depósitos y Aseguramientos", borderwidth=2, relief="groove")
estimacion_frame.grid(row=28, column=0, sticky="ew", padx=10, pady=5)
estimacion_frame.grid_remove()  # Ocultar por defecto
trabajador_empresa_var = tk.StringVar(value="Trabajador")
ttk.Radiobutton(estimacion_frame, text="Trabajador", variable=trabajador_empresa_var, value="Trabajador").pack(anchor="w", pady=2)
ttk.Radiobutton(estimacion_frame, text="Empresa", variable=trabajador_empresa_var, value="Empresa").pack(anchor="w", pady=2)
ttk.Label(estimacion_frame, text="Revocación").pack(anchor="w", padx=10, pady=(5,0))
revocacion_var = tk.StringVar()
revocacion_entry = ttk.Entry(estimacion_frame, textvariable=revocacion_var, width=40)
revocacion_entry.pack(anchor="w", padx=10, pady=(0,5))
revocacion_entry.bind("<Button-3>", lambda event: mostrar_menu(event, revocacion_entry))

# Subframe y widgets para "Estimación parcial"
estimacion_parcial_frame = ttk.LabelFrame(left_scrollable_frame, text="Depósitos y Aseguramientos", borderwidth=2, relief="groove")
estimacion_parcial_frame.grid(row=29, column=0, sticky="ew", padx=10, pady=5)
estimacion_parcial_frame.grid_remove()  # Ocultar por defecto
ttk.Radiobutton(estimacion_parcial_frame, text="Trabajador", variable=trabajador_empresa_var, value="Trabajador").pack(anchor="w", pady=2)
ttk.Radiobutton(estimacion_parcial_frame, text="Empresa", variable=trabajador_empresa_var, value="Empresa").pack(anchor="w", pady=2)
ttk.Label(estimacion_parcial_frame, text="Revocación parcial").pack(anchor="w", padx=10, pady=(5,0))
revocacion_parcial_var = tk.StringVar()
revocacion_parcial_entry = ttk.Entry(estimacion_parcial_frame, textvariable=revocacion_parcial_var, width=40)
revocacion_parcial_entry.pack(anchor="w", padx=10, pady=(0,5))
revocacion_parcial_entry.bind("<Button-3>", lambda event: mostrar_menu(event, revocacion_parcial_entry))

# Widgets para "Desestimación con costas"
cuantia_costas_var = tk.StringVar(value="800")
cuantia_costas_label = ttk.Label(left_scrollable_frame, text="Cuantía de la condena en costas")
cuantia_costas_label.grid(row=30, column=0, sticky="ew", padx=10, pady=(5,0))
cuantia_costas_label.grid_remove()  # Ocultar por defecto
cuantia_costas_entry = ttk.Entry(left_scrollable_frame, textvariable=cuantia_costas_var, width=40)
cuantia_costas_entry.grid(row=31, column=0, sticky="ew", padx=10, pady=(0,5))
cuantia_costas_entry.grid_remove()  # Ocultar por defecto
cuantia_costas_entry.bind("<Button-3>", lambda event: mostrar_menu(event, cuantia_costas_entry))

# Subframe y widgets para "Desestimación sin costas"
razon_var = tk.StringVar(value="Beneficiario de justicia gratuita")
razones_frame = ttk.LabelFrame(left_scrollable_frame, text="Razones para su no imposición", borderwidth=2, relief="groove")
razones_frame.grid(row=32, column=0, sticky="ew", padx=10, pady=5)
razones_frame.grid_remove()  # Ocultar por defecto
razones_options = [
    "Beneficiario de justicia gratuita",
    "Sindicato",
    "Funcionario público",
    "Personal laboral"
]
for razon in razones_options:
    rb = ttk.Radiobutton(razones_frame, text=razon, variable=razon_var, value=razon)
    rb.pack(anchor="w", pady=2)


def toggle_costas_frames():
    """Función para mostrar/ocultar frames según la selección de costas."""
    # Primero, "olvidamos" todos los frames/subframes
    estimacion_frame.grid_remove()
    estimacion_parcial_frame.grid_remove()
    cuantia_costas_label.grid_remove()
    cuantia_costas_entry.grid_remove()
    razones_frame.grid_remove()

    # Luego, mostramos el frame/subframe correspondiente
    if costas_var.get() == "Estimación":
        estimacion_frame.grid()
    elif costas_var.get() == "Estimación parcial":
        estimacion_parcial_frame.grid()
    elif costas_var.get() == "Desestimación con costas":
        cuantia_costas_label.grid()
        cuantia_costas_entry.grid()
    elif costas_var.get() == "Desestimación sin costas":
        razones_frame.grid()

# Vincular la función toggle_costas_frames al cambio de la variable costas_var
costas_var.trace("w", lambda *_: toggle_costas_frames())

# Iniciar la función una vez para mostrar el frame/subframe inicial
toggle_costas_frames()

# Asegurarse de que update_text se llama cuando cambian los valores de los nuevos checkboxes
cuestiones_var.trace("w", update_text)
inadmisibilidad_var.trace("w", update_text)
infraccion_var.trace("w", update_text)
sintesis_var.trace("w", update_text)
revision_var.trace("w", update_text)
sustantivas_var.trace("w", update_text)
articulos_var.trace("w", update_text)
jurisprudencia_var.trace("w", update_text)
trabajador_empresa_var.trace("w", update_text)
revocacion_var.trace("w", update_text)
revocacion_parcial_var.trace("w", update_text)
cuantia_costas_var.trace("w", update_text)
razon_var.trace("w", update_text)


for var in tipos_vars:
    var.trace("w", update_text)


# Buttons
button_frame = ttk.Frame(left_scrollable_frame)
button_frame.grid(row=100, column=0, sticky="ew", padx=10, pady=2)

# Configuramos el grid del button_frame para alinear mejor los widgets
button_frame.grid_columnconfigure(0, weight=1)
button_frame.grid_columnconfigure(1, weight=1)


# Botón "Guardar"
save_button = ttk.Button(button_frame, text="Guardar", command=save_to_doc)
# Colocamos el botón "Guardar" en la primera columna
save_button.grid(row=0, column=0, padx=2, pady=2, sticky="ew")

# Crear atajo de teclado para crear texto

root.bind('<Control-Return>', lambda event: save_button.invoke())

# Botón "Limpiar"
clear_button = ttk.Button(button_frame, text="Limpiar", command=clear_all)
# Colocamos el botón "Limpiar" en la segunda columna
clear_button.grid(row=0, column=1, padx=2, pady=2, sticky="ew")

# Crear atajo de teclado para limpiar

root.bind('<Control-BackSpace>', lambda event: clear_button.invoke())

# Barra de progreso
progress_bar = Progressbar(button_frame, length=100, mode='indeterminate')
# Colocamos la barra de progreso en la tercera columna pero inicialmente no la muestres
progress_bar.grid(row=0, column=2, padx=3, pady=2, sticky="ew")
progress_bar.grid_remove()  # Inicia con la barra de progreso oculta

def start_progress_bar():
    """Función para comenzar la animación de la barra de progreso."""
    progress_bar.pack()
    progress_bar.start()

def stop_progress_bar():
    """Función para detener y ocultar la barra de progreso."""
    progress_bar.stop()
    progress_bar.pack_forget()

# Right text widget
right_text = tk.Text(right_frame, wrap="word", padx=10, pady=10, font=("SegoeUI", 9))
right_text.grid(row=0, column=0, sticky="nsew")
right_text.bind("<Button-3>", lambda event: mostrar_menu(event, right_text))
right_frame.grid_rowconfigure(0, weight=1)
right_frame.grid_columnconfigure(0, weight=1)


# Trigger updates
partido_var.trace("w", update_text)
juzgado_var.trace("w", update_text)
autos_var.trace("w", update_text)
autos_year_var.trace("w", update_text)
recurrente_var.trace("w", update_text)
revisiones_var.trace("w", update_text)
censuras_var.trace("w", update_text)
impugnante_var.trace("w", update_text)

for widget in left_scrollable_frame.winfo_children():
    bind_to_mousewheel(widget)

# Ajustamos el embalaje del Canvas y el Scrollbar para reducir el espacio entre los widgets y el scrollbar.
left_scrollbar.pack(side="right", fill="y", expand=False)
left_canvas.pack(side="left", fill="both", expand=True, padx=(10, 0), pady=0)  # Agregar un padding de 10 píxeles a la izquierda

root.grid_columnconfigure(1, weight=3)
root.grid_rowconfigure(0, weight=1)

# Centra la ventana en la pantalla
center_window(root)

# Imagen de inicio # Cargar imagen
image_path = "login2.png"  # Asegúrate de que este sea el nombre de tu archivo y que esté en el mismo directorio que tu script
image = PhotoImage(file=image_path)


# Crear un label y mostrar la imagen
image_label = tk.Label(right_frame, image=image, bg='white')
image_label.grid(row=0, column=0, sticky="nsew")  # Usa grid para colocar el label en la ventana principal
right_frame.grid_rowconfigure(0, weight=1)
right_frame.grid_columnconfigure(0, weight=1)

# Función para ocultar la imagen
def hide_image():
    image_label.grid_forget()  # Eliminar el label de la ventana con grid_forget()

# Establecer un temporizador para llamar a hide_image después de 5 segundos (5000 milisegundos)
root.after(5000, hide_image)

# Inicia la verificación del portapapeles
verificar_portapapeles()

root.mainloop()