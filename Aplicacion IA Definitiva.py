from tkcalendar import DateEntry
import datetime
import PyPDF2
import itertools
import os
import tkinter as tk
from tkinter import filedialog, messagebox, PhotoImage, ttk
from tkinter.ttk import Progressbar
from PyPDF2 import PdfReader
import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import threading
import subprocess
import pyperclip
import re
import json


archivo_conteo = 'conteo_uso_opciones.json'

# Inicializar el diccionario de conteo
if os.path.exists(archivo_conteo):
    with open(archivo_conteo, 'r') as archivo:
        conteo_uso_opciones = json.load(archivo)
else:
    conteo_uso_opciones = {}


def guardar_conteo():
    with open(archivo_conteo, 'w') as archivo:
        json.dump(conteo_uso_opciones, archivo)


# Historial del portapapeles y configuración
historial_portapapeles = []
MAX_ANCHO_MENU = 30  # Máximo número de caracteres por elemento en el menú

# FUNCIONES PARA LA CREACIÓN DE LA VENTANA DE RESOLUCIÓN DE LA REVISIÓN FÁCTICA

reglas = open('revisor_factico.txt', 'r', encoding="utf8")


# Función para que la ventana aparezca en el centro de la pantalla
def center_window_revisor(root):
    # Actualiza la ventana para asegurar que tkinter tenga las dimensiones correctas
    root.update_idletasks()

    # Obtiene las dimensiones de la pantalla
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()

    # Obtiene las dimensiones de la ventana
    window_width = root.winfo_width()
    window_height = root.winfo_height()

    # Calcula las coordenadas x e y para centrar la ventana
    x = (screen_width // 2) - (window_width // 2)
    y = (screen_height // 2) - (window_height // 2)

    # Establece la posición de la ventana
    root.geometry(f"+{x}+{y}")


# Creamos funciones para preparar el texto

def eliminar_salto_inicial(texto):
    while texto[0] == "\n":
        texto = texto[1:]
    return texto


def borrar_saltos_linea_vacios(texto):
    return texto.replace("\n\n", "\n")


def tabular_parrafos(texto):
    parrafos = texto.split("\n")
    parrafos_tabular = []
    for parrafo in parrafos:
        parrafos_tabular.append("\t" + parrafo)
    return "\n".join(parrafos_tabular)


def preparar_texto(texto):
    preparado1 = eliminar_salto_inicial(texto)
    preparado2 = borrar_saltos_linea_vacios(preparado1)
    preparado3 = tabular_parrafos(preparado2)
    return preparado3


def get_response_revision(prompt):
    response = openai.ChatCompletion.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system",
             "content": reglas.read() + "####\nEstas son todas las reglas que tienes que seguir. Piensa paso a paso, analiza el texto del recurso, las palabras usadas en el texto alternativo propuesto, y en los documentos que se proponen.\nNunca menciones que estás siguiendo 'reglas', sino en todo caso señalar la jurisprudencia de la Sala Cuarta del Tribunal Supremo o 'de conformidad con el art. 193.b) LRJS'.\nEste es el recurso que tienes que resolver:"},
            {"role": "user", "content": prompt},
        ]
    )
    pyperclip.copy(response['choices'][0]['message']['content'])
    return response['choices'][0]['message']['content']


def analizar_texto():
    try:
        # Obtener el texto seleccionado del cuadro de texto
        selected_text = right_text.get(tk.SEL_FIRST, tk.SEL_LAST)
        abrir_nueva_ventana(selected_text)
    except tk.TclError:
        print("No hay texto seleccionado")


def abrir_nueva_ventana(texto_seleccionado):
    def thread_safe_valorar_texto():
        valorar_button.config(state=tk.DISABLED)  # Deshabilita el botón mientras se ejecuta la función
        valorar_texto()
        valorar_button.config(state=tk.NORMAL)  # Habilita el botón una vez que finaliza la función

    def valorar_texto():
        input_text = left_text.get(1.0, tk.END).strip()
        response = get_response_revision(input_text)

        # Mostrar la respuesta en el cuadro de texto derecho
        right_text.delete(1.0, tk.END)
        right_text.insert(tk.END, response)
        root.clipboard_clear()
        root.clipboard_append(preparar_texto(response))

    def toggle_instructions_left():
        if instructions_label_left.winfo_viewable():
            instructions_label_left.grid_remove()
        else:
            instructions_label_left.grid(row=1, column=1, sticky=tk.N)

    def toggle_instructions_right():
        if instructions_label_right.winfo_viewable():
            instructions_label_right.grid_remove()
        else:
            instructions_label_right.grid(row=1, column=2, sticky=tk.N)

    def toggle_both_instructions():
        toggle_instructions_left()
        toggle_instructions_right()

    def abrir_ventana_feedback():
        def enviar_feedback():
            feedback = cuadro_feedback.get("1.0", "end-1c")
            texto_original = left_text.get("1.0", "end-1c")
            regla = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system",
                     "content": "Eres una IA que asume el rol de juez que resuelve recursos de suplicación. El siguiente texto es una corrección ante un recurso que se te ha planteado y lo has resuelto mal. La corrección corresponde al usuario, que ha considerado que el resultado que tú has dado es incorrecto. Estudia cuál es el recurso y la corrección del usuario, extrae de este texto una regla clara, completa y precisa que puedas utilizar a partir de ahora para resolver los recursos. Escríbela de forma que tú la comprendas y puedas evitar cometer el mismo error la próxima vez. Esta regla se incluirá dentro de un archivo donde están el resto de reglas que utilizarás para resolver los recursos. La estructura es 'Regla:' y la explicación. Escribe la regla a partir del siguiente texto:"},
                    {"role": "user",
                     "content": f"Petición de revisión: {texto_original}\n Solución resumida: {feedback}"},
                ]
            )
            with open("revisor_factico.txt", "a", encoding='utf-8') as f:
                f.write(borrar_saltos_linea_vacios(
                    eliminar_salto_inicial(regla['choices'][0]['message']['content'])) + "\n##\n")
            ventana_feedback.destroy()

        ventana_feedback = tk.Toplevel(root)
        ventana_feedback.title("Retroalimentación")
        ventana_feedback['bg'] = "#ffffff"
        cuadro_feedback = tk.Text(ventana_feedback, wrap=tk.WORD, width=45, height=10, font=("SegoeUI", 9),
                                  relief='groove', bd=2)
        cuadro_feedback.pack(expand=True, fill=tk.BOTH, padx=10)
        boton_cerrar = tk.Button(ventana_feedback, text="Cerrar", command=enviar_feedback, bg="white",
                                 activebackground="light grey", relief="groove")
        boton_cerrar.pack(pady=10, padx=10)

    ventana = tk.Toplevel(root)
    ventana.title("Revisor fáctico suplicacional")
    ventana['bg'] = "#ffffff"
    ventana.resizable(width=False, height=False)

    # Aquí puedes agregar todos los elementos de la interfaz de usuario que describiste en tu segunda ventana
    frame = tk.Frame(ventana, padx="10", pady="10")
    frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
    frame.config(bg="white")

    label_titulo = tk.Label(ventana, text="Revisor fáctico suplicacional", background="#ffffff", font=45)
    label_titulo.grid(row=0, column=0, pady=10, columnspan=3)

    instructions_label_left = tk.Label(frame,
                                       text="Este es el cuadro en el que tendrá que introducir la revisión fáctica interesada en el recurso, introduciendo el texto de la propuesta del recurrente tal y como se plantea.",
                                       bg="white", wraplength=310)
    instructions_label_left.grid(row=1, column=1, sticky=tk.N)
    instructions_label_left.grid_remove()

    instructions_label_right = tk.Label(frame,
                                        text="Este es el cuadro en el que se mostrará el Dictamen de la IA, sobre la revisión fáctica. El texto se copiará automáticamente en el portapapeles.",
                                        bg="white", wraplength=330)
    instructions_label_right.grid(row=1, column=2, sticky=tk.N)
    instructions_label_right.grid_remove()
    # Por ahora, solo agregaré el cuadro de texto izquierdo donde se colocará el texto seleccionado

    left_text = tk.Text(frame, wrap=tk.WORD, width=45, height=26, font=("SegoeUI", 9), relief='groove', bd=2)
    left_text.grid(row=0, column=1, padx=(0, 10), pady=(0, 10))

    # Colocar el texto seleccionado en el cuadro de texto izquierdo
    left_text.insert(tk.END, texto_seleccionado)

    # Agrega aquí el resto de los componentes de tu interfaz de usuario según tu necesidad
    right_text = tk.Text(frame, wrap=tk.WORD, width=45, height=26, font=("SegoeUI", 9), relief='groove', bd=2)
    right_text.grid(row=0, column=2, padx=(10, 0), pady=(0, 10))

    valorar_button = tk.Button(frame, text="Valorar",
                               command=lambda: threading.Thread(target=thread_safe_valorar_texto).start(),
                               bg="white", activebackground="light grey", relief="groove")
    valorar_button.grid(row=2, column=1, columnspan=2, pady=10)

    toggle_instructions_button_left = tk.Button(frame, bitmap="info", command=toggle_both_instructions, bg="white",
                                                activebackground="white", width="20", relief="groove", bd=0)
    toggle_instructions_button_left.grid(row=2, column=1, sticky=(tk.W))

    toggle_instructions_button_right = tk.Button(frame, bitmap="warning", command=abrir_ventana_feedback, bg="white",
                                                 activebackground="white", width="20", relief="groove", bd=0)
    toggle_instructions_button_right.grid(row=2, column=2, sticky=(tk.E))

    ventana.bind('<Control-Return>', lambda event: valorar_button.invoke())
    ventana.bind('<Control-BackSpace>', lambda event: right_text.delete('1.0', tk.END))

    # Centrar la ventana en la pantalla (puedes usar la función 'center_window' que ya tienes)
    center_window_revisor(ventana)

    ventana.mainloop()


def menu_contextual_resolver_revision_factica(event):
    try:
        # Mostrar el menú contextual solo si hay texto seleccionado
        right_text.selection_get()
        menu_contextual.post(event.x_root, event.y_root)
    except tk.TclError:
        pass


# FIN DEL MENU PARA RESOLVER REVISIONES FÁCTICAS

# FUNCIONES PARA EXTRAER TEXTO CONCRETO
def extract_text_between_phrases(text, start_phrase, end_phrase):
    """
    Extrae y devuelve el texto que se encuentra entre dos frases específicas dentro de una cadena de texto.
    Si alguna de las frases no se encuentra o el orden es incorrecto, devuelve un mensaje indicándolo.
    """
    try:
        start_index = text.find(start_phrase)
        end_index = text.find(end_phrase, start_index)

        if start_index != -1 and end_index != -1 and start_index < end_index:
            return text[start_index + len(start_phrase):end_index].strip()
        else:
            return "No se encontró el texto especificado o el orden de las frases es incorrecto."
    except Exception as e:
        return f"Error al procesar el texto: {e}"

#crear hilos
def ejecutar_en_hilo(funcion):
    """Ejecuta una función dada en un hilo separado."""
    hilo = threading.Thread(target=funcion)
    hilo.start()
    return hilo

#Centrar la ventana
def center_window(root):
    try:
        root.update_idletasks()
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        window_width = root.winfo_width()
        window_height = root.winfo_height()
        x = (screen_width // 2) - (window_width // 2)
        y = (screen_height // 2) - (window_height // 2)
        y -= 34
        root.geometry(f"+{x}+{y}")
    except Exception as e:
        print(f"Error al centrar la ventana: {e}")


# Credenciales de OpenAI API
openai.api_key = os.getenv('OPENAI_API_KEY')

#Preparar texto
def borrar_saltos_linea_vacios(texto):
    return texto.replace("\n\n", "\n")


def tabular_parrafos(texto):
    parrafos = texto.split("\n")
    parrafos_tabular = []
    for parrafo in parrafos:
        parrafos_tabular.append("\t" + parrafo)
    return "\n".join(parrafos_tabular)


def texto_preparado_para_procesador(texto):
    preparado1 = borrar_saltos_linea_vacios(texto)
    preparado2 = tabular_parrafos(preparado1)
    return preparado2


# Definir una función para generar rangos de números como cadenas
def rango_numerico_como_cadenas(inicio, fin):
    return [str(i) for i in range(inicio, fin + 1)]

# Partido judiciales
partido_juzgado_map = {
    "Galdar": ['1'],
    "Arrecife": ['1', '2'],
    "Puerto del Rosario": ['3', '4'],
    "Las Palmas de Gran Canaria": rango_numerico_como_cadenas(1, 11)
}

def word_to_number(word):
    word_number_mapping = {
        "uno": 1, "dos": 2, "tres": 3, "cuatro": 4, "cinco": 5,
        "seis": 6, "siete": 7, "ocho": 8, "nueve": 9,
        "diez": 10, "once": 11, "doce": 12
    }
    return word_number_mapping.get(word.lower(), word)

def get_response(prompt, model):
    response = openai.ChatCompletion.create(
        model=model,
        messages=[
            {"role": "system", "content": prompt},
        ]
    )
    return response['choices'][0]['message']['content']


def extract_data_from_response(response):
    # Uso de expresiones regulares para encontrar todos los textos entre corchetes
    data_matches = re.findall(r'\[(.*?)\]', response)

    if len(data_matches) == 6:  # Verificamos que hay 5 datos como esperado
        numero_procedimiento = data_matches[0]
        fecha_sentencia = data_matches[1]
        numero_juzgado = data_matches[2]
        nombre_demandante = data_matches[3]
        nombre_demandado = data_matches[4]
        partido_judicial = data_matches[5]

        return numero_procedimiento, fecha_sentencia, numero_juzgado, nombre_demandante, nombre_demandado, partido_judicial
    else:
        print("Error: El formato de la respuesta no es el esperado.")
        return None

def extract_full_text_from_pdf(reader):
    return ''.join(page.extract_text() for page in reader.pages)

def extract_first_page_text(reader):
    return reader.pages[0].extract_text()

def extract_info_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        full_text = extract_full_text_from_pdf(reader)
        extracted_text = extract_text_between_phrases(full_text, "DERECHO", "FALLO")

        prompt_for_summary = create_summary_prompt(extracted_text)
        summary_response = get_response(prompt_for_summary, "gpt-4-1106-preview")
        print(summary_response)

        first_page_text = extract_first_page_text(reader)
        combined_prompt = create_data_extraction_prompt(first_page_text)
        response = get_response(combined_prompt, "gpt-3.5-turbo")

        response = convert_response_to_numbers(response)
        extracted_data = extract_data_from_response(response)

        if extracted_data:
            display_extracted_data(extracted_data)
            update_ui_components(extracted_data, summary_response)

def create_summary_prompt(extracted_text):
    return (
        f"Eres un Juez experto en Derecho Laboral y de la Seguridad Social, actua como tal, tienes que hacer un resumen de 370 palabras del sentido de una sentencia y las razones para estimar o desestimar. Habla siempre en pasado. Nunca hables de la posibilidad de recurso. No digas si ha posibilidad de recurso. Utiliza a lo largo del texto expresiones como 'la sentencia de instancia/la resolución combatida/el pronunciamiento impugnado   apreció/consideró probado/resolvió/entendió'. Empieza el resumen con la expresión 'La sentencia de instancia estimaba/desestimaba la demanda'. La sentencia es la siguiente:\n: {extracted_text}"
    )

def create_data_extraction_prompt(first_page_text):
    return (
        f"A partir del siguiente texto: '{first_page_text}':\n"
        "- Indica el número de procedimiento sin ceros al principio.\n"
        "- Proporciona la fecha de la sentencia, con este formato '%d/%m/%Y'.\n"
        "- Menciona el número del juzgado en formato número, 1...12.\n"
        "- Lista el o los nombres del demandante (si hay más de uno, sepáralos con 'y').\n"
        "- Especifica el nombre del demandado (excluye 'el FOGASA' o 'Fondo de Garantía Salarial').\n"
        "- Especifíca el nombre del partido judicial, puede ser: Las Palmas de Gran Canaria, Arrecife, Galdar o Puerto del Rosario.\n"
        "- Sólo quiero que pongas los datos con el siguiente formato [número del procedimiento sin ceros al principio] [fecha de la sentencia] [número del juzgado] [nombre del demandante] [nombre del demandado] [nombre del partido judicial].\n"
        "- Pon cada dato entre corchetes."
    )

def convert_response_to_numbers(response):
    return " ".join(word_to_number(part) for part in response.split(" "))

def display_extracted_data(extracted_data):
    # Imprime los datos extraídos
    numero_procedimiento, fecha_sentencia, numero_juzgado, nombre_demandante, nombre_demandado, partido_judicial = extracted_data
    print(f"Número de Procedimiento: {numero_procedimiento}")
    print(f"Fecha de Sentencia: {fecha_sentencia}")
    print(f"Nombre del Partido Judicial: {partido_judicial}")
    print(f"Número de Juzgado: {numero_juzgado}")
    print(f"Nombre del Demandante: {nombre_demandante.title()}")
    print(f"Nombre del Demandado: {nombre_demandado}")

def update_ui_components(extracted_data, summary_response):
    # Actualizar componentes de la UI
    numero_procedimiento, fecha_sentencia, numero_juzgado, _, _, partido_judicial = extracted_data
    partido_var.set(partido_judicial)
    update_juzgado()
    juzgado_var.set(numero_juzgado)
    fecha_cal.set_date(datetime.datetime.strptime(fecha_sentencia, '%d/%m/%Y').date())
    autos_var.set(numero_procedimiento.split('/')[0])
    autos_year_var.set(numero_procedimiento.split('/')[1])
    sintesis_var.set(texto_preparado_para_procesador(summary_response))
    update_text()


def cambiar_estado_boton(boton, estado):
    boton['state'] = estado


pdf_path = None  # Esto se define en la parte superior del script


def cargar_pdf_sentencia():
    global pdf_path  # Indicar que vamos a usar la variable global
    progress_bar.grid(row=0, column=2, padx=2, pady=2, sticky="ew")  # Mostras la barra de progreso
    progress_bar.start(5)  # Iniciar la animación de la barra
    # Cargar el último directorio desde el archivo JSON
    directorio_inicial = cargar_ultimo_directorio_cargar_sentencia_recurso()
    # Abre la ventana de diálogo para seleccionar un archivo PDF
    filepath = filedialog.askopenfilename(
        title="Seleccionar archivo PDF",
        filetypes=[("Archivos PDF", "*.pdf")],
        initialdir=directorio_inicial  # Usar el directorio inicial cargado
    )

    if filepath:
        # Guardar el nuevo último directorio en el archivo JSON
        ultimo_directorio = os.path.dirname(filepath)
        guardar_ultimo_directorio_cargar_sentencia_recurso(ultimo_directorio)
        pdf_path = filepath  # Aquí asignas el valor seleccionado a la variable global
        extract_info_from_pdf(filepath)
        print(f"Archivo seleccionado: {filepath}")
        # Por ejemplo, puedes almacenar la ruta del archivo en una variable,
        # abrir y leer el archivo, etc.
    cambiar_estado_boton(recurso_button, 'normal')
    progress_bar.stop()  # Detener la animación
    progress_bar.grid_remove()  # Ocultar la barra de progreso



def update_juzgado(*args):
    partido = partido_var.get()
    juzgados = partido_juzgado_map.get(partido, [])
    juzgado_var.set(juzgados[0] if juzgados else '')
    juzgado_options['values'] = juzgados

def fecha_a_texto(fecha):
    meses = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
        7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    dia = fecha.day
    mes = meses[fecha.month]
    año = fecha.year
    return f"{dia} de {mes} de {año}"


# Pre-compilar expresiones regulares
motivo_ordinal_pattern = re.compile(r"(único|primer|segundo|tercer|cuarto|quinto|sexto|séptimo|octavo|noveno|décimo) motivo")
comillas_pattern = re.compile(r'"([^"]*?)"|“(.*?)”|\'([^\']*?)\'|«(.*?)»')

def save_to_doc():
    # Verificación de los campos de recurso
    numero_recurso = recurso_var.get()
    año_recurso = recurso_year_var.get()

    if not numero_recurso:
        tk.messagebox.showwarning("Advertencia", "Por favor, ingrese el número de recurso.")
        recurso_entry.focus()
        return
    elif not año_recurso:
        tk.messagebox.showwarning("Advertencia", "Por favor, seleccione el año del recurso.")
        recurso_year_options.focus()
        return
    resumen = on_button_click()
    doc = configurar_documento()

    text_lines = right_text.get("1.0", tk.END).split('\n')
    procesar_lineas(text_lines, doc)

    # Guardar y abrir el documento
    filename = guardar_documento(doc, resumen)
    abrir_documento(filename)

def configurar_documento():
    doc = Document()
    # Configuración del estilo...
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return doc

def procesar_lineas(text_lines, doc):
    for line in text_lines:
        # Procesamiento de cada línea
        if line.startswith(("PRIMERO.-", "SEGUNDO.-", "TERCERO.-", "CUARTO.-", "QUINTO.-", "SEXTO.-", "SÉPTIMO.-", "OCTAVO.-")):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = p.add_run(line)
            run.bold = True
        elif "FALLAMOS" in line:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(line)
            run.bold = True
        elif '"' in line:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            parts = re.split(r'(".*?")', line)
            for part in parts:
                if part.startswith('"') and part.endswith('"'):
                    run = p.add_run(part)
                    run.italic = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph()
            start = 0
            for match in sorted(itertools.chain(comillas_pattern.finditer(line), motivo_ordinal_pattern.finditer(line)), key=lambda m: m.start()):
                p.add_run(line[start:match.start()])
                run = p.add_run(match.group(0))
                if comillas_pattern.match(match.group(0)):
                    run.italic = True
                elif motivo_ordinal_pattern.match(match.group(0)):
                    run.bold = True
                start = match.end()
            p.add_run(line[start:])


# Nombre del archivo JSON para almacenar la última ruta
ultimo_directorio_cargar_sentencia_recurso = "ultimo_directorio.json"

def guardar_ultimo_directorio_cargar_sentencia_recurso(directorio):
    with open(ultimo_directorio_cargar_sentencia_recurso, 'w') as archivo_json:
        json.dump({'ultimo_directorio': directorio}, archivo_json)

def cargar_ultimo_directorio_cargar_sentencia_recurso():
    try:
        with open(ultimo_directorio_cargar_sentencia_recurso, 'r') as archivo_json:
            data = json.load(archivo_json)
            return data.get('ultimo_directorio', '/')
    except FileNotFoundError:
        return '/'

# Ruta para almacenar la última ruta utilizada para guardar docuemntos
ultimo_directorio_guardar_archivo = os.path.join(os.path.expanduser('~'), 'last_path.json')

def cargar_ultimo_directorio_guardar_documento():
    # Intenta cargar la última ruta desde el archivo
    if os.path.exists(ultimo_directorio_guardar_archivo):
        with open(ultimo_directorio_guardar_archivo, 'r') as file:
            data = json.load(file)
            return data.get('last_path', os.path.expanduser('~/Documents'))
    else:
        return os.path.expanduser('~/Documents')

def guardar_ultimo_directorio_guardar_documento(path):
    # Guarda la última ruta en un archivo
    with open(ultimo_directorio_guardar_archivo, 'w') as file:
        json.dump({'last_path': path}, file)

def guardar_documento(doc, resumen):
    meses = {
        "January": "enero", "February": "febrero", "March": "marzo", "April": "abril",
        "May": "mayo", "June": "junio", "July": "julio", "August": "agosto",
        "September": "septiembre", "October": "octubre", "November": "noviembre", "December": "diciembre"
    }
    now = datetime.datetime.now()
    numero_recurso = recurso_var.get()
    año_recurso = recurso_year_var.get()

    # Cargar la última ruta utilizada, o usar la carpeta Documentos por defecto
    initial_dir = cargar_ultimo_directorio_guardar_documento()

    # Definir el nombre del archivo por defecto
    default_filename = f"Rec. {numero_recurso}-{año_recurso} - {resumen}.docx"

    # Abrir el cuadro de diálogo para guardar
    filetypes = [('Documentos de Word', '*.docx')]
    filename = filedialog.asksaveasfilename(
        title='Guardar archivo',
        initialdir=initial_dir,
        initialfile=default_filename,
        filetypes=filetypes
    )

    if filename:  # Si el usuario no canceló la operación
        if not filename.endswith('.docx'):
            filename += '.docx'
        doc.save(filename)
        print(f"Documento guardado en: {filename}")

        # Guardar la última ruta utilizada
        guardar_ultimo_directorio_guardar_documento(os.path.dirname(filename))

        return filename
    else:
        print("Guardado cancelado.")
        return None

def abrir_documento(filename):
    try:
        if os.name == 'nt':
            os.startfile(filename)
        elif os.name == 'posix':
            subprocess.call(('open', filename))
        else:
            subprocess.call(('xdg-open', filename))
    except Exception as e:
        print(f"Error abriendo el archivo: {e}")

def on_button_click():
    prompt = "Haz un resumen de 10 palabras del siguiente texto. No hables de la sentencia, habla del caso concreto, del litigo. Pon las palabras más significativas del caso. El resumen debe ser telegráfico:\n" + sintesis_var.get() # Obtener texto del widget de entrada
    response = get_response(prompt, "gpt-3.5-turbo")
    # Eliminar caracteres específicos de la respuesta
    response = response.replace('.', '').replace(',', '').replace(';', '').replace(':', '')
    return response


def clear_all():
    reset_text_variables([partido_var, recurrente_var, autos_var, autos_year_var, sintesis_var,
                          articulos_var, jurisprudencia_var, revocacion_var, revocacion_parcial_var,
                          razon_var, juzgado_var, revisiones_var, censuras_var, recurso_year_var, recurso_var,
                          impugnante_var])

    reset_checkboxes([cuestiones_var, inadmisibilidad_var, infraccion_var,
                      revision_var, sustantivas_var])

    reset_radiobuttons()

    clear_text_fields([right_text])

    # Actualizar interfaz y elementos dependientes
    fecha_cal.set_date(datetime.date.today())
    update_ui_elements()


def reset_text_variables(variables):
    for var in variables:
        var.set('')


def reset_checkboxes(checkbox_vars):
    for chk_var in checkbox_vars:
        chk_var.set(False)


def reset_radiobuttons():
    costas_var.set(None)
    trabajador_empresa_var.set("Trabajador")
    for tipo_var in tipos_vars:
        tipo_var.set(False)


def clear_text_fields(text_widgets):
    for text_widget in text_widgets:
        text_widget.delete("1.0", tk.END)


def update_ui_elements():
    update_juzgado()
    toggle_infraccion()
    toggle_costas_frames()
    toggle_censuras()


# Función que inserta la respuesta en la sección adecuada
def insertar_respuesta_revision(respuesta_texto):
    # Encuentra la posición en el texto donde queremos insertar la respuesta
    busqueda_frase = "Sentado lo anterior, pasamos a analizar las pretensiones concretas."
    content = right_text.get("1.0", tk.END)

    # Encuentra el final de la frase donde agregaremos la respuesta
    inicio_index = content.find(busqueda_frase)
    if inicio_index != -1:
        fin = inicio_index + len(busqueda_frase)
        pos = right_text.search(busqueda_frase, "1.0", tk.END)
        if pos:
            # Obtiene la posición de la línea y la columna para insertar texto
            linea, columna = pos.split('.')
            posicion = f"{linea}.{int(columna) + len(busqueda_frase)}"
            # Agrega un salto de línea antes y después del texto insertado
            right_text.insert(posicion, "\n\n" + respuesta_texto)
            right_text.see(posicion)  # Desplaza el texto para que la inserción sea visible
    else:
        print("No se encontró la frase para insertar la respuesta.")

def update_text(*args):
    fecha_seleccionada = fecha_a_texto(fecha_cal.get_date())

    ordinals = ["SEGUNDO.-", "TERCERO.-", "CUARTO.-", "QUINTO.-", "SEXTO.-", "SÉPTIMO.-", "OCTAVO.-", "NOVENO.-"]


    content_sections = [f"PRIMERO.- Síntesis de la litis\n{sintesis_var.get()}"]

    index = 0

    if cuestiones_var.get():
        content_sections.append(f"""{ordinals[index]} Cuestiones de orden público
	Con carácter prioritario, debemos recordar en los términos de la Sentencia del Tribunal Supremo de 11 de mayo de 2018 (Pleno, rec. 1800/2016), reiterada por otras posteriores, que la cuestión de la recurribilidad de la Sentencia de instancia ha de examinarse de oficio por afectar al orden público procesal y a la propia competencia funcional, incluso “antes de llevar a cabo cualquier pronunciamiento sobre la contradicción o sobre el fondo del asunto”.
	El artículo 190 de la Ley Reguladora de la Jurisdicción Social (en adelante LRJS) establece con carácter general el ámbito de la suplicación, precisando de forma positiva y negativa el artículo 191 LRJS las resoluciones susceptibles de ser combatidas a través del citado recurso extraordinario.
	En el presente caso...
""")
        index += 1

    if inadmisibilidad_var.get():
        content_sections.append(f"""{ordinals[index]} Inadmisibilidad del recurso 
	Por razones de sistemática hemos de comenzar con el análisis y resolución del motivo de inadmisibilidad alegado por la parte recurrida, cuya estimación haría innecesaria la resolución del resto de motivos impugnatorios. 
	El artículo 190 de la Ley Reguladora de la Jurisdicción Social (en adelante LRJS) establece con carácter general el ámbito de la suplicación, precisando de forma positiva y negativa el artículo 191 LRJS las resoluciones susceptibles de ser combatidas a través del citado recurso extraordinario. 
	En el presente caso…
""")
        index += 1

    if infraccion_var.get():
        infraccion_content = [f"""{ordinals[index]} Infracción de normas procesales
	La nulidad de actuaciones sólo puede articularse por la vía del art. 193.a) LRJS, que exige como requisitos para que pueda prosperar el motivo aducido, los siguientes: 
		1) Identificar el precepto procesal que se entienda infringido o doctrina emanada del TS, TC o los órganos jurisdiccionales instituidos en los Tratados y Acuerdos internacionales en materia de derechos humanos y libertades fundamentales ratificados por España, o la del TJUE (art. 219.2 LRJS y art. 1.6 CC) 
		2) La infracción debe haber provocado un perjuicio real sobre los derechos de defensa del interesado, irrogándole indefensión (STC 168/2002). 
		3) El defecto procesal no puede alegarse por la parte que lo provocó. 
		4) Es preciso que la parte perjudicada haya formulado protesta en tiempo y forma.
	Ha de recordarse que es criterio reiterado de las diversas salas de suplicación que la declaración de nulidad de actuaciones es un remedio excepcional, que ha de aplicarse con criterio restrictivo, pues una interpretación amplia de la posibilidad de anulación podría incluso vulnerar el derecho a la tutela judicial efectiva, proclamado en el artículo 24 CE, en su vertiente del derecho a un proceso público sin dilaciones indebidas. A este respecto, no solamente es la celeridad uno de los principios orientadores de la interpretación y aplicación de las normas reguladoras del proceso laboral ordinario (artículo 74.1 LJS) sino que la Ley Reguladora de la Jurisdicción Social aumenta considerablemente las posibilidades de la Sala de suplicación de examinar el fondo del asunto, tanto por vulneración de las normas reguladoras de la sentencia (artículo 202.2) como por no haber entrado la resolución de instancia en el fondo al haber apreciado alguna circunstancia obstativa, como puede ser una excepción procesal (artículo 202.3), siempre que el relato de hechos probados de la sentencia y demás antecedentes no cuestionados sean suficientes para realizar tal pronunciamiento sobre el fondo."""]
        for i, var in enumerate(tipos_vars):
            if var.get():
                infraccion_content.append(tipos_options[i])
        content_sections.append("\n".join(infraccion_content))
        content_sections.append("	En el presente caso...\n")
        index += 1
    else:
        for var in tipos_vars:
            var.set(False)  # If "Infracción de normas procesales" is unchecked, uncheck all related checkboxes

    if revision_var.get():
        content_sections.append(f"""{ordinals[index]} Revisión de hechos probados 
	La parte recurrente interesa la revisión fáctica de la sentencia, al amparo del art. 193.b) LRJS. 
	En primer lugar, debe razonarse que, con carácter general, el órgano que conoce del recurso extraordinario de suplicación no puede efectuar una nueva ponderación de la prueba, pues es a quien ha presidido el acto del juicio en la instancia a quien corresponde apreciar los elementos de convicción para establecer la verdad procesal intentando que la misma se acerque lo más posible a la verdad material. Ahora bien, tal principio debe ser matizado en el sentido de que el tribunal ad quem está autorizado para revisar las conclusiones fácticas cuando los documentos o pericias citados por la parte recurrente pongan de manifiesto de manera patente e incuestionable el error en el que ha incurrido la sentencia recurrida, o la irracionalidad o arbitrariedad de sus conclusiones. De otra forma, carecería de sentido la previsión del artículo 193.b) LRJS. 
	Además, debe señalarse que la jurisprudencia viene exigiendo con reiteración, hasta el punto de constituir doctrina pacifica, que para estimar este motivo es necesario que concurran los siguientes requisitos: 
		1.- Que se señale con precisión y claridad cuál es el hecho afirmado, negado u omitido, que la parte recurrente considera equivocado, contrario a lo acreditado o que consta con evidencia y no ha sido incorporado al relato fáctico. 
		2.- Que se ofrezca un texto alternativo concreto para figurar en la narración fáctica calificada de errónea, bien sustituyendo a alguno de sus puntos, bien complementándolos. 
		3.- Que se citen pormenorizadamente los documentos o pericias de los que se considera se desprende la equivocación de la sentencia, sin que sea dable admitir su invocación genérica, ni plantearse la revisión de cuestiones fácticas no discutidas a lo largo del proceso; la ley señala que el error debe ponerse de manifiesto precisamente merced a las pruebas documentales o periciales practicadas en la instancia. 
		4.- Que esos documentos o pericias pongan de manifiesto el error de manera clara, evidente, directa y patente; sin necesidad de acudir a conjeturas, suposiciones o argumentaciones más o menos lógicas, naturales y razonables, de modo que sólo son admisibles para poner de manifiesto el error de hecho, los documentos que ostenten un decisivo valor probatorio, tengan concluyente poder de convicción por su eficacia, suficiencia, fehaciencia o idoneidad. 
		5.- Que la revisión pretendida sea trascendente a la parte dispositiva de la sentencia, con efectos modificadores de ésta, pues el principio de economía procesal impide incorporar hechos cuya inclusión a nada práctico conduciría, si bien cabrá admitir la modificación fáctica cuando no siendo trascendente en esta instancia pudiera resultarlo en otras superiores. 
		6.- Que no se trate de una nueva valoración global de la prueba incorporada al proceso. 
	Estos criterios han sido reafirmados, entre otras, por la reciente Sentencia 90/2022 del Tribunal Supremo, de 1 de febrero (rec. 2429/2019). 
	Sentado lo anterior, pasamos a analizar las pretensiones concretas.""")
        if revisiones_var.get():
            content_sections.append(tabular_parrafos(revisiones_var.get()))
        else:
            content_sections.append("")
        index += 1

    if sustantivas_var.get():
        content_sections.append(f"""{ordinals[index]} Infracción de normas sustantivas y de la jurisprudencia """)
        articulos = articulos_var.get().strip()
        jurisprudencia = jurisprudencia_var.get().strip()

        if articulos and jurisprudencia:
            content_sections.append(
                f"""	La parte recurrente interesa la revocación de la sentencia, al amparo del art. 193.c) LRJS, alegando la infracción de los arts. {articulos}, así como la infracción de la jurisprudencia sentada por {jurisprudencia}. 
	Con carácter previo debemos indicar – por lo que respecta a las normas citadas por la parte recurrente en su recurso como infringidas – que el recurso de suplicación no es una apelación o segunda instancia, sino un recurso extraordinario sujeto a motivos tasados en cuya formulación se han de respetar los requisitos legales.
	Los motivos basados en el apartado c) del art. 193 LRJS se destinan a la impugnación del fallo por error in iudicando, y el recurrente tiene la carga de:
		a)	Citar debidamente el precepto o preceptos sustantivos y en su caso la jurisprudencia que, a su juicio, han sido vulnerados por el fallo de la sentencia, articulando motivos separados para cada precepto o grupo de preceptos que guarden unidad temática.
		b)	Razonar la pertinencia y fundamentación de los motivos (art. 196.2 LRJS) lo cual exige argumentar la conexión entre el contenido normativo de las normas o jurisprudencia citadas y el litigio, mostrando cómo su correcta aplicación debería haber llevado a dar distinta solución al debate.
	Dicha doctrina jurisprudencial, así mismo señala que no basta que el recurso cite la disposición legal conculcada si contiene diversos artículos, sino que es preciso que se señale el específico precepto que se entiende vulnerado, y si el precepto contiene varios apartados resulta igualmente indispensable señalar expresamente cuál de ellos se reputa infringido.""")
            if censuras_var.get():
                content_sections.append(tabular_parrafos(censuras_var.get()))
        elif articulos:
            content_sections.append(f"""	La parte recurrente interesa la revocación de la sentencia, al amparo del art. 193.c) LRJS, alegando la infracción de los arts. {articulos}.
	Con carácter previo debemos indicar – por lo que respecta a las normas citadas por la parte recurrente en su recurso como infringidas – que el recurso de suplicación no es una apelación o segunda instancia, sino un recurso extraordinario sujeto a motivos tasados en cuya formulación se han de respetar los requisitos legales.
	Los motivos basados en el apartado c) del art. 193 LRJS se destinan a la impugnación del fallo por error in iudicando, y el recurrente tiene la carga de:
		a)	Citar debidamente el precepto o preceptos sustantivos y en su caso la jurisprudencia que, a su juicio, han sido vulnerados por el fallo de la sentencia, articulando motivos separados para cada precepto o grupo de preceptos que guarden unidad temática.
		b)	Razonar la pertinencia y fundamentación de los motivos (art. 196.2 LRJS) lo cual exige argumentar la conexión entre el contenido normativo de las normas o jurisprudencia citadas y el litigio, mostrando cómo su correcta aplicación debería haber llevado a dar distinta solución al debate.
	Dicha doctrina jurisprudencial, así mismo señala que no basta que el recurso cite la disposición legal conculcada si contiene diversos artículos, sino que es preciso que se señale el específico precepto que se entiende vulnerado, y si el precepto contiene varios apartados resulta igualmente indispensable señalar expresamente cuál de ellos se reputa infringido.""")
            if censuras_var.get():
                content_sections.append(tabular_parrafos(censuras_var.get()))
        index += 1

    # Añadir el texto de "Costas, depósitos y consignaciones"
    if costas_var.get() == "Estimación":
        content_sections.append(f"""{ordinals[index]} Costas, depósitos y consignaciones""")
        if trabajador_empresa_var.get() == "Trabajador":
            content_sections.append(f"""	La estimación del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, implica que no haya lugar a hacer pronunciamiento alguno sobre las costas del recurso. 


	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	ESTIMAR el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, revocando la misma en el sentido de que:
""")
        else:
            content_sections.append(f"""	La estimación del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, implica que no haya lugar a hacer pronunciamiento alguno sobre las costas del recurso.
	Se decreta también, como preceptúa el artículo 203.1 LRJS, la devolución de la totalidad del depósito constituido para recurrir y dar el destino legal a los aseguramientos si los hubiere, todo ello firme que sea la presente resolución. 


	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	ESTIMAR el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, revocando la misma en el sentido de que:
""")
        content_sections.append(f'"{revocacion_entry.get()}"')
        content_sections.append(f"""\n	Sin costas.""")
    elif costas_var.get() == "Estimación parcial":
        content_sections.append(f"""{ordinals[index]} Costas, depósitos y consignaciones""")
        if trabajador_empresa_var.get() == "Trabajador":
            content_sections.append(f"""	La estimación parcial del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, implica que no haya lugar a hacer pronunciamiento alguno sobre las costas del recurso. 


	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	ESTIMAR PARCIALMENTE el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, revocando la misma en el sentido de que:
""")
        else:
            content_sections.append(f"""	La estimación parcial del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, implica que no haya lugar a hacer pronunciamiento alguno sobre las costas del recurso.
	Se decreta también, como preceptúa el artículo 203.1 LRJS, la devolución de la totalidad del depósito constituido para recurrir y dar el destino legal a los aseguramientos si los hubiere, todo ello firme que sea la presente resolución. 


	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	ESTIMAR PARCIALMENTE el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, revocando la misma en el sentido de que:
""")
        content_sections.append(f'"{revocacion_parcial_entry.get()}"')
        content_sections.append(f"""\n	Sin costas.""")
    elif costas_var.get() == "Desestimación con costas":
        content_sections.append(f"""{ordinals[index]} Costas, depósitos y consignaciones
	La desestimación del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, implica la expresa imposición de condena en las costas del recurso a la parte vencida en el mismo, condena que incluye el pago de la Minuta de Honorarios del Letrado de la parte impugnante, en la cuantía que esta Sala, y dentro de los límites legales, señala en {cuantia_costas_var.get()} euros.
	Se decreta también, como preceptúa el artículo 204.4 LRJS, la pérdida del depósito constituido para recurrir que, una vez sea firme la sentencia, se ingresará al Tesoro público; y por fin, se condena a la misma parte vencida a la pérdida de la cantidad consignada a la que se dará el destino legal.

	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	DESESTIMAR el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, confirmando la misma en su integridad.

	Se imponen las costas a la parte recurrente, en la cuantía de {cuantia_costas_var.get()} euros. Se decreta la pérdida del depósito y las consignaciones efectuadas, en su caso, para recurrir.""")
    elif costas_var.get() == "Desestimación sin costas":
        content_sections.append(f"""{ordinals[index]} Costas, depósitos y consignaciones
	La desestimación del recurso, de conformidad con lo establecido en el artículo 235.1 LRJS, no implica en el presente caso la expresa imposición de condena en las costas del recurso a la parte vencida en el mismo, al ser {razon_var.get()}.

	Vistos los preceptos legales citados, sus concordantes y demás disposiciones de general y pertinente aplicación.

FALLAMOS

	DESESTIMAR el recurso de suplicación interpuesto por {recurrente_var.get()} contra la sentencia del Juzgado de lo Social nº {juzgado_var.get()} de los de {partido_var.get()}, de fecha {fecha_a_texto(fecha_cal.get_date())}, dictada en autos nº {autos_var.get()}/{autos_year_var.get()}, confirmando la misma en su integridad.

	Sin costas.""")

    final_text = "\n".join(content_sections)
    right_text.delete("1.0", tk.END)
    right_text.insert(tk.END, final_text)

# FUNCIONES PARA LA CARGA DEL RECURSO

pdf_path_for_text_extraction = None  # Definición de la nueva variable global
def load_pdf_and_get_text():
    global pdf_path_for_text_extraction  # Usar la nueva variable global
    # Cargar el último directorio desde el archivo JSON
    directorio_inicial = cargar_ultimo_directorio_cargar_sentencia_recurso()
    file_path = filedialog.askopenfilename(
        title="Seleccionar archivo PDF",
        filetypes=[("Archivos PDF", "*.pdf")],
        initialdir=directorio_inicial  # Usar el directorio inicial cargado
    )
    if file_path:
        # Guardar el nuevo último directorio en el archivo JSON
        ultimo_directorio = os.path.dirname(file_path)
        guardar_ultimo_directorio_cargar_sentencia_recurso(ultimo_directorio)
        pdf_path_for_text_extraction = file_path  # Almacenar la ruta del archivo PDF en la nueva variable global
        reader = PdfReader(file_path)
        text = ''
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    else:
        return None

def numero_a_palabra(n):
    numeros = ["cero", "un", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve", "diez"]
    return numeros[n] if n < len(numeros) else str(n)

def get_response3(pdf_text):
    contenido_pdf = f"ESTE ES EL RECURSO:\n{pdf_text}\n###########\n"
    response = openai.ChatCompletion.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": contenido_pdf},
            {"role": "user", "content": """Eres un experto analizador de recursos de suplicación. Quiero que analices el Recurso. Los recursos de suplicación pueden interponerse por el art. 193.a) LRJS, art. 193.b) LRJS o art. 193.c) LRJS.
            - El apartado a) del art. 193 LRJS regula un motivo de recurso que pretende reponer los autos al estado en el que se encontraban en el momento de cometerse una infracción de normas o garantías del procedimiento que haya producido indefensión.
            - El apartado b) del art. 193 LRJS regula un motivo de recurso que pretende revisar los hechos declarados probados, a la vista de las pruebas documentales y periciales practicadas.
            - El apartado c) del art. 193 LRJS regula un motivo de recurso que pretende examinar las infracciones de normas sustantivas o de la jurisprudencia.
            - Tienes que decirme, cuántos motivos suplicacionales hay de la letra a), cuántos motivos suplicacionales hay de la letra b) y cuántos motivos suplicacionales hay de la letra c).
            - Extrae también el nombre del recurrente, NO PONGAS EL NOMBRE DEL ABOGADO, asegurate.
            - Si hay algún motivo del apartado c), examinalo detenidamente, y haz una lista de los artículos que dice que se han infringido. Haz una relación de ellos. Por ejemplo [art. 1 ET, art. 24 RD 1457/2023, art. 9 LGSS].
            - El formato de la respuesta será el siguiente: [nombre recurrente] [numero a)] [numero b)] [numero c)] [artículos infringidos], por ejemplo [Pedro Rodríguez] [0] [4] [2] [art. 1 ET, art. 24 RD 1457/2023, art. 9 LGSS].
            - Respeta el formato, tienes prohibido cambiar el formato.
            - El formato de la respuesta será el siguiente: [nombre recurrente] [numero a)] [numero b)] [numero c)] [artículos infringidos], por ejemplo [Pedro Rodríguez] [0] [4] [2] [art. 1 ET, art. 24 RD 1457/2023, art. 9 LGSS].
            - Respeta el formato, tienes prohibido cambiar el formato.
            - Take a deep breath and work on this problem step by step."""}
        ]
    )
    return response['choices'][0]['message']['content']

motivos_b = 0  # Definiendo la variable a nivel global
motivos_c = 0  # Definiendo la variable a nivel global
def interpret_response(response_str):
    global motivos_b  # Declarando que vamos a usar la variable global
    global motivos_c  # Declarando que vamos a usar la variable global
    elements = response_str.strip('[]').split('] [')
    if len(elements) == 5:
        nombre_recurrente = elements[0].title()
        recurrente_var.set(nombre_recurrente)
        motivos_a = int(elements[1])
        infraccion_var.set(motivos_a > 0)
        motivos_b = int(elements[2])
        revision_var.set(motivos_b > 0)
        motivos_c = int(elements[3])
        sustantivas_var.set(motivos_c > 0)
        articulos_infringidos = elements[4]
        if articulos_infringidos:
            articulos_var.set(f'{articulos_infringidos}')
            toggle_censuras()
        else:
            articulos_var.set('NO SE HAN PODIDO IDENTIFICAR LOS ARTÍCULOS, INTRODÚCELOS')
            toggle_censuras()

        partes_frase = []
        if motivos_a > 0:
            partes_frase.append(
                f"{numero_a_palabra(motivos_a)} motivo{'s' if motivos_a > 1 else ''} de infracción de normas o garantías procesales")
        if motivos_b > 0:
            partes_frase.append(
                f"{numero_a_palabra(motivos_b)} motivo{'s' if motivos_b > 1 else ''} de revisión fáctica")
        if motivos_c > 0:
            partes_frase.append(
                f"{numero_a_palabra(motivos_c)} motivo{'s' if motivos_c > 1 else ''} de censura jurídica")

        # Construir la frase base
        frase_base = f"\nDisconforme la parte actuante, {nombre_recurrente}, interpone el presente recurso de suplicación articulando "

        # Unir las frases de motivos
        if len(partes_frase) > 1:
            frase_motivos = ', '.join(partes_frase[:-1]) + f" y {partes_frase[-1]}"
        elif len(partes_frase) == 1:
            frase_motivos = partes_frase[0]
        else:
            frase_motivos = ""

        # Completar la frase final
        if frase_motivos:
            frase_final = frase_base + frase_motivos + ", pretendiendo la revocación de la Sentencia con la condena de la parte demandada. "
        else:
            frase_final = frase_base + "pretendiendo la revocación de la Sentencia con la condena de la parte demandada. "

        if impugnante_var.get():
            frase_final += "El recurso fue impugnado por la representación letrada de " + impugnante_var.get() + "\n"
        else:
            frase_final += "\n"
        sintesis_entry.insert(tk.END, tabular_parrafos(frase_final))
        update_text()
        pass
    else:
        return "El formato de respuesta no es correcto."


def cargar_pdf_recurso():
    progress_bar.grid(row=0, column=2, padx=2, pady=2, sticky="ew")  # Mostras la barra de progreso
    progress_bar.start(10)  # Iniciar la animación de la barra
    pdf_text = load_pdf_and_get_text()
    if pdf_text:
        gpt_response = get_response3(pdf_text)
        formatted_response = interpret_response(gpt_response)  # Usa la función para formatear la respuesta

    else:
        messagebox.showinfo("Información", "No se ha cargado ningún archivo PDF.")
    cambiar_estado_boton(analizar_revisiones_button, 'normal')
    cambiar_estado_boton(analizar_censuras_button, 'normal')
    progress_bar.stop()  # Detener la animación
    progress_bar.grid_remove()  # Ocultar la barra de progreso

#FIN DE LAS FUNCIONES PARA LA CARGA DEL RECURSO

#ANALISIS DE LAS REVISIONES FÁCTICAS

# Función para unir los PDFs y añadir los textos necesarios
def unir_pdfs_y_anadir_texto(pdf1, pdf2):
    texto_unido = "ESTA ES LA SENTENCIA:\n\n"
    texto_unido += leer_pdf(pdf1)
    texto_unido += "\n\nESTE ES EL RECURSO:\n\n"
    texto_unido += leer_pdf(pdf2)
    return texto_unido

# Función para leer el contenido de un archivo PDF
def leer_pdf(archivo):
    with open(archivo, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        texto = ""
        for page in reader.pages:
            texto += page.extract_text() + "\n"
    return texto

# Función para obtener la respuesta de ChatGPT
def get_response4(prompt, texto_unido):
    response = openai.ChatCompletion.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": texto_unido},
            {"role": "user", "content": prompt},
        ]
    )
    return response['choices'][0]['message']['content']

# FUNCIÓN PARA CONSTRUIR EL PROMPT

def reemplazar_motivos(texto, numero_de_repeticiones):
    # Mapeo de los ordinales
    ordinales = {
        "primero": "primer",
        "segundo": "segundo",
        "tercero": "tercer",
        "cuarto": "cuarto",
        "quinto": "quinto",
        "sexto": "sexto",
        "séptimo": "séptimo",
        "octavo": "octavo",
        # Añade más si es necesario
    }

    # Reemplazar "motivo primero" por "único motivo" si numero_de_repeticiones es 1
    if numero_de_repeticiones == 1:
        texto = texto.replace("motivo primero", "único motivo")
    else:
        # Reemplazar cada "motivo [ordinal]" por "[ordinal] motivo"
        for ordinal in ordinales:
            texto = texto.replace(f"motivo {ordinal}", f"{ordinales[ordinal]} motivo")

    return texto

def generar_respuestas(prompt_base, numero_de_repeticiones, pdf2):
    ordinales = ["primero", "segundo", "tercero", "cuarto", "quinto", "sexto", "séptimo", "octavo"]
    respuesta_final = ""

    for i in range(numero_de_repeticiones):
        # Reemplazar [numero] por el número de repeticiones y ajustar singular o plural
        if numero_de_repeticiones == 1:
            prompt_con_numero = prompt_base.replace("[numero]", "1 motivo")
        else:
            prompt_con_numero = prompt_base.replace("[numero]", f"{numero_de_repeticiones} motivos")

        # Reemplazar todas las instancias de 'primero' por el ordinal correspondiente
        prompt_modificado = prompt_con_numero.replace("primero", ordinales[i])

        # Generar la frase con los ordinales para la frase 'No analices ni el motivo...'
        frase_ordinales = ", ni el ".join(ordinales[:i])
        if frase_ordinales:
            frase_ordinales = "No analices el motivo " + frase_ordinales + ", ya los analizaste antes. "

        # Insertar la frase en el lugar correcto del prompt modificado
        partes_prompt = prompt_modificado.split("El motivo puede ser una modificación de un hecho probado o la adición de un nuevo hecho probado o de un nuevo párrafo a un hecho probado ya existente.")
        prompt_modificado = partes_prompt[0] + frase_ordinales + "El motivo puede ser una modificación de un hecho probado o la adición de un nuevo hecho probado o de un nuevo párrafo a un hecho probado ya existente." + partes_prompt[1]
        print(prompt_modificado)
        respuesta = get_response4(prompt_modificado, pdf2)
        pyperclip.copy(respuesta)
        respuesta_final += f"{respuesta}\n"

    return respuesta_final

# Función para manejar
def analizar_revisiones_facticas():
    progress_bar.grid(row=0, column=2, padx=2, pady=2, sticky="ew")  # Mostras la barra de progreso
    progress_bar.start(5)  # Iniciar la animación de la barra
    with open(pdf_path, 'rb') as pdf_file:
        reader = PyPDF2.PdfReader(pdf_path)
        full_text = ''
        for page in reader.pages:
            full_text += page.extract_text()

        start_phrase = "PROBADOS"
        end_phrase = "FUNDAMENTOS"
        extracted_text = extract_text_between_phrases(full_text, start_phrase, end_phrase)

        # Añadir texto al inicio y al final de extracted_text
        extracted_text = 'ESTOS SON LOS HECHOS PROBADOS\n\n' + extracted_text + '\n\nESTE ES EL RECURSO\n\n'

    # Leer el segundo PDF y añadir su contenido a extracted_text
    pdf2_content = leer_pdf(pdf_path_for_text_extraction)
    extracted_text += pdf2_content
    print(extracted_text)
    numero_de_repeticiones = motivos_b  # Puedes cambiar este número según lo necesites
    respuesta_final = generar_respuestas(promptrevision, numero_de_repeticiones, extracted_text)
    respuesta_final = reemplazar_motivos(respuesta_final, numero_de_repeticiones)
    progress_bar.stop()  # Detener la animación
    progress_bar.grid_remove()  # Ocultar la barra de progreso

# Asegúrate de definir la variable 'prompt' fuera de la función, como una variable global o directamente en la función.
    revisiones_var.set(respuesta_final)  # Devuelve la respuesta completa para su uso posterior)


promptrevision = """Eres un experto analizador de recursos de suplicación. Quiero que analices el Recurso. Los recursos de suplicación pueden interponerse por el art. 193.a) LRJS, art. 193.b) LRJS o art. 193.c) LRJS.
            - El apartado a) del art. 193 LRJS regula un motivo de recurso que pretende reponer los autos al estado en el que se encontraban en el momento de cometerse una infracción de normas o garantías del procedimiento que haya producido indefensión.
            - El apartado b) del art. 193 LRJS regula un motivo de recurso que pretende revisar los hechos declarados probados, a la vista de las pruebas documentales y periciales practicadas.
            - El apartado c) del art. 193 LRJS regula un motivo de recurso que pretende examinar las infracciones de normas sustantivas o de la jurisprudencia.
            - Tienes que buscar los motivos suplicacionales de la letra b).
            - En el documento hay [numero] motivos de la letra b).
            - Una vez que los encuentres centrate sólo en el motivo primero. Eres un experto contanto motivos. Sólo quiero que analices el motivo primero, no te confundas. El motivo puede ser una modificación de un hecho probado o la adición de un nuevo hecho probado o de un nuevo párrafo a un hecho probado ya existente. 
            - Si el primero de los motivos se refiere a la modificación de un hecho probado de la sentencia tienes que escribir lo siguiente:

"Como motivo primero de revisión fáctica, la parte recurrente, interesa la modificación del Hecho Probado [decir qué número de hecho probado es], cuya redacción original es:

[ve a los HECHOS PROBADOS y copia exactamente el hecho probado que se pretende modificar]

La redacción que se propone sería la siguiente:

[ve al RECURSO y copia exactamente la modificación que se pretende]

Para ello, el recurrente se apoya en [hacer una relación de las pruebas en las que el recurrente se basa para proponer la modificación]."

            - Tienes que seguir esta estructura. Tienes prohibido utilizar otra estructura, me enfadaría si utilizaras otra estructura.
            - Si el primero de los motivos se refiere a la adición de un nuevo hecho probado o de un nuevo párrafo a un hecho probado ya existente, tienes que escribir lo siguiente:

"Como motivo primero de revisión fáctica, la parte recurrente, interesa la adición de un nuevo Hecho Probado [decir qué número de hecho probado es]/de un nuevo párrafo al Hecho Probado [decir qué número de hecho probado es], cuya redacción sería la siguiente:

[ve al RECURSO y copia exactamente la adición que se pretende]

Para ello, el recurrente se apoya en [hacer una relación de las pruebas en las que el recurrente se basa para proponer la adición]."

            - Tienes que seguir esta estructura. Tienes prohibido utilizar otra estructura, me enfadaría si utilizaras otra estructura.
            - Take a deep breath and work on this problem step by step."""

# FIN DE ANALISIS DE REVISIONES FÁCTICAS
# ANALISIS DE LAS CENSURAS JURÍDICAS

def generar_respuestas2(prompt_base, numero_de_repeticiones, pdf2):
    ordinales = ["primero", "segundo", "tercero", "cuarto", "quinto", "sexto", "séptimo", "octavo"]
    respuesta_final = ""

    for i in range(numero_de_repeticiones):
        # Reemplazar [numero] por el número de repeticiones y ajustar singular o plural
        if numero_de_repeticiones == 1:
            prompt_con_numero = prompt_base.replace("[numero]", "1 motivo")
        else:
            prompt_con_numero = prompt_base.replace("[numero]", f"{numero_de_repeticiones} motivos")

        # Reemplazar todas las instancias de 'primero' por el ordinal correspondiente
        prompt_modificado = prompt_con_numero.replace("primero", ordinales[i])

        # Generar la frase con los ordinales para la frase 'No analices ni el motivo...'
        frase_ordinales = ", ni el ".join(ordinales[:i])
        if frase_ordinales:
            frase_ordinales = "No analices el motivo " + frase_ordinales + ", ya los analizaste antes. "
        # Insertar la frase en el lugar correcto del prompt modificado

        partes_prompt = prompt_modificado.split("Think step by step.")
        prompt_modificado = partes_prompt[0] + frase_ordinales + "Think step by step." + partes_prompt[1]

        print(prompt_modificado)
        respuesta = get_response4(prompt_modificado, pdf2)
        pyperclip.copy(respuesta)

        # Agregar el prompt modificado a la respuesta final
        respuesta_final += f"{respuesta}\n"
        respuesta_final = reemplazar_motivos(respuesta_final, numero_de_repeticiones)

    return respuesta_final


def analizar_censuras_juridicas():
    progress_bar.grid(row=0, column=2, padx=2, pady=2, sticky="ew")  # Mostras la barra de progreso
    progress_bar.start(5)  # Iniciar la animación de la barra
    pdf2 = leer_pdf(pdf_path_for_text_extraction)
    numero_de_repeticiones = motivos_c  # Puedes cambiar este número según lo necesites
    respuesta_final = generar_respuestas2(promptcensura, numero_de_repeticiones, pdf2)
    progress_bar.stop()  # Detener la animación
    progress_bar.grid_remove()  # Ocultar la barra de progreso

    # Asegúrate de definir la variable 'prompt' fuera de la función, como una variable global o directamente en la función.
    censuras_var.set(respuesta_final)  # Devuelve la respuesta completa para su uso posterior)

promptcensura = """Eres un experto analizador de recursos de suplicación. Quiero que analices el Recurso. Los recursos de suplicación pueden interponerse por el art. 193.a) LRJS, art. 193.b) LRJS o art. 193.c) LRJS.
            - El apartado a) del art. 193 LRJS regula un motivo de recurso que pretende reponer los autos al estado en el que se encontraban en el momento de cometerse una infracción de normas o garantías del procedimiento que haya producido indefensión.
            - El apartado b) del art. 193 LRJS regula un motivo de recurso que pretende revisar los hechos declarados probados, a la vista de las pruebas documentales y periciales practicadas.
            - El apartado c) del art. 193 LRJS regula un motivo de recurso que pretende examinar las infracciones de normas sustantivas o de la jurisprudencia.
            - Tienes que buscar los motivos suplicacionales de la letra c).
            - En el documento hay [numero] de la letra c).
            - Una vez que los encuentres centrate sólo en el motivo primero. Eres un experto contanto motivos. Sólo quiero que analices el motivo primero.
            - Think step by step. 
            - Tienes que hacer un resumen de 200 palabras del motivo primero, de las razones por las cuales el recurrente recurre la sentencia.
            - Sigue la siguiente estructura. Sólo del motivo primero de la letra c):

"Como motivo primero de censura jurídica, la parte recurrente interesa la revocación de la sentencia de instancia al considerar que se ha producido la infracción [decir los artículos que el recurrente dice que se han infringido en el primero motivo de la letra c)], a saber [aquí empieza el resumen]."

            - Tienes que seguir esta estructura. Tienes prohibido utilizar otra estructura, me enfadaría si utilizaras otra estructura.
            - Take a deep breath and work on this problem step by step."""

# FIN DEL ANALISIS DE LA CENSURA JURÍDICA

def on_mousewheel(event):
    left_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

def bind_to_mousewheel(widget):
    widget.bind("<MouseWheel>", on_mousewheel)

# Funciones relacionadas con el portapapeles
def verificar_portapapeles():
    try:
        ultimo_copiado = root.clipboard_get()
        if ultimo_copiado and (ultimo_copiado not in historial_portapapeles):
            historial_portapapeles.append(ultimo_copiado)
    except tk.TclError:
        pass
    root.after(1000, verificar_portapapeles)

def pegar_en_texto(elemento, lugar):
    conteo_uso_opciones[elemento] = conteo_uso_opciones.get(elemento, 0) + 1
    guardar_conteo()
    cursor_pos = lugar.index(tk.INSERT)
    lugar.insert(cursor_pos, elemento)


def truncar_texto(texto, max_ancho):
    return (texto[:max_ancho] + '...') if len(texto) > max_ancho else texto

def mostrar_menu(event, lugar):
    menu = tk.Menu(root, tearoff=0)
    for elemento in historial_portapapeles:
        texto_menu = truncar_texto(elemento, MAX_ANCHO_MENU)
        menu.add_command(label=texto_menu, command=lambda e=elemento, l=lugar: pegar_en_texto(e, l))

    try:
        # Coloca el menú en la posición del cursor
        menu.tk_popup(event.x_root, event.y_root)
    finally:
        # Asegura que el menú se libere después de ser utilizado
        menu.grab_release()

def abrir_ventana_modificacion(lugar):
    modificacion = tk.Toplevel(root)
    modificacion.title("Modificación de Hechos Probados")
    modificacion.configure(bg='white')

    # Set a minimum size for the window
    modificacion.minsize(670, 274)
    modificacion.resizable(False, False)

    # Create frames for the layout
    left_frame = tk.Frame(modificacion, bg='white')
    left_frame.pack(side=tk.LEFT, fill=tk.Y)
    left_frame.pack_propagate(False)

    right_frame = tk.Frame(modificacion, bg='white')
    right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
    right_frame.pack_propagate(False)

    # Create a style
    style = ttk.Style()
    style.configure('TFrame', background='white')
    style.configure('TLabel', background='white')
    style.configure('TRadiobutton', background='white')

    # Variables to store form data
    revision_number = tk.StringVar()
    new_paragraph_or_fact = tk.StringVar(value='Nuevo Hecho')
    hp_number = tk.StringVar()
    new_proven_fact = tk.StringVar()
    old_proven_fact = tk.StringVar()
    supporting_basis = tk.StringVar()

    # Function to update the constructed message
    def update_message(*args):
        # Obtener los valores de los campos
        revision = revision_number.get()
        hp_num = hp_number.get()
        old_fact = old_fact_text.get().strip()
        new_fact = new_fact_text.get().strip()
        basis = basis_entry.get()

        # Construir el mensaje
        full_message = (
            f"La {revision} revisión fáctica de la parte recurrente, pretende la modificación del HP {hp_num}. "
            f"En la Sentencia de Instancia, el HP {hp_num}º tiene la siguiente redacción:\n\n"
            f"«{old_fact}»\n\n"
            f"La redacción que se propone sería la siguiente:\n\n"
            f"«{new_fact}»\n\n"
            f"Para dicha revisión fáctica, la parte recurrente se apoya en {basis}."
        )

        # Actualizar el widget de texto
        message_text.configure(state='normal')  # Habilita la edición del texto
        message_text.delete(1.0, tk.END)  # Elimina el contenido actual
        message_text.insert(tk.END, full_message)  # Inserta el nuevo mensaje
        message_text.configure(state='disabled')  # Deshabilita la edición del texto

    # Add a Text widget and a Scrollbar widget for the constructed message
    message_text = tk.Text(right_frame, font=("SegoeUI", 9), wrap='word', state='disabled', borderwidth=1,
                           relief='groove', bd=2)
    message_text.pack(side="top", fill="both", expand=True, padx=5, pady=10)
    message_scroll = ttk.Scrollbar(right_frame, orient='vertical', command=message_text.yview)
    message_text['yscrollcommand'] = message_scroll.set
    message_text.bind("<Button-3>", lambda event: mostrar_menu(event, message_text))

    # Dropdown options
    revision_options = ['primera', 'segunda', 'tercera', 'cuarta', 'quinta', 'sexta', 'única']

    # Create form fields
    ttk.Label(left_frame, text="Número de revisión").grid(row=0, column=0, sticky='w', padx=5, pady=2)
    revision_menu = ttk.Combobox(left_frame, textvariable=revision_number, values=revision_options, state="readonly")
    revision_menu.grid(row=1, column=0, sticky='ew', padx=5, pady=2)
    revision_menu.bind('<<ComboboxSelected>>', update_message)

    vcmd = (modificacion.register(solo_numeros), '%P')

    ttk.Label(left_frame, text="Número del HP").grid(row=5, column=0, sticky='w', padx=5, pady=2)
    hp_entry = ttk.Entry(left_frame, textvariable=hp_number, validate="key", validatecommand=vcmd)
    hp_entry.grid(row=6, column=0, sticky='ew', padx=5, pady=2)
    hp_number.trace_add('write', update_message)

    ttk.Label(left_frame, text="Hecho probado antiguo").grid(row=7, column=0, sticky='w', padx=5, pady=2)
    old_fact_text = ttk.Entry(left_frame, textvariable=old_proven_fact)
    old_fact_text.grid(row=8, column=0, sticky='ew', padx=5, pady=2)
    old_proven_fact.trace_add('write', update_message)
    old_fact_text.bind("<Button-3>", lambda event: mostrar_menu(event, old_fact_text))

    ttk.Label(left_frame, text="Hecho probado nuevo").grid(row=9, column=0, sticky='w', padx=5, pady=2)
    new_fact_text = ttk.Entry(left_frame, textvariable=new_proven_fact)
    new_fact_text.grid(row=10, column=0, sticky='ew', padx=5, pady=2)
    new_proven_fact.trace_add('write', update_message)
    new_fact_text.bind("<Button-3>", lambda event: mostrar_menu(event, new_fact_text))

    ttk.Label(left_frame, text="En qué se apoya").grid(row=11, column=0, sticky='w', padx=5, pady=2)
    basis_entry = ttk.Entry(left_frame, textvariable=supporting_basis)
    basis_entry.grid(row=12, column=0, sticky='ew', padx=5, pady=2)
    supporting_basis.trace_add('write', update_message)
    basis_entry.bind("<Button-3>", lambda event: mostrar_menu(event, basis_entry))

    # Function to copy text to clipboard
    def copy_text_to_clipboard():
        # Ensure the text widget is enabled before trying to copy
        message_text.configure(state='normal')
        modificacion.clipboard_clear()  # Clear the clipboard
        text_to_copy = message_text.get("1.0", tk.END)  # Get text from the text widget
        modificacion.clipboard_append(texto_preparado_para_procesador(text_to_copy))  # Append text to the clipboard
        message_text.configure(state='disabled')  # Disable the text widget again

    # Create the copy button
    copy_button = ttk.Button(left_frame, text="Copiar Texto", command=copy_text_to_clipboard)
    copy_button.grid(row=13, column=0, pady=10, padx=5)

    # Pack the Scrollbar widget
    message_scroll.pack(side="right", fill="y")

    # Call update_message on any variable change
    revision_number.trace_add('write', update_message)
    new_paragraph_or_fact.trace_add('write', update_message)
    hp_number.trace_add('write', update_message)
    new_proven_fact.trace_add('write', update_message)
    old_proven_fact.trace_add('write', update_message)
    supporting_basis.trace_add('write', update_message)

    # Inicia la verificación del portapapeles
    verificar_portapapeles()
    center_window_revisor(modificacion)

    # Run the Tkinter event loop
    modificacion.mainloop()

def abrir_ventana_adicion(lugar):
    adicion = tk.Toplevel(root)
    adicion.title("Adición de Hechos Probados")
    adicion.configure(bg='white')

    # Set a minimum size for the window
    adicion.minsize(670, 274)
    adicion.resizable(False, False)

    # Create frames for the layout
    left_frame = tk.Frame(adicion, bg='white')
    left_frame.pack(side=tk.LEFT, fill=tk.Y)
    left_frame.pack_propagate(False)

    right_frame = tk.Frame(adicion, bg='white')
    right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
    right_frame.pack_propagate(False)

    # Create a style
    style = ttk.Style()
    style.configure('TFrame', background='white')
    style.configure('TLabel', background='white')
    style.configure('TRadiobutton', background='white')

    # Variables to store form data
    revision_number = tk.StringVar()
    new_paragraph_or_fact = tk.StringVar(value='Nuevo Hecho')
    hp_number = tk.StringVar()
    new_proven_fact = tk.StringVar()
    supporting_basis = tk.StringVar()

    # Function to update the constructed message
    def update_message(*args):
        # Construye el mensaje
        revision = revision_number.get()
        paragraph_fact = 'adición de un nuevo' if new_paragraph_or_fact.get() == 'Nuevo Hecho' else 'adición de un nuevo párrafo al'
        hp_num = 'HP ' + hp_number.get() + "º"
        new_fact = new_proven_fact.get()
        basis = supporting_basis.get()
        full_message = (
            f"La {revision} revisión fáctica de la parte recurrente, pretende la {paragraph_fact} "
            f"{hp_num}, cuya redacción sería la siguiente:\n\n«{new_fact}»\n\n"
            f"Para dicha revisión fáctica, la parte recurrente se apoya en {basis}."
        )

        # Actualiza el Text widget
        message_text.configure(state='normal')  # Habilita la edición del texto
        message_text.delete(1.0, tk.END)  # Elimina el contenido actual
        message_text.insert(tk.END, full_message)  # Inserta el nuevo mensaje
        message_text.configure(state='disabled')  # Deshabilita la edición del texto

    # Add a Text widget and a Scrollbar widget for the constructed message
    message_text = tk.Text(right_frame, font=("SegoeUI", 9), wrap='word', state='disabled', borderwidth=1,
                           relief='groove', bd=2)
    message_text.pack(side="top", fill="both", expand=True, padx=5, pady=10)
    message_scroll = ttk.Scrollbar(right_frame, orient='vertical', command=message_text.yview)
    message_text['yscrollcommand'] = message_scroll.set
    message_text.bind("<Button-3>", lambda event: mostrar_menu(event, message_text))

    # Dropdown options
    revision_options = ['primera', 'segunda', 'tercera', 'cuarta', 'quinta', 'sexta', 'única']

    # Create form fields
    ttk.Label(left_frame, text="Número de revisión").grid(row=0, column=0, sticky='w', padx=5, pady=2)
    revision_menu = ttk.Combobox(left_frame, textvariable=revision_number, values=revision_options, state="readonly")
    revision_menu.grid(row=1, column=0, sticky='ew', padx=5, pady=2)
    revision_menu.bind('<<ComboboxSelected>>', update_message)

    ttk.Label(left_frame, text="Nuevo párrafo o nuevo hecho").grid(row=2, column=0, sticky='w', padx=5, pady=2)
    ttk.Radiobutton(left_frame, text="Nuevo Hecho", variable=new_paragraph_or_fact, value='Nuevo Hecho').grid(row=3,
                                                                                                              column=0,
                                                                                                              sticky='w',
                                                                                                              padx=15,
                                                                                                              pady=2)
    ttk.Radiobutton(left_frame, text="Nuevo Párrafo", variable=new_paragraph_or_fact, value='Nuevo Párrafo').grid(row=4,
                                                                                                                  column=0,
                                                                                                                  sticky='w',
                                                                                                                  padx=15,
                                                                                                                  pady=2)
    new_paragraph_or_fact.trace_add('write', update_message)

    # Configuración del validador
    vcmd = (adicion.register(solo_numeros), '%P')

    ttk.Label(left_frame, text="Número del HP").grid(row=5, column=0, sticky='w', padx=5, pady=2)
    hp_entry = ttk.Entry(left_frame, textvariable=hp_number, validate="key", validatecommand=vcmd)
    hp_entry.grid(row=6, column=0, sticky='ew', padx=5, pady=2)
    hp_number.trace_add('write', update_message)

    ttk.Label(left_frame, text="Párrafo o Hecho probado nuevo").grid(row=7, column=0, sticky='w', padx=5, pady=2)
    fact_entry = ttk.Entry(left_frame, textvariable=new_proven_fact)
    fact_entry.grid(row=8, column=0, sticky='ew', padx=5, pady=2)
    new_proven_fact.trace_add('write', update_message)
    fact_entry.bind("<Button-3>", lambda event: mostrar_menu(event, fact_entry))

    ttk.Label(left_frame, text="En qué se apoya").grid(row=9, column=0, sticky='w', padx=5, pady=2)
    basis_entry = ttk.Entry(left_frame, textvariable=supporting_basis)
    basis_entry.grid(row=10, column=0, sticky='ew', padx=5, pady=2)
    supporting_basis.trace_add('write', update_message)
    basis_entry.bind("<Button-3>", lambda event: mostrar_menu(event, basis_entry))

    # Function to copy text to clipboard
    def copy_text_to_clipboard():
        # Ensure the text widget is enabled before trying to copy
        message_text.configure(state='normal')
        adicion.clipboard_clear()  # Clear the clipboard
        text_to_copy = message_text.get("1.0", tk.END)  # Get text from the text widget
        adicion.clipboard_append(texto_preparado_para_procesador(text_to_copy))  # Append text to the clipboard
        message_text.configure(state='disabled')  # Disable the text widget again

    # Create the copy button
    copy_button = ttk.Button(left_frame, text="Copiar Texto", command=copy_text_to_clipboard)
    copy_button.grid(row=11, column=0, pady=10, padx=5)

    # Pack the Scrollbar widget
    message_scroll.pack(side="right", fill="y")

    # Call update_message on any variable change
    revision_number.trace_add('write', update_message)
    new_paragraph_or_fact.trace_add('write', update_message)
    hp_number.trace_add('write', update_message)
    new_proven_fact.trace_add('write', update_message)
    supporting_basis.trace_add('write', update_message)

    # Inicia la verificación del portapapeles
    verificar_portapapeles()

    center_window_revisor(adicion)

    # Run the Tkinter event loop
    adicion.mainloop()

def mostrar_menu_derecha(event, lugar):
    menu = tk.Menu(root, tearoff=0)

    # Sección para historial de portapapeles
    for elemento in historial_portapapeles:
        texto_menu = truncar_texto(elemento, MAX_ANCHO_MENU)
        menu.add_command(label=texto_menu, command=lambda e=elemento: pegar_en_texto(e, lugar))

    # Separador
    menu.add_separator()

    # Sección de Censura Jurídica
    menu_censura_juridica = tk.Menu(menu, tearoff=0)
    # Aquí agregas las opciones para Censura Jurídica
    opciones_censura = [("Desestima", "Por ende, se desestima este motivo de censura jurídica."),
                        ("Desestima y confirma", "Por ende, se desestima este motivo de censura jurídica y se confirma la sentencia de instancia."),
                        "SEPARATOR",
                        ("Censura jurídica defectuosa", """	La inexistente explicación de por qué se entienden infringidos los preceptos {#forminput -text ¿Qué se alega como infringido? -type 6 -items Preceptos legales|Preceptos constitucionales -vals legales|contitucionales -variablename Qué se alega como infringido} que cita, no cumple con las exigencias establecidas en el artículo 196.2 in fine LRJS acerca de la justificación de la pertinencia y fundamentación del motivo. 
	Adolece, en consecuencia, el recurso de suplicación de un requisito esencial, cual es que no razona la pertinencia y fundamentación del motivo de manera que se pueda deducir cuál sería el alcance de la infracción y su adecuación al presente supuesto. Siendo así que esta Sala no puede colaborar de oficio en la construcción del recurso, ya que ello atentaría contra el principio de seguridad jurídica y colocaría a la recurrida en indefensión, ni puede conocer, so pena de romper el principio de igualdad entre las partes, de violaciones jurídicas no acusadas por y en el recurso de suplicación o defectuosamente justificadas en orden a su pertinencia y fundamentación jurídica, con la única salvedad -que no es el caso de autos- de que, por afectar al orden público, cupiera actuar de oficio.
	Como señala el Tribunal Constitucional, en su Sentencia 294/1993, de 18 octubre 1993, "el recurso de suplicación no es un recurso de apelación ni una segunda instancia, sino un recurso extraordinario, de objeto limitado, en el que el Tribunal ad quem no puede valorar ex novo toda la prueba practicada ni revisar el derecho aplicable, sino que debe limitarse a las concretas cuestiones planteadas por las partes, en especial la recurrente, que por ello mismo debe respetar una serie de requisitos formales impuestos por la ley y concretados por la jurisprudencia". En la misma línea, la STS de 7 de mayo de 1996 exige que "en el escrito de interposición se expongan con suficiente precisión y claridad, el motivo o los motivos en que se ampare, citándose las normas del ordenamiento jurídico o la jurisprudencia que se consideren infringidas".
	Todo ello implica que el Tribunal de suplicación solo puede examinar aquellas infracciones legales que hayan sido aducidas por la parte recurrente con una adecuada justificación de la pertinencia y fundamentación del motivo, sin que le sea hacedero abordar las infracciones no denunciadas o complementar la justificación de la pertinencia o fundamentación de los motivos. Si esas específicas exigencias no se cumplen, no es viable el recurso o el alegato concreto que adolece de este defectuoso planteamiento. De ahí que, cuando, cual ocurre en el presente caso, la mención que se hace no atiende a esas exigencias, la omisión compromete el derecho de la parte contraria a la defensa y aboca a la Sala a una inadmisible construcción ex officio del recurso, siendo que dicha actividad está reservada a la recurrente.
 	El motivo articulado en el presente caso carece de toda motivación, argumentación o justificación. Desconocemos en qué medida pretende el recurrente anudar la infracción de los {#cond -if {#forminput -text ¿Qué se alega como infringido? -type 6 -items Preceptos legales|Preceptos constitucionales -vals legales|contitucionales -variablename Qué se alega como infringido} == legales -then preceptos legales -else derechos fundamentales -variablename legales o fundamentales} invocados con la pretensión deducida en el suplico del recurso. El recurso adolece de una defectuosa técnica procesal que ha de impedir su examen por esta Sala. no corresponde a jueces y tribunales cubrir las carencias técnico-procesales de las partes. Una cosa es que determinadas deficiencias formales no puedan comportar efectos preclusivos o desestimatorios por meros formalismos; otra, muy distinta, como en el presente supuesto, que el órgano judicial tenga que interpretar el motivo de recurso y constituirse en abogado de parte, en tanto que con dicho proceder se estaría afectando al derecho a la tutela judicial efectiva de la contraparte.
	Consecuentemente, procede la desestimación de este motivo de censura jurídica."""),
                        ("Obiter dicta", """	El motivo suplicacional ha de ser desestimado al centrarse en un "obiter dicta".
	El Tribunal Supremo, en Sentencia de 21 de Septiembre de 2022 (rec. 4018/2021) señala que “Es sabido que las conclusiones constitutivas de "obiter dicta" carecen de virtualidad a los efectos de la contradicción entre sentencias que prevé el art. 219 LRJS, como se deduce de lo declarado en las sentencias de esta Sala de 22 de septiembre del 2005 (rec. nº. 3454/2004), 23 de marzo del 2005 (rec. nº. 5344/2003) y 26 de abril del 2004 (rec. nº. 2098/2003), entre otras. Los "obiter dicta" no pueden fundar la admisión del recurso que nos ocupa por supuesta contradicción de doctrina ya que, la contradicción sólo existe cuando en supuestos de hecho similares con fundamentos y pretensiones sustancialmente idénticos, se dictan sentencias contradictorias, esto es cuando existe una contradicción real y no hipotética. (STS 25/6/2008, rec 2150/2007; 23/9/2008, rec 2370/07)". 
	En este mismo sentido, la Sentencia del Tribunal Supremo de 13 de mayo de 2014 (rec. 109/2013), advierte que "como tal, el obiter no puede ser objeto de impugnación (únicamente se pueden recurrir las decisiones y su "ratio decidendi"), habiendo declarado la Sala muy reiteradamente que "... la equivalencia de resultado a falta de efecto útil de la casación o de sus motivos conduce a su desestimación cuando la hipotética estimación no incidiría en la modificación del fallo de la sentencia recurrida... lo que obviamente es el caso examinado obiter dictum". 
	Esta doctrina es plenamente extensiva al recurso de suplicación. Consecuentemente, procede la desestimación de este motivo de censura jurídica."""),
                        ("Indemnización por daños morales", """Como recientemente ha recordado la STS de 9 de marzo de 2022 (rec. 2269/2019), los daños morales resultan indisolublemente unidos a la vulneración del derecho fundamental, y al ser especialmente difícil su estimación detallada, deben flexibilizarse las exigencias normales para la determinación de la indemnización. 
	En multitud de ocasiones el recurso a la utilización de los elementos que ofrece la cuantificación de las sanciones de la LISOS no resulta, por sí mismo, suficiente para cumplir con relativa precisión la doble función de resarcir el daño y de servir de elemento disuasorio para impedir futuras vulneraciones del derecho fundamental. Ello es debido a que la horquilla de la cuantificación de las sanciones en la LISOS para un mismo tipo de falta (leve, grave, muy grave) resulta ser excesivamente amplía. Por ello, el recurso a las sanciones de la LISOS debe ir acompañado de una valoración de las circunstancias concurrentes en el caso concreto. Aspectos tales como la antigüedad del trabajador en la empresa, la persistencia temporal de la vulneración del derecho fundamental, la intensidad del quebrantamiento del derecho, las consecuencias que se provoquen en la situación personal o social del trabajador o del sujeto titular del derecho infringido, la posible reincidencia en conductas vulneradoras, el carácter pluriofensivo de la lesión, el contexto en el que se haya podido producir la conducta o una actitud tendente a impedir la defensa y protección del derecho transgredido, entre otros que puedan valorarse atendidas las circunstancias de cada caso, deben constituir elementos a tener en cuenta en orden a la cuantificación de la indemnización."""),
                        ("Afectación general", """Por Auto del Tribunal Supremo de 18 de Febrero de 2020 (rec. 2207/19), se señala que “[...] Por lo que se refiere a la "afectación general", hemos de recordar: a) que este supuesto excepcional de interposición del recurso de suplicación "responde a un interés abstracto: la defensa del 'ius constitucionis' y la garantía de la uniformidad de la doctrina legal en todo el territorio nacional como principal expresión del principio constitucional de igualdad en la aplicación de la Ley" […] b) que no puede confundirse con la posible proyección general de un litigio sobre la interpretación de una norma, sino que requiere que "esa proyección se traduzca en un nivel de litigiosidad relevante y actual sobre el problema que se debate".”"""),
                        ("Límites al escrito de impugnación", """Por Sentencia del Tribunal Supremo de 15 de Octubre de 2013, (rec. 1195/2013), se señala que tomando como punto de partida la jurisprudencia precedente y la redacción del art. 197 LRJS, se llega a la convicción de que en el escrito de impugnación del recurso de suplicación el impugnante puede limitarse a oponerse al recurso de suplicación, alegar motivos de inadmisibilidad del recurso, rectificaciones de hechos, o causas de oposición subsidiarias. Pero en modo alguno puede este escrito ser el cauce adecuado para solicitar la anulación o revocación total o parcial de la sentencia impugnada. Conclusión que resulta del tenor literal del precepto, que no establece que en el escrito de impugnación se pueda solicitar la revocación de la sentencia impugnada; de los arts. 202 y 203.1 y 2 LRJS; así del art. 202.3 LRJS no resulta que de estimarse las alegaciones contenidas en el escrito de impugnación, proceda alterar el contenido del fallo; de la propia naturaleza del escrito de impugnación, y porque la jurisprudencia constitucional inspiradora de la reforma legal admite la posibilidad de revisar los hechos probados y aducir nuevos fundamentos jurídicos en el escrito de impugnación, pero siempre limitados a la inadmisión o desestimación del recurso, no a la revocación de la sentencia impugnada."""),
                        ("Documentos nuevos", """La doctrina de la Sala Cuarta respecto a la admisión de nuevos documentos, ex art. 233 LRJS, según recuerda, por todos, el ATS de 30 de septiembre de 2019 (rec. 4947/2018) o el ATS de 19 de abril de 2022 (rcud. 4056/2021), es de este tenor:
	“La doctrina de la Sala, relativa al alcance del art. 233.1 LRJS, viene diciendo lo siguiente:
1. En los recursos extraordinarios de suplicación y casación, incluido el de casación para la unificación de doctrina, los únicos documentos que podrán ser admitidos durante su tramitación serán los que tengan la condición formal de sentencias o resoluciones judiciales o administrativas firmes y no cualesquiera otros diferentes de aquellos.
2. La admisión de dichos documentos viene igualmente condicionada a que: a) Las sentencias o resoluciones hayan sido dictadas o notificadas en fecha posterior al momento en que se llevaron a cabo las conclusiones en el juicio laboral de instancia. b) Que serán admisibles si, además, por su objeto y contenido aparecieran como condicionantes o decisivas para resolver la cuestión planteada en la instancia o en el recurso. y c) En el caso de que no se trate de documentos de tal naturaleza o calidad, deberán ser rechazados de plano, y serán devueltos a la parte que los aportó, sin que puedan por lo tanto ser tenidos en cuenta para la posterior resolución que haya de dictar la Sala.
3. Los documentos que por reunir aquellos requisitos previos hayan sido admitidos y unidos a los autos producirán el efecto pretendido por la parte sólo en el caso de que la producción, obtención o presentación de los mismos no tenga su origen en una actuación dolosa, fraudulenta o negligente de la propia parte que pretende aportarlos; lo cual será valorado en la resolución (auto o sentencia) que proceda adoptar en definitiva.
4. Cuando el documento o documentos aportados reúna todas las anteriores exigencias la Sala valorará en cada caso su alcance en la propia sentencia o auto que haya de dictar”"""),
                        ("Subsanación de defectos en el recurso", """En este caso, habida cuenta de la deficiente técnica suplicacional utilizada, resulta aplicable, la doctrina sentada por el Tribunal Constitucional entre otras, en su Sentencia 230/2000, de 2 de octubre, según la cual el órgano de suplicación está obligado, en aras a la efectividad del derecho a la tutela judicial efectiva, a llevar a cabo una adecuada ponderación de las eventuales irregularidades cometidas en el escrito de interposición del recurso, atendiendo a su entidad y a su incidencia en la consecución de la finalidad perseguida con los requisitos omitidos, así como a su trascendencia para las garantías procesales de la contraparte, lo que implica que el dato al que debe atender a la hora de adoptar su decisión no sea la forma o técnica del escrito sino su contenido, no debiendo rechazar "a limine" su estudio cuando de forma suficientemente precisa exponga los hechos o argumentos que estime erróneos y cuáles los que debieran ser tenidos por correctos, como ha hecho la parte recurrente para oponerse a la resolución de instancia."""),
                        ("Causa de inadmisión es causa de desestimación", """El Tribunal Supremo, en su Sentencia 696/2022, de 26 de Julio, en atención a una causa de inadmisión del recurso señala que: 
“La posibilidad de que un recurso admitido a trámite finalice con una resolución mediante la cual se concluye que concurre una causa de inadmisión es acorde con nuestra jurisprudencia. Recordemos que cualquier causa que pudiese motivar en su momento la inadmisión del recurso, una vez que se llega a la fase de sentencia queda transformada en causa de desestimación.
Se trata de una consecuencia que no puede considerarse contraria al derecho a la tutela judicial efectiva. Con arreglo a reiterada jurisprudencia constitucional "La comprobación de los presupuestos procesales para la viabilidad de la acción puede volverse a abordar o reconsiderarse en la sentencia, de oficio o a instancia de parte, dando lugar, en su caso, a un pronunciamiento de inadmisión por falta de tales presupuestos" (Por todas, STC 200/2012, de 12 de noviembre).”"""),
                        ("Introducción de cuestiones nuevas", """En relación con esta cuestión, hemos de recordar, entre otras, la Sentencia del Tribunal Supremo 104/2022, de 02 de Febrero de 2022 que señala “Como regla general, en los recursos extraordinarios de suplicación y casación no pueden examinarse las denominadas "cuestiones nuevas" porque el objeto de los citados recursos no consiste en volver a enjuiciar la cuestión de fondo sino en examinar la corrección de la sentencia de instancia a la vista de los motivos previstos en la LRJS, lo que excluye que puedan suscitarse cuestiones que no fueron alegadas ni examinadas por el órgano judicial de instancia.”"""),
                        ("Petición de principio", """El recurso no puede cuestionar, restar, adicionar o partir de una realidad diversa de la judicialmente apreciada. De lo contrario incurriría en un defecto procesal que lo abocaría al fracaso. Al construir su razonamiento sobre premisas que no coinciden con la crónica judicial incurriría en una petición de principio, haciendo supuesto de determinada cuestión. Al actuar de ese modo se propicia el vicio procesal de la llamada "petición de principio" o "hacer supuesto de la cuestión", que se produce cuando se parte de premisas fácticas distintas a las de la resolución recurrida. Este defecto se produce cuando el recurrente parte de unas premisas fácticas distintas a las que declara probadas la resolución recurrida (STS 141/2021 de 2 febrero, rec. 128/2019 y las citadas en ella). Esta circunstancia puede ocurrir en dos casos, cuando el recurso parte de hechos nunca declarados probados, cuando el recurso parte de hechos probados cuya adición no se ha estimado."""),
                        ("Perspectiva de género", """En numerosas ocasiones se ha advertido por el Tribunal Supremo que ninguna duda cabe sobre la obligación de jueces y tribunales de incorporar la perspectiva de género en lo que constituye su actuación como Poder del Estado, esto es, en la interpretación y aplicación de las normas. Así lo ha sostenido, por ejemplo, en las SSTS de 26 septiembre 2018 (rcud. 1352/2017), 13 noviembre 2019 (rcud. 75/2018) y 23 de Junio de 2022 (rcud. 646/2021)."""),
                        ("Incapacidad sin revisión fáctica", """Tiene esta Sala reiteradamente establecido que nunca podrá prosperar una censura jurídica si no varía el relato de hechos probados de resolución recurrida cuando entre una y otra dimensión de la sentencia exista una íntima relación de ambos presupuestos, doctrina que resulta de aplicación a supuestos como el que aquí se enjuicia, en los que la censura jurídico-sustantiva tendría como presupuesto necesario la modificación de la narración fáctica. Las limitaciones funcionales consignadas en la narración fáctica no han sufrido modificación alguna por el cauce procesal adecuado, siendo compatible/incompatile, como consideró el juzgador a quo, con actividades laborales de carácter retribuido."""),
                        ("No indemnización conciliación", """	En relación con la indemnización por daños y perjuicios derivados de la vulneración de derechos fundamentales, el Tribunal Supremo en sentencia de 25 de mayo de 2023 (rec. 1602/2020) señala lo siguiente:
“Esta Sala ha venido sosteniendo que las medidas normativas dirigidas a la conciliación de la vida laboral y familiar de las personas trabajadoras, ya desde la perspectiva de no discriminación por razón de sexo o por circunstancias personales, así como desde la protección a la familia e infancia, que proclama nuestra constitución, gozan de una dimensión constitucional de forma que, a la hora de proceder a la interpretación de las mismas deben tenerse presente esos derechos fundamentales para lo cual han de ponderarse todas las circunstancias concurrentes para poder cohonestar los intereses en juego, tal y como indica el citado art. 37.7 del ET y el art. 139 de la LRJS cuando dispone que las discrepancias entre empresa y trabajador deberán solventarse llevando al proceso judicial las propuestas y alternativas para la concreción del horario.
Ahora bien, ello no significa que toda decisión sobre concreción horaria implique necesariamente un trato discriminatorio por razón de sexo.
Es cierto que la STC 233/2007 y otras posteriores afirmaron que la restricción de estos derechos o la atribución de efectos laborales desfavorables a su ejercicio "conecta íntimamente con la prohibición de discriminación por razón de sexo de las trabajadoras" y que "constituye discriminación indirecta el tratamiento formalmente neutro o no discriminatorio, del que se deriva, por las diversas condiciones fácticas que se dan entre los trabajadores de uno y otro sexo, un impacto adverso sobre los miembros de un determinado sexo. Obviamente, salvo que este tratamiento responda a una finalidad legítima y utilice medios proporcionados, adecuados y necesarios para conseguirla" y que "la incorporación de la discriminación indirecta como contenido proscrito por el art. 14 CE repercute en la forma de abordar el análisis de este tipo de discriminaciones, pues deberá atenderse necesariamente a los datos revelados por la estadística; en su caso a datos estadísticos actualizados de oficio por el tribunal ( STC 253/2004, FJ 8).
Y así lo reitera la STC 153/2021, en relación con la conexión intima con la prohibición de discriminación por razón de sexto de las trabajadoras cuando ven restringidos sus derechos a la conciliación de la vida familiar y laboral pero ello no significa que no deba examinarse si dicha trabajadora ha sido objeto de dicho trato discriminatorio directo o indirecto. Y en el caso que afronta considera que aunque la medida sea neutra -en aquel caso se trataba de una reducción de jornada y concreción horaria que la empresa otorgó cambiando de servicio a la trabajadora- en cuanto que no se produce por ser mujer, afirma que " No obstante, siguen siendo las mujeres las que se acogen mayoritariamente a dicha medida de conciliación, tal y como apreciamos recientemente en la STC 71/2020, FJ 4 y reitera ahora el Ministerio Fiscal. En efecto, según datos recientes del Instituto Nacional de Estadística (en adelante, INE), en 2018 se acogieron a la reducción de jornada por guarda legal el 1,10 por 100 de las mujeres trabajadoras (146300 trabajadoras, en términos absolutos), frente al 0,5 por 100 de los hombres trabajadores (68900 trabajadores); en definitiva, más del doble de mujeres que de hombres (INE, "Encuesta de población activa. Módulo sobre conciliación entre la vida laboral y la familiar. Año 2018"). Ello pese a que la tasa de empleo de las mujeres es aun significativamente inferior a la de los hombres: ese mismo año la tasa de empleo de las mujeres era de 43,1 por 100, frente al 54,2 por 100 la de los hombres, con una brecha de género de 11,1 por 100 (INE, "Tasa de empleo. Brecha de género 2018"). Y añade que " Estos datos confirman que, aunque pueda haber habido un ligero avance, sigue siendo una realidad la conclusión que extrajimos en la STC 3/2007: cualquier menoscabo de este derecho perjudica fundamentalmente a las mujeres. Por consiguiente, y en tanto en cuanto no se alcance en nuestra sociedad un reparto equilibrado entre mujeres y hombres de las tareas de cuidado familiar que coadyuve a la consecución de la igualdad en el ámbito laboral, incurre en discriminación indirecta por razón de sexo el tratamiento que implique una restricción o la asignación de consecuencias laborales negativas al ejercicio por las mujeres trabajadoras de estos derechos de conciliación de la vida laboral y familiar, siempre que no pueda probarse que responden a razones o factores objetivos, ajenos a toda discriminación por razón de sexo, susceptibles de legitimar la medida en cuestión".
No obstante, el TC, al valorar las circunstancias del caso niega que la decisión empresarial suponga para la trabajadora un trato peyorativo en las condiciones de trabajo o limitaciones de sus expectativas profesionales.
Y también descarta que exista una discriminación indirecta, que solo se desvanece cuando el tratamiento responda a una finalidad legitima y se utilicen medios proporcionados, adecuados y necesarios para conseguirla, porque en ese caso se ofrecieron razones objetivas, vinculadas a los requerimientos de los servicios a atender, lo que elimina la existencia de trato discriminatorio por alguno de los factores del art. 14 CE.
Pues bien, bajo esas premisas doctrinales. en el presente caso, en orden a los derechos fundamentales, la sentencia recurrida como la de contraste, como viene diciendo la doctrina constitucional y la jurisprudencia, parten de la necesidad de que la denuncia de su vulneración deba ir acompañada de la presentación de indicios que pongan de manifiesto que la decisión empresarial impugnada pretende desconocer y transgredir el derecho fundamental cuestionado y ante ello es la empresa la que debe acreditar que su decisión es ajena a ese propósito. A tal efecto y, en relación con la materia que nos ocupa, del derecho a la conciliación de la vida familiar al que responde la concreción horaria del art. 37.7 del ET, y al margen de las circunstancias que deban valorarse a la hora de determinar o ponderar los derechos de las partes en el ejercicio de los respectivos derechos en el ámbito de la legalidad ordinaria, resulta que así como la sentencia recurrida aprecia la vulneración del derecho a no discriminación por razón de sexo para la referencial esos mismos hechos no tiene ese alcance y esta última conclusión es la que esta Sala entiende que es la ajustada a derecho.
La mera denegación de la concreción horaria que interesa la persona trabajadora, con indicación de las causas que lo impiden no implica, por si sólo, que se esté vulnerando el derecho de no discriminación por razón de sexo ni siquiera por discriminación indirecta.”
	En síntesis, la Sala sostiene que las medidas normativas dirigidas a la conciliación de la vida laboral y familiar de las personas trabajadoras tienen una dimensión constitucional y deben ser interpretadas teniendo en cuenta los derechos fundamentales. Sin embargo, esto no significa que toda decisión sobre concreción horaria implique necesariamente un trato discriminatorio por razón de sexo. Se argumenta que aunque las mujeres son mayoritariamente las que se acogen a estas medidas, no todas las decisiones sobre horarios implican trato discriminatorio. La discriminación indirecta es aquella en la que, a pesar de no ser explícita, afecta desproporcionadamente a un género y no responde a necesidades legítimas o proporcionadas. Se cita una sentencia del Tribunal Constitucional en la que se analiza un caso en el que se niega que la decisión empresarial implique trato peyorativo o discriminación indirecta, dado que se ofrecieron razones objetivas vinculadas a los requerimientos de los servicios. La discriminación indirecta se da cuando un tratamiento aparentemente neutro tiene un impacto adverso en un determinado sexo, salvo que responda a una finalidad legítima y utilice medios proporcionados, adecuados y necesarios para conseguirla. La jurisprudencia establece que para denunciar una vulneración de derechos fundamentales, como la no discriminación por razón de sexo, se deben presentar indicios que demuestren que la decisión empresarial busca desconocer y transgredir dicho derecho fundamental. El Tribunal Supremo concluye que la denegación de la concreción horaria solicitada por la trabajadora, aunque aparentemente perjudique a las mujeres, no implica en sí misma la vulneración del derecho a no discriminación, incluso indirecta.
	En la presente causa, no hay dato del que obtener que la denegación de la concreción horaria que solicitaba la parte actora, cambiando el turno de tarde al de mañana y dejar de trabajar en fines de semana, se base en un factor relacionado con el sexo - por el hecho de ser mujer- sino que, si acaso, se estaría -siguiendo aquella doctrina constitucional- ante una denegación neutra -afectante tanto a hombres como mujeres- aunque, si tenemos en consideración la realidad social que ya ha tomado en cuenta la doctrina constitucional, existente precisamente en el tiempo en el que aquí se interesaron las concreciones horarias por la trabajadora, esa denegación podría valorarse como indicio de discriminación indirecta, lo que, en el presente caso, tampoco concurriría.
	En lo que se refiere a la perspectiva de género, como criterio interpretativo, no hay razón alguna para que deba justificar una solución como la alcanzada en la sentencia recurrida en tanto que lo que se está cuestionando es si la denegación de la concreción horaria ha tenido como real base la de discriminar a la trabajadora y ello implica valorar hechos y no interpretar las normas en juego.
"""),
                        ]
    # Dividir las opciones de Revisión Fáctica en dos listas
    indice_separador = opciones_censura.index("SEPARATOR")
    opciones_antes_separador = opciones_censura[:indice_separador]
    opciones_despues_separador = opciones_censura[indice_separador + 1:]

    # Ordenar las opciones por conteo de uso
    opciones_despues_separador_ordenadas = sorted(opciones_despues_separador, key=lambda x: conteo_uso_opciones.get(x[1], 0), reverse=True)

    # Añadir las opciones ordenadas al menú
    for opcion, texto in opciones_antes_separador:
        menu_censura_juridica.add_command(label=opcion, command=lambda t=texto: pegar_en_texto(t, lugar))

    # Añadir el separador
    menu_censura_juridica.add_separator()

    # Añadir las opciones después del separador
    for opcion, texto in opciones_despues_separador_ordenadas:
        menu_censura_juridica.add_command(label=opcion, command=lambda t=texto: pegar_en_texto(t, lugar))
        
    menu.add_cascade(label="Fundamentación Jurídica", menu=menu_censura_juridica)

    # Sección de Revisión Fáctica
    menu_revision_factiva = tk.Menu(menu, tearoff=0)
    # Aquí agregas las opciones para Revisión Fáctica
    opciones_revision = [("Procede", "La documental en la que se apoya es literosuficiente y dada la posible trascendencia de tal circunstancia, ora en esta instancia, ora en otra superior, se admite la revisión planteada."),
                         ("No procede", "Expuesto lo que antecede, no procede la estimación de la revisión fáctica propuesta por la parte recurrente."),
                         # Separador entre "Procede" y "Hecho Indirecto"
                        "SEPARATOR",
                         ("Hecho indirecto", "Así, lo que se quiere incluir en el factum de la sentencia no son hechos directos, sino indirectos. El relato fáctico debe contener la convicción judicial sobre los hechos controvertidos o necesitados de prueba (art. 97.2 en relación con el art. 90.1, ambos de la LJRS) Es decir, la versión judicial de lo ocurrido y no el contenido de los medios de prueba (los informes) que es lo que se quiere introducir. La inclusión en los hechos probados de hechos indirectos induce a confusión por falta de claridad, infringiendo así el mandato que contiene el art. 218.1 LEC, que es instrumental respecto al deber constitucional de motivación ex art. 24 y 120.3 CE, pues no se sabe a ciencia cierta si da por acreditado un hecho o el contenido de un medio de prueba, que es cosa bien distinta."),
                         ("Hechos negativos", "Ahora bien, pretendiéndose la adición, en relación a la ausencia de acreditación de un determinado extremo, de un hecho negativo, no ha lugar a aquélla, conforme a reiterada doctrina jurisprudencial (STS de 21 de enero de 2021, rec. 158/2019)."),
                         ("Obstrucción negativa", "La petición ha de ser denegada por cuanto utiliza la llamada “obstrucción negativa”, es decir, se limita a decir que el hecho no está probado, o no está suficientemente probado. Así, la “obstrucción negativa”, carece de eficacia revisora en suplicación, dadas las amplias facultades que el artículo 97.2 LRJS otorga al Juzgador a quo para la apreciación de los elementos de convicción, que es un concepto más extenso que el de medios de prueba pues no sólo abarca a los que enumera el art. 299 LEC, sino también el comportamiento de las partes en el transcurso del proceso e incluso sus omisiones, de manera que, en atención a lo expuesto, no ha de tener éxito la pretensión de revisión a que se contrae el motivo del recurso entablado frente a la resolución de instancia."),
                         ("Hecho conclusivo-valorativo", "La modificación propuesta presenta un claro signo conclusivo-valorativo más que meramente fáctico, resultando de la interpretación de los documentos que hace la parte recurrente, debiendo recordarse a este respecto que la prueba ha de ser fehaciente, es decir, debe reflejar la verdad por sí sola, y con ello el error del Juzgador, sin otras consideraciones colaterales, hipótesis o conjeturas, porque, si éstas se admitieran, la Sala se encontraría suplantando al Juez de lo Social en la valoración de la prueba, como si de un nuevo juicio se tratara (una mera apelación) y no resolviendo un recurso que tiene naturaleza extraordinaria."),
                         ("Hecho predeterminante del fallo", "Como señalan, entre otras, las Sentencia del Tribunal Supremo de 11 y 19 de Junio de 1985, “hecho predeterminante del fallo” no equivale a hecho trascendente para el resultado del pleito a favor de uno u otro litigante, sino a un hecho descrito en forma no meramente fáctica, sino que incorpora la regla jurídica determinante de la solución. Lo definitorio es que un supuesto hecho implique en sí mismo la solución del caso discutido. O como ha expresado en otras palabras el Tribunal Supremo, en Sentencias de 19 de diciembre de 2012 y 8 de abril de 2014, las afirmaciones predeterminantes del fallo que no pueden figurar en el relato de hechos probados son aquellas que implican la previa celebración de un juicio de valor, de una calificación jurídica que debe hacerse en la fundamentación de derecho, pero no las que describen un hecho cuya probanza pueda ser determinante del fallo."),
                         ("Testifical Documentada", "Como señala el Tribunal Supremo, Sentencia de 24 de enero de 2020 (rec. 3962/2016), es doctrina de esta Sala, que los escritos en que se reflejan manifestaciones de terceros no pueden ser considerados, en principio, como prueba documental a los efectos de fundamentar la revisión fáctica en suplicación a que se refiere el art. 193.b) LRJS , '... ni tampoco el error de hecho en casación ordinaria (art. 207.d) LRJS), al no tratarse de un auténtico documento sino de meras manifestaciones testimoniales formuladas por escrito que por ello no pierden su verdadero carácter de prueba testifical o de una denominada prueba testifical impropia, que solo habría adquirido todo su valor procesal como tal prueba testifical de haber sido ratificada en juicio por sus firmantes, cuya valoración queda a la libre apreciación del juzgador de instancia, como se deduce, además, palmariamente de la redacción literal de los preceptos procesales reguladores...'  (STS citada de 15 de octubre de 2014 (rcud. 1654/2013)."),
                         ("Literosuficiencia", """El Tribunal Supremo en su Sentencia 4/2015, de 10 de julio, define los documentos literosuficientes como aquellos que basten por sí mismos para llegar a la conclusión acreditativa que se pretende, evidenciando el objeto de prueba sin necesidad de acudir a otras fuentes probatorias o a complejos desarrollos argumentales. El documento o documentos en que se base la petición de revisión deben gozar de literosuficiencia, esto es: "[...] que de los documentos transcritos se desprenda, sin necesidad de razonamientos o conjeturas, el pretendido error del Juzgador [...]" (STS de 22 de junio de 2022, rec. 15/2022), requisitos que no se dan cita en este caso."""),
                         ("Autosuficiencia", """Tradicionalmente, en los recursos extraordinarios, la apreciación del error de hecho fundado en prueba documental requería la concurrencia del requisito de autonomía, que se plasmaba en los requisitos de “autosuficiencia” (que el o los documentos hicieran prueba por sí solos, sin necesidad de acudir a otros elementos probatorios), y de “literosuficiencia” (que el o los documentos hicieran prueba por si solos, sin necesidad de acudir a hipótesis o conjeturas), así lo señalaban las Sentencias de la Sala Primera del Tribunal Supremo de 22 de Octubre y 14 de Diciembre de 1979. Al respecto, la Sala de lo Social del Tribunal Supremo ha exigido, para atribuir virtualidad revisora en casación a la prueba documental, que el error de hecho denunciado quede evidenciado, de forma patente, clara y directa, por los propios documentos alegados, sin necesidad de llevar a cabo deducciones o cálculos, ni conjeturas o suposiciones. La mentada doctrina, aplicable al recurso extraordinario de suplicación, hace hincapié en el requisito de evidencia del error, pero conectándolo con el principio de autonomía, de forma que es preciso que el error probatorio se infiera de los propios documentos, por sí solos. Por ello, hemos de reiterar que el éxito de la revisión histórica fundada en prueba documental en el recurso extraordinario de suplicación requiere que el o los documentos invocados por la parte recurrente acrediten por sí solos, sin necesidad de acudir a hipótesis o conjeturas, el error fáctico de instancia."""),
                         ("Valoración prueba por el juez a quo", "Como reiteradamente se ha puesto de manifiesto, el Juzgador/a ostenta una amplia facultad para valorar todo el material probatorio practicado en la instancia, de modo que puede obtener y deducir una interpretación distinta a aquella que obtiene la parte, ya que, ante posibles contradicciones debe prevalecer el criterio del órgano jurisdiccional, que actúa de manera imparcial y objetiva frente al interés de una parte, correspondiendo al Juzgador/a la facultad privativa sobre la valoración de todas las pruebas aportadas al proceso, de acuerdo con el artículo 97.2 de la LRJS (STS 18/11/1999)."),
                         ("Valoración de la prueba en otro sentido", """En este sentido el Tribunal Constitucional ha declarado que el derecho a la tutela judicial efectiva no ampara la valoración de la prueba conforme a las pretensiones ejercitadas por la recurrente, sino el derecho de los litigantes a “una valoración de la prueba que no sea manifiestamente irrazonable o totalmente infundada, absurda o notoriamente errónea” (STC nº 484/1984 de 26 de julio), y que únicamente cabe apreciar indefensión por la valoración de la prueba cuando exista “una defectuosa utilización de las reglas rectoras de la carga de la prueba” (STC nº 140/1994 de 9 de mayo), o “por prescindir de la contemplación racional de la prueba de una de las partes” (STC nº 63/1993 de 1 de marzo), circunstancias que no concurren en el presente caso, en que el juzgador de instancia ha valorado la prueba aportada."""),
                         ("Valoración basada en testifical/interrogatorio", """La recurrente ampara la revisión en la testifical/el interrogatorio, lo que no es admisible pues la revisión de hechos probados sólo tiene cabida en virtud de prueba documental o pericial, según los artículos 193.b) y 196.3 LRJS, mientras que la valoración probatoria del interrogatorio de parte corresponde exclusivamente al magistrado de instancia, conforme al principio de inmediación y a lo establecido en el artículo 97.2 de dicha Ley. 
	Así pues, la Sentencia del Tribunal Supremo de 25 marzo 2014 (rec. 161/2013), en esta misma línea ha rechazado que la modificación fáctica pueda ampararse en la prueba testifical o en el interrogatorio de parte, por cuanto el magistrado de instancia ya valoró esta prueba en conjunción con el resto de la practicada. En efecto es reiterada la doctrina judicial que declara que la prueba testifical o de interrogatorio, por ser de libre valoración por el juez a quo, no es controlable ni revisable por la Sala, al ser de libre valoración con arreglo a la sana crítica, dada la naturaleza extraordinaria del recurso de suplicación, que no es un recurso de apelación, por lo que la Sala no puede analizar la prueba de interrogatorio practicada."""),
                         ("Documentos electrónicos", "De conformidad con la tesis sentada por el Tribunal Supremo, a partir de su Sentencia 706/2020, de 23 de Julio, la LEC contiene preceptos favorables al concepto amplio de prueba documental, comprensiva de la electrónica. El avance tecnológico ha hecho que muchos documentos se materialicen y presenten a juicio a través de los nuevos soportes electrónicos, lo que no debe excluir su naturaleza de prueba documental, con las necesarias adaptaciones (por ejemplo, respecto de la prueba de autenticación). Si no se postula un concepto amplio de prueba documental, llegará un momento en que la revisión fáctica casacional quedará vaciada de contenido si se limita a los documentos escritos, cuyo uso será exiguo."),
                         ("Folios de whatsapp", """La postura favorable a la revisión fáctica suplicacional basada en dichos medios de prueba debería sustentarse en la citada STS de 29 de enero de 2019.
	Conforme a dicha tesis, si una parte procesal aporta unos folios impresos, afirmando que se corresponden con unos correos electrónicos o mensajes SMS o WhatsApp y la parte procesal contraria solicita la revisión fáctica suplicacional basada en dichos folios, podrían tener eficacia revisora en suplicación.
	Por el contrario, si la misma parte procesal que aportó dichos folios es quien postula la revisión fáctica suplicacional basada en ellos, en tal caso el TSJ deberá comprobar si se ha impugnado su autenticidad. En caso de que la parte contraria haya impugnado la autenticidad de esos folios y no se haya practicado prueba de autenticación, en ningún caso tendrán eficacia revisora. Pero si la parte contraria no ha impugnado su autenticidad, conforme a los citados pronunciamientos del TS favorables a atribuirles eficacia revisora casacional, no habría obstáculo para que pudieran tener eficacia revisora suplicacional."""),
                         ("Revisión fáctica sin censura jurídica", """La Sentencia del Tribunal Supremo de 3 de noviembre de 2020 (rec. 42/2019), recuerda una reitera doctrina relativa al recurso de casación en la que se pone de manifiesto que dicho recurso debe estar fundado en un motivo de infracción de las normas del ordenamiento jurídico, no siendo viable uno que solo persiga alterar el relato de los hechos declarados probados en la sentencia recurrida.
En el presente caso, la parte recurrente no formula denuncia jurídica alguna, siendo así que la revisión fáctica es meramente instrumental de la denuncia jurídica.
En efecto, señala dicha sentencia “acerca del cumplimiento del requisito consistente en fundamentar la infracción legal denunciada señalando que el recurso de casación, como extraordinario que es, debe estar fundado en un motivo de infracción de ley; y esta exigencia no se cumple con sólo indicar los preceptos que se consideren aplicables, sino que es requisito ineludible para su correcta observancia razonar de forma expresa y clara sobre la pertinencia y fundamentación del recurso en relación con la infracción o infracciones que son objeto de denuncia. Como se afirmó en la STS de 18 de octubre de 2007 (Rec. 110/2006), uno de los requisitos esenciales del recurso de casación, como recurso extraordinario, es que el mismo se fundamente, al menos en un motivo de violación legal o jurisprudencial de los señalados en el actual artículo 207 e) LRJS. Esa necesidad de concreción y precisión de la norma que se considera infringida era exigida ya en el artículo 1707 LEC anterior; y en LEC 2000 que en su artículo 479.3 y 4 prescribió que el recurso deberá indicar la infracción legal que se considere cometida. El requisito de fundamentación de la infracción legal cometida en la sentencia impugnada, además de derivar directamente del artículo 210 LRJS, es una consecuencia lógica del carácter casacional del recurso, puesto que, sin ella, se transferiría a la Sala, en contra del principio de equilibrio procesal, el examen de oficio del ajuste de la sentencia recurrida a la legalidad. El recurso de casación, por su carácter de extraordinario no permite su viabilidad con el simple apoyo de que la resolución impugnada perjudica al recurrente, ni tampoco puede fundarse, exclusivamente, en la modificación de los hechos probados de la sentencia recurrida, pues tal modificación, por relevante que sea, no permite modificar el fallo de la sentencia, sin una alegación -precisa y clara- de las normas o de la jurisprudencia infringidas que posibilitarían analizar la aplicación del derecho efectuado por la sentencia recurrida y, en su caso, su casación y anulación”.
Y sigue diciendo que “La omisión de la infracción legal denunciada o de su fundamentación ha sido calificada reiteradamente como insubsanable ( STS de 14 de noviembre de 2003, Rcud. 4641/2002 y las que en ella se citan y las que le han seguido), de forma que la falta de este requisito exime a la Sala del conocimiento del fondo del litigio.
La descrita exigencia no pretende aplicar al recurrente un rigorismo puramente formal, que sería contrario al derecho a la tutela judicial efectiva que consagra el art. 24.1 CE, sino, bien al contrario, garantizar ese mismo derecho a los recurridos, en tanto que la estimación de un recurso que ha sido defectuosamente planteado supondría dejar en indefensión a la parte contraria, en cuanto obligaría al Tribunal a adoptar postura de parte para subsanar de oficio los defectos en los que haya podido incurrir el recurrente, viéndose obligado a identificar las normas legales que no han sido invocadas en el recurso cuya posible infracción pudiera dar lugar a casar la sentencia, y al hilo de ello construir los argumentos jurídicos que conducirían a su vulneración que no fueron articulados por la recurrente, privando de esta forma a la recurrida de la posibilidad de defenderse de unas alegaciones que ni tan siquiera habían sido esgrimidas en el escrito de recurso ( STS de 26 de febrero de 2020, Rec. 160/2019)”.
Esta doctrina es reiterada por la reciente Sentencia del Tribunal Supremo de 12 de enero de 2023 (rec. 18/2021), en la que se señala que “La aplicación de nuestra doctrina en el presente recurso, de conformidad con lo informado por el Ministerio Fiscal, nos lleva a la conclusión de que el mismo incurre en causa de inadmisión que nos lleva, en este momento procesal, a su desestimación porque el escrito de interposición se destina en exclusiva a la revisión de uno de los hechos declarados probados sin justificar en modo alguno en qué infracción normativa ha podido incurrir la sentencia de instancia para poder alterar el signo del fallo.”
Por todo ello, siendo de aplicación lo señalado al caso presente, siendo claro que el recurso de la demandante no cumple con las exigencias que impone el art. 196.2 LRJS, ya que no contiene motivo de infracción jurídica ni cita de las normas del ordenamiento jurídico o la jurisprudencia que se consideren infringidas, el recurso no puede prosperar porque el fallo está en el defectuoso planteamiento del recurso."""),
                         ("Documentos contradictorios", """En este sentido, el TC ha sostenido que cuando existan documentos contradictorios, en la medida que de los mismos puedan extraerse conclusiones contradictorias e incompatibles, debe prevalecer la solución fáctica de instancia, al ser el órgano judicial a quo el órgano jurisdiccional soberano para la apreciación de la prueba (SSTC n.º 444/1989, de 20 de febrero y n.º 24/1990, de 15 de febrero). El TS ha reiterado este criterio, con la salvedad de que la apreciación de instancia tiene que ser razonable (STS de 11 de junio de 1986).
Más recientemente el TS ha insistido en que la revisión fáctica casacional solo procederá cuando haya un error evidente (SSTS de 8 de noviembre de 2017, recurso 134/2017; 16 de enero de 2018, recurso 262/2016 y 17 de julio de 2018, recurso 170/2017).
En los citados pronunciamientos subyace un principio de respeto de la valoración probatoria de instancia, salvo que se acredite cumplidamente (evidentemente) la existencia de error, de forma que cuando existen dos medios de prueba contradictorios con un valor probatorio semejante, si el juez de instancia ha otorgado credibilidad a uno de ellos y la parte que recurre invoca el otro, el TSJ tiende a respetar la valoración probatoria de instancia."""),
                         ("Pericial contradictoria", """En este sentido, la doctrina del Tribunal Supremo, entre otras: SSTS 12 marzo, 3, 17 y 31 de mayo, 21 y 25 de junio y 17 de diciembre de 1990 y 25 de enero de 1991, establece que ante dictámenes médicos contradictorios, excepto la concurrencia de circunstancias especiales, se ha de atender a la valoración realizada por el juzgador de instancia en virtud de las competencias que le asignan el art. 97.2 LRJS, 218.2 LEC y 120.3 CE.
En los citados pronunciamientos subyace un principio de respeto de la valoración probatoria de instancia, salvo que se acredite cumplidamente (evidentemente) la existencia de error, de forma que cuando existen dos medios de prueba contradictorios con un valor probatorio semejante, si el juez de instancia ha otorgado credibilidad a uno de ellos y la parte que recurre invoca el otro, el TSJ tiende a respetar la valoración probatoria de instancia."""),
                         ]
    # Dividir las opciones de Revisión Fáctica en dos listas
    indice_separador = opciones_revision.index("SEPARATOR")
    opciones_antes_separador = opciones_revision[:indice_separador]
    opciones_despues_separador = opciones_revision[indice_separador + 1:]

    # Ordenar las opciones por conteo de uso
    opciones_despues_separador_ordenadas = sorted(opciones_despues_separador, key=lambda x: conteo_uso_opciones.get(x[1], 0), reverse=True)

    # Añadir las opciones ordenadas al menú
    for opcion, texto in opciones_antes_separador:
        menu_revision_factiva.add_command(label=opcion, command=lambda t=texto: pegar_en_texto(t, lugar))

    menu_revision_factiva.add_separator()

    menu_revision_factiva.add_command(label="Adición", command=lambda: abrir_ventana_adicion(lugar))
    menu_revision_factiva.add_command(label="Modificación", command=lambda: abrir_ventana_modificacion(lugar))


    # Añadir el separador
    menu_revision_factiva.add_separator()

    # Añadir las opciones después del separador
    for opcion, texto in opciones_despues_separador_ordenadas:
        menu_revision_factiva.add_command(label=opcion, command=lambda t=texto: pegar_en_texto(t, lugar))


    menu.add_cascade(label="Revisión Fáctica", menu=menu_revision_factiva)

    # Mostrar menú
    try:
        menu.tk_popup(event.x_root, event.y_root)
    finally:
        menu.grab_release()

# Función de validación
def solo_numeros(P):
    # Permite solo si el texto es un número o está vacío (para permitir borrado)
    return P.isdigit() or P == ""

# Función de recordar el tamaño de la ventana
# Ruta del archivo de configuración
config_file = 'window_config.json'

def save_config():
    if root.state() == 'zoomed':
        config = {'maximized': True}
    else:
        config = {
            'size': '1300x800',
            'maximized': False
        }
    with open(config_file, 'w') as file:
        json.dump(config, file)

def on_closing():
    save_config()
    root.destroy()

def load_config():
    if os.path.exists(config_file):
        with open(config_file, 'r') as file:
            return json.load(file)
    else:
        return None

root = tk.Tk()
root.title('Asistente a la Suplicación')
root.geometry("1300x800")  # Establece el tamaño de la ventana a 1300x800
root.iconbitmap('icon.ico')


# todo blanco
style = ttk.Style()
style.configure('White.TFrame', background='white')
style.configure('White.TButton', background='white')
style.configure('White.TLabel', background='white')
style.configure('White.TLabelframe', background='white')
style.configure('White.TLabelframe.Label', background='white')  # Para el título del LabelFrame
style.configure('White.TCheckbutton', background='white')
style.configure('White.TRadiobutton', background='white')

# Agrega más configuraciones según sea necesario


# Cargar la configuración guardada
config = load_config()
if config:
    if config['maximized']:
        root.state('zoomed')
    else:
        root.geometry(config['size'])

# Configurar el evento de cierre para guardar la configuración de tamaño
root.protocol("WM_DELETE_WINDOW", on_closing)

# Left frame for inputs
left_frame = ttk.Frame(root, width=200, style='White.TFrame')
left_frame.grid(row=0, column=0, sticky="ns")

# Right frame for text display
right_frame = ttk.Frame(root, style='White.TFrame')
right_frame.grid(row=0, column=1, sticky="nsew")


# Left frame content
left_canvas = tk.Canvas(left_frame, highlightthickness=0, background='white')
left_scrollbar = ttk.Scrollbar(left_frame, orient="vertical", command=left_canvas.yview)
left_scrollable_frame = ttk.Frame(left_canvas, style='White.TFrame')
# En Windows, el evento MouseWheel puede ser usado directamente.
left_canvas.bind_all("<MouseWheel>", lambda e: left_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))




# Añadir un padding interno al canvas para centrar los widgets
left_scrollable_frame.bind(
    "<Configure>",
    lambda e: left_canvas.configure(
        scrollregion=left_canvas.bbox("all"),
        width=e.width + 10  # Agregar un padding interno de 10 píxeles a cada lado
    )
)

left_canvas.create_window((0, 0), window=left_scrollable_frame, anchor="nw")
left_canvas.configure(yscrollcommand=left_scrollbar.set)

# Frame que contiene los botones
opciones_frame = ttk.Frame(left_scrollable_frame, style='White.TFrame')
opciones_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=2)

# Configuramos el weight de las columnas del frame para que se expandan adecuadamente
opciones_frame.grid_columnconfigure(0, weight=1)
opciones_frame.grid_columnconfigure(1, weight=1)

# Botón "cargar" dentro del frame
cargar_button = ttk.Button(opciones_frame, text="Cargar Sentencia", command=lambda: ejecutar_en_hilo(cargar_pdf_sentencia), style='White.TButton')
cargar_button.grid(row=0, column=0, sticky="ew", padx=2, pady=2)

# Botón "resolver" dentro del frame
recurso_button = ttk.Button(opciones_frame, text="Cargar Recurso", state='disabled', command=lambda: ejecutar_en_hilo(cargar_pdf_recurso), style='White.TButton')
recurso_button.grid(row=0, column=1, sticky="ew", padx=2, pady=2)


# Partido de procedencia checkboxes
partido_var = tk.StringVar()
partido_frame = ttk.LabelFrame(left_scrollable_frame, text="Partido de procedencia", style='White.TLabelframe')
partido_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=5)

for partido in ["Las Palmas de Gran Canaria", "Arrecife", "Puerto del Rosario", "Galdar"]:
    check = ttk.Checkbutton(partido_frame, text=partido, variable=partido_var, onvalue=partido, offvalue='', command=update_juzgado, style='White.TCheckbutton')
    check.pack(side="top", anchor="w", padx=5, pady=2)

# Juzgado de procedencia dropdown
juzgado_var = tk.StringVar()
juzgado_label = ttk.Label(left_scrollable_frame, text="Juzgado de procedencia", style='White.TLabel')
juzgado_label.grid(row=2, column=0, sticky="ew", padx=10, pady=(5,0))
juzgado_options = ttk.Combobox(left_scrollable_frame, textvariable=juzgado_var)
juzgado_options.grid(row=3, column=0, sticky="ew", padx=10, pady=(0,5))

# Fecha calendar
fecha_label = ttk.Label(left_scrollable_frame, text="Fecha de la sentencia recurrida", style='White.TLabel')
fecha_label.grid(row=4, column=0, sticky="ew", padx=10, pady=(5,0))
fecha_cal = DateEntry(left_scrollable_frame)
fecha_cal.grid(row=5, column=0, sticky="ew", padx=10, pady=(0,5))

# Autos de procedencia
autos_var = tk.StringVar()
autos_year_var = tk.StringVar()
autos_frame = ttk.Frame(left_scrollable_frame, style='White.TFrame')
autos_frame.grid(row=6, column=0, sticky="ew", padx=10, pady=5)

autos_label = ttk.Label(autos_frame, text="Autos", style='White.TLabel')
autos_label.grid(row=0, column=0, columnspan=2, sticky="w")

# Configuración del validador
vcmd = (root.register(solo_numeros), '%P')

autos_entry = ttk.Entry(autos_frame, textvariable=autos_var, validate="key", validatecommand=vcmd)
autos_entry.grid(row=1, column=0, sticky="ew", padx=(0,10))
autos_entry.bind("<Button-3>", lambda event: mostrar_menu(event, autos_entry))
autos_year_options = ttk.Combobox(autos_frame, textvariable=autos_year_var, values=['2021', '2022', '2023'])
autos_year_options.grid(row=1, column=1, sticky="ew")

autos_frame.columnconfigure(0, weight=3)
autos_frame.columnconfigure(1, weight=1)

# Numero de recurso
recurso_var = tk.StringVar()
recurso_year_var = tk.StringVar()
recurso_frame = ttk.Frame(left_scrollable_frame, style='White.TFrame')
recurso_frame.grid(row=7, column=0, sticky="ew", padx=10, pady=5)

recurso_label = ttk.Label(recurso_frame, text="Recurso", style='White.TLabel')
recurso_label.grid(row=0, column=0, columnspan=2, sticky="w")

recurso_entry = ttk.Entry(recurso_frame, textvariable=recurso_var, validate="key", validatecommand=vcmd)
recurso_entry.grid(row=1, column=0, sticky="ew", padx=(0,10))
recurso_entry.bind("<Button-3>", lambda event: mostrar_menu(event, recurso_entry))
recurso_year_options = ttk.Combobox(recurso_frame, textvariable=recurso_year_var, values=['2021', '2022', '2023'])
recurso_year_options.grid(row=1, column=1, sticky="ew")

recurso_frame.columnconfigure(0, weight=3)
recurso_frame.columnconfigure(1, weight=1)

# Recurrente entry

recurrente_var = tk.StringVar()
recurrente_label = ttk.Label(left_scrollable_frame, text="Recurrente", style='White.TLabel')
recurrente_label.grid(row=8, column=0, sticky="ew", padx=10, pady=(5,0))
recurrente_entry = ttk.Entry(left_scrollable_frame, textvariable=recurrente_var)
recurrente_entry.grid(row=9, column=0, sticky="ew", padx=10, pady=(0,5))
recurrente_entry.bind("<Button-3>", lambda event: mostrar_menu(event, recurrente_entry))

# Impugnante

impugnante_var = tk.StringVar()
impugnante_label = ttk.Label(left_scrollable_frame, text="Impugnante", style='White.TLabel')
impugnante_label.grid(row=10, column=0, sticky="ew", padx=10, pady=(5,0))
impugnante_entry = ttk.Entry(left_scrollable_frame, textvariable=impugnante_var)
impugnante_entry.grid(row=11, column=0, sticky="ew", padx=10, pady=(0,5))
impugnante_entry.bind("<Button-3>", lambda event: mostrar_menu(event, impugnante_entry))


def update_sintesis_with_impugnante(*args):
    # Base de la frase para la actualización.
    frase_base_impugnacion = " El recurso fue impugnado por la representación letrada de "

    # Extraer el contenido actual de la síntesis.
    sintesis_text = sintesis_var.get().strip()

    # Obtener el nombre del impugnante del Entry asociado.
    impugnante_name = impugnante_var.get().strip()

    # Si hay un nombre del impugnante, verificamos su presencia y actualizamos el nombre.
    if impugnante_name:
        # Si ya contiene la frase base, actualizamos solo el nombre del impugnante.
        if frase_base_impugnacion.strip() in sintesis_text:
            # Encontrar el comienzo de la frase y actualizar solo esa parte.
            start = sintesis_text.find(frase_base_impugnacion.strip())
            end = sintesis_text.find("\n", start + len(frase_base_impugnacion)) if "\n" in sintesis_text[
                                                                                           start:] else None
            sintesis_text = (sintesis_text[
                             :start].strip() + frase_base_impugnacion + impugnante_name + "." + "\n") if end is None else (
                        sintesis_text[:start] + frase_base_impugnacion + impugnante_name + "." + "\n" + sintesis_text[end:])
        else:
            # Si no, añadimos la frase al final.
            sintesis_text += frase_base_impugnacion + impugnante_name
    else:
        # Si no hay impugnante, eliminamos la frase si existe.
        if frase_base_impugnacion.strip() in sintesis_text:
            # Encontrar y eliminar la línea entera que contiene la frase.
            start = sintesis_text.find(frase_base_impugnacion.strip())
            end = sintesis_text.find("\n", start + len(frase_base_impugnacion)) if "\n" in sintesis_text[
                                                                                           start:] else None
            sintesis_text = (sintesis_text[:start].rstrip()) if end is None else (
                        sintesis_text[:start] + sintesis_text[end + 1:])

    # Actualizamos el Entry de la síntesis con el nuevo o ajustado contenido.
    sintesis_var.set(sintesis_text)

# Asociar la nueva función al cambio de la variable impugnante_var
impugnante_var.trace("w", update_sintesis_with_impugnante)


# Revisiones variable
revisiones_var = tk.StringVar()

# Censuras variable
censuras_var = tk.StringVar()

# Add a new entry widget for 'Sintesis de la litis' after 'Recurrente'
sintesis_var = tk.StringVar()
sintesis_label = ttk.Label(left_scrollable_frame, text="Síntesis de la litis", style='White.TLabel')
sintesis_label.grid(row=12, column=0, sticky="ew", padx=10, pady=(5,0))
sintesis_entry = ttk.Entry(left_scrollable_frame, textvariable=sintesis_var, width=40)
sintesis_entry.grid(row=13, column=0, sticky="ew", padx=10, pady=(0,5))
sintesis_entry.bind("<Button-3>", lambda event: mostrar_menu(event, sintesis_entry))

# Variable del resumen de la sentencia
resumen_caso_var = tk.StringVar()

# Checkbox 'Cuestiones de orden público'
cuestiones_var = tk.BooleanVar(value=False)
cuestiones_check = ttk.Checkbutton(left_scrollable_frame, text="Cuestiones de orden público", variable=cuestiones_var, style='White.TCheckbutton')
cuestiones_check.grid(row=14, column=0, sticky="w", padx=8, pady=2)

# Checkbox 'Inadmisibilidad del recurso'
inadmisibilidad_var = tk.BooleanVar(value=False)
inadmisibilidad_check = ttk.Checkbutton(left_scrollable_frame, text="Inadmisibilidad del recurso", variable=inadmisibilidad_var, style='White.TCheckbutton')
inadmisibilidad_check.grid(row=15, column=0, sticky="w", padx=8, pady=2)

# Frame para los checkboxes que dependen de 'Infracción de las normas procesales'
dynamic_frame = ttk.Frame(left_scrollable_frame, style='White.TFrame')

# Checkbox 'Infracción de las normas procesales'
infraccion_var = tk.BooleanVar(value=False)
infraccion_check = ttk.Checkbutton(left_scrollable_frame, text="Infracción de las normas procesales", variable=infraccion_var, style='White.TCheckbutton')
infraccion_check.grid(row=16, column=0, sticky="w", padx=8, pady=2)


# Se define la función toggle_infraccion más arriba en tu código
def toggle_infraccion():
    if infraccion_var.get():
        # Mostrar el frame con los checkboxes
        dynamic_frame.grid(row=17, column=0, sticky="ew", padx=15, pady=2)
    else:
        # Ocultar el frame con los checkboxes
        dynamic_frame.grid_forget()

infraccion_var.trace("w", lambda *_: toggle_infraccion())

# Definición de los checkboxes dentro del frame 'tipos_frame'
tipos_vars = [tk.BooleanVar(value=False) for _ in range(6)]
tipos_options = [
    "Denegación de prueba",
    "Falta de motivación pidiendo nulidad",
    "Falta de motivación sin pedir nulidad",
    "Incongruencia Extrapetitum",
    "Incongruencia Infrapetitum",
    "Defectuosa grabación"
]

for i, option in enumerate(tipos_options):
    chk = ttk.Checkbutton(dynamic_frame, text=option, variable=tipos_vars[i], style='White.TCheckbutton')
    chk.pack(anchor="w", pady=2)

tipos_options[0] = """	La parte recurrente interesa la nulidad de la Sentencia, al amparo del art. 193.a) LRJS, porque en el acto del juicio oral se le denegó la práctica de prueba propuesta; la parte recurrente formuló protesta ante la denegación de dicha práctica y justificó que era necesaria.
	La parte recurrente plantea que la denegación de dichas pruebas le produce indefensión.
	El derecho a un proceso con todas las garantías, previsto en el artículo 24.2 de la Constitución Española, ha sido interpretado por el Tribunal Constitucional entre otras en su Sentencia 199/1992, de 19 de noviembre, en el sentido de que «la indefensión que prohíbe el artículo 24.1 de la Constitución no nace de la sola y simple infracción de las reglas procesales por parte de los órganos competentes, sino que se produce cuando la vulneración de las normas procesales lleva consigo la prohibición del derecho a la defensa, con el consiguiente perjuicio real y efectivo para los intereses del afectado».
	En cuanto al derecho a la prueba, la Sala Cuarta del Tribunal Supremo en Sentencia de 25 de enero de 2018, (rec. 1648/2016) señala que “la tutela judicial efectiva presupone la garantía de las personas de que pueden acudir al Juez, ser oídas, proponer y practicar pruebas, a una resolución fundada en derecho y a los recursos […]”."""
tipos_options[1] = """	La parte recurrente interesa la nulidad de la Sentencia, al amparo del art. 193.a) LRJS, por falta de motivación.
	El derecho a un proceso con todas las garantías, previsto en el artículo 24.2 de la Constitución Española, ha sido interpretado por el Tribunal Constitucional entre otras en su Sentencia 199/1992, de 19 de noviembre, en el sentido de que «la indefensión que prohíbe el artículo 24.1 de la Constitución no nace de la sola y simple infracción de las reglas procesales por parte de los órganos competentes, sino que se produce cuando la vulneración de las normas procesales lleva consigo la prohibición del derecho a la defensa, con el consiguiente perjuicio real y efectivo para los intereses del afectado».
	La doctrina del TC respecto a la necesidad de motivación de las sentencias como se razona en la STC 80/2000, de 27 de marzo, establece que: «el requisito de motivación de las sentencias, como parte integrante del derecho a la tutela judicial efectiva, no impone que la resolución ofrezca una exhaustiva descripción del proceso intelectual llevado a cabo por el juzgador para resolver, ni una pormenorizada respuesta a todas las alegaciones de las partes, ni un determinado alcance o entidad en el razonamiento empleado, ni siquiera la corrección jurídica interna de la fundamentación empleada; basta que la argumentación vertida, con independencia de su parquedad o concentración, cumpla la doble finalidad de exteriorizar el motivo de la decisión, su ‘ratio decidendi’ (STC 5/1995, de 10 de enero), excluyente de un mero voluntarismo selectivo o de la pura arbitrariedad, y que permita su eventual revisión jurisdiccional a través del efectivo ejercicio de los recursos establecidos, como aquí sucede con este recurso de amparo».
	La doctrina constitucional y también la jurisprudencia del Tribunal Supremo vienen siendo constantes en señalar que la motivación debe cumplir con dos requisitos: el primero es que debe ser suficiente, lo que no quiere decir exhaustiva, y en segundo lugar fundamentada o argumentada en derecho y no en otros criterios del juez/a o tribunal."""
tipos_options[2] = """	La parte recurrente interesa invoca el art. 193.a) LRJS, por falta de motivación. Sin embargo, no plantea la nulidad de la sentencia.
	El planteamiento resulta interesante, pero dándose la circunstancia de que no se solicita la nulidad de la sentencia, entendemos que debería haberse articulado al amparo de la letra c) del artículo 193 LRJS y por tanto en este punto carece de sentido entrar a analizar más profundamente lo planteado, sin perjuicio de su análisis más adelante cuando tratemos los motivos impugnatorios del art. 193.c) LRJS."""
tipos_options[3] = """	La parte recurrente interesa la nulidad de la Sentencia, al amparo del art. 193.a) LRJS, por incongruencia "extra petitum".
	Sobre la figura de la congruencia, se ha ocupado el Tribunal Constitucional en numerosas ocasiones. Así, desde la STC 20/1982, de 5 de mayo, viene señalando que “la misma se mide por el ajuste o adecuación entre lo resuelto y los términos en que las partes han formulado sus pretensiones y peticiones, de manera tal que no puede la sentencia otorgar más de lo que se hubiera pedido en la demanda, ni menos de lo que hubiera sido admitido por el demandado, ni otorgar otra cosa diferente, que no hubiera sido pretendida”. 
	El Tribunal Supremo, en ST de 22 de diciembre de 2016 señala que la congruencia constituye “un ajuste sustancial entre lo pedido y lo resuelto que, por tanto, no exige del fallo una conformidad literal y rígida con las peticiones de las partes, sino racional y flexible”.
	La incongruencia "extra petitum", es la que se origina cuando se resuelve sobre cuestiones distintas y ajenas a lo solicitado por las partes, lo que implica una invasión frontal del derecho de defensa contradictorio, privando a los litigantes de la facultad de alegar lo que a su derecho proceda o lo que estimen conveniente a sus legítimos intereses."""
tipos_options[4] = """	La parte recurrente interesa la nulidad de la Sentencia, al amparo del art. 193.a) LRJS, por incongruencia omisiva.
	Sobre la figura de la congruencia, se ha ocupado el Tribunal Constitucional en numerosas ocasiones. Así, desde la STC 20/1982, de 5 de mayo, viene señalando que “la misma se mide por el ajuste o adecuación entre lo resuelto y los términos en que las partes han formulado sus pretensiones y peticiones, de manera tal que no puede la sentencia otorgar más de lo que se hubiera pedido en la demanda, ni menos de lo que hubiera sido admitido por el demandado, ni otorgar otra cosa diferente, que no hubiera sido pretendida”.
	El Tribunal Supremo, en ST de 22 de diciembre de 2016 señala que la congruencia constituye “un ajuste sustancial entre lo pedido y lo resuelto que, por tanto, no exige del fallo una conformidad literal y rígida con las peticiones de las partes, sino racional y flexible”.
	La incongruencia "infra petitum", es la que se origina por falta de respuesta judicial razonada a alguna de las cuestiones o elementos esenciales de la pretensión sometidos por las partes a la consideración del Juzgador/a, y cuyo conocimiento y decisión sea trascendente para fijar el fallo.
	No todos los casos de ausencia de respuesta judicial expresa producen una indefensión constitucionalmente relevante, debiendo ponderarse las circunstancias concurrentes en cada caso para determinar:
		a)	Si el silencio de la resolución judicial constituye una auténtica lesión del art. 24.1 CE o si, por el contrario, puede razonablemente interpretarse como una desestimación tácita.
		b)	Si efectivamente se ha planteado la cuestión cuyo conocimiento se afirma eludido por el Tribunal.
		c)	Si la incongruencia omisiva apreciada causó un efectivo perjuicio de los derechos de defensa de quien se queja en amparo."""
tipos_options[5] = """	La parte recurrente interesa la nulidad de la Sentencia, al amparo del art. 193.a) LRJS, por defectos en la grabación de la vista.
	La defectuosa grabación del juicio oral no conlleva necesariamente la anulación de las actuaciones de instancia cuando no se ha causado indefensión a la parte recurrente. En este sentido, el Tribunal Constitucional entre otras en su Sentencia 199/1992 señala que «la indefensión que prohíbe el artículo 24.1 de la Constitución no nace de la sola y simple infracción de las reglas procesales por parte de los órganos competentes, sino que se produce cuando la vulneración de las normas procesales lleva consigo la prohibición del derecho a la defensa, con el consiguiente perjuicio real y efectivo para los intereses del afectado».
	Igualmente, la STC n.º 4/2004, de 14 de enero, explica que la destrucción del acta de juicio no es en sí una causa de nulidad del juicio mismo si ese juicio materialmente se celebró con todas las garantías y en él las partes pudieron ejercer sus derechos de alegación y prueba sin limitación."""

# Añadir nuevos checkboxes debajo de 'Infracción de normas procesales'
revision_var = tk.BooleanVar(value=False)
revision_check = ttk.Checkbutton(left_scrollable_frame, text="Revisión de hechos probados", variable=revision_var, style='White.TCheckbutton')
revision_check.grid(row=18, column=0, sticky="w", padx=8, pady=2)

# Botón para 'Analizar revisiones fácticas'
analizar_revisiones_button = ttk.Button(left_scrollable_frame, text="Analizar revisiones fácticas", state='disabled', command=lambda: ejecutar_en_hilo(analizar_revisiones_facticas))
analizar_revisiones_button.grid(row=19, column=0, padx=10, pady=10, sticky="ew")
analizar_revisiones_button.grid_remove()  # Ocultar por defecto

def toggle_revision_button():
    if revision_var.get():
        analizar_revisiones_button.grid()  # Mostrar el botón
    else:
        analizar_revisiones_button.grid_remove()  # Ocultar el botón

# Vincular la función toggle_revision_button al cambio de estado del checkbox 'Revisión de hechos probados'
revision_var.trace("w", lambda *args: toggle_revision_button())

# Llamar a la función una vez para configurar el estado inicial del botón
toggle_revision_button()

sustantivas_var = tk.BooleanVar(value=False)
sustantivas_check = ttk.Checkbutton(left_scrollable_frame, text="Infracción de normas sustantivas", variable=sustantivas_var, style='White.TCheckbutton')
sustantivas_check.grid(row=20, column=0, sticky="w", padx=8, pady=2)

# Cuadros de texto para 'Infracción de normas sustantivas'
articulos_var = tk.StringVar()
jurisprudencia_var = tk.StringVar()

# Artículos infringidos
articulos_label = ttk.Label(left_scrollable_frame, text="Artículos infringidos", style='White.TLabel')
articulos_label.grid(row=22, column=0, sticky="ew", padx=10, pady=(5,0))
articulos_entry = ttk.Entry(left_scrollable_frame, textvariable=articulos_var, width=40)
articulos_entry.grid(row=23, column=0, sticky="ew", padx=10, pady=(0,5))
articulos_entry.bind("<Button-3>", lambda event: mostrar_menu(event, articulos_entry))

# Jurisprudencia infringida
jurisprudencia_label = ttk.Label(left_scrollable_frame, text="Jurisprudencia infringida", style='White.TLabel')
jurisprudencia_label.grid(row=24, column=0, sticky="ew", padx=10, pady=(5,0))
jurisprudencia_entry = ttk.Entry(left_scrollable_frame, textvariable=jurisprudencia_var, width=40)
jurisprudencia_entry.grid(row=25, column=0, sticky="ew", padx=10, pady=(0,5))
jurisprudencia_entry.bind("<Button-3>", lambda event: mostrar_menu(event, jurisprudencia_entry))

# Botón para 'Analizar revisiones fácticas'
analizar_censuras_button = ttk.Button(left_scrollable_frame, text="Analizar censuras jurídicas", state='disabled', command=lambda: ejecutar_en_hilo(analizar_censuras_juridicas))
analizar_censuras_button.grid(row=26, column=0, padx=10, pady=10, sticky="ew")
analizar_censuras_button.grid_remove()  # Ocultar el botón

def toggle_censuras():
    if sustantivas_var.get():
        articulos_label.grid()
        articulos_entry.grid()
        jurisprudencia_label.grid()
        jurisprudencia_entry.grid()
        analizar_censuras_button.grid()  # Mostrar el botón
    else:
        articulos_label.grid_remove()
        articulos_entry.grid_remove()
        jurisprudencia_label.grid_remove()
        jurisprudencia_entry.grid_remove()
        analizar_censuras_button.grid_remove()  # Ocultar el botón

# Vincular la función toggle_revision_button al cambio de estado del checkbox 'Revisión de hechos probados'
sustantivas_var.trace("w", lambda *args: toggle_censuras())

# Llamar a la función una vez para configurar el estado inicial del botón
toggle_censuras()

# Frame "Costas, depósitos y consignaciones"
costas_frame = ttk.LabelFrame(left_scrollable_frame, text="Costas, depósitos y consignaciones", borderwidth=1, relief="groove", style='White.TLabelframe')
costas_frame.grid(row=27, column=0, sticky="ew", padx=10, pady=5)

# Variables para los radiobuttons
costas_var = tk.StringVar(value=None)

opciones_costas = [
    "Estimación",
    "Estimación parcial",
    "Desestimación sin costas",
    "Desestimación con costas"
]

# Crear Radiobuttons usando la lista
for opcion in opciones_costas:
    ttk.Radiobutton(costas_frame, text=opcion, value=opcion, variable=costas_var, command=update_text, style='White.TRadiobutton').pack(anchor="w", padx=5, pady=2)


# Subframe y widgets para "Estimación"
estimacion_frame = ttk.LabelFrame(left_scrollable_frame, text="Depósitos y Aseguramientos", borderwidth=2, relief="groove", style='White.TLabelframe')
estimacion_frame.grid(row=28, column=0, sticky="ew", padx=10, pady=5)
estimacion_frame.grid_remove()  # Ocultar por defecto
trabajador_empresa_var = tk.StringVar(value="Trabajador")
ttk.Radiobutton(estimacion_frame, text="Trabajador", variable=trabajador_empresa_var, value="Trabajador", style='White.TRadiobutton').pack(anchor="w", pady=2)
ttk.Radiobutton(estimacion_frame, text="Empresa", variable=trabajador_empresa_var, value="Empresa", style='White.TRadiobutton').pack(anchor="w", pady=2)
ttk.Label(estimacion_frame, text="Revocación", style='White.TLabel').pack(anchor="w", padx=10, pady=(5,0))
revocacion_var = tk.StringVar()
revocacion_entry = ttk.Entry(estimacion_frame, textvariable=revocacion_var, width=40)
revocacion_entry.pack(anchor="w", padx=10, pady=(0,5))
revocacion_entry.bind("<Button-3>", lambda event: mostrar_menu(event, revocacion_entry))

# Subframe y widgets para "Estimación parcial"
estimacion_parcial_frame = ttk.LabelFrame(left_scrollable_frame, text="Depósitos y Aseguramientos", borderwidth=2, relief="groove", style='White.TLabelframe')
estimacion_parcial_frame.grid(row=29, column=0, sticky="ew", padx=10, pady=5)
estimacion_parcial_frame.grid_remove()  # Ocultar por defecto
ttk.Radiobutton(estimacion_parcial_frame, text="Trabajador", variable=trabajador_empresa_var, value="Trabajador", style='White.TRadiobutton').pack(anchor="w", pady=2)
ttk.Radiobutton(estimacion_parcial_frame, text="Empresa", variable=trabajador_empresa_var, value="Empresa", style='White.TRadiobutton').pack(anchor="w", pady=2)
ttk.Label(estimacion_parcial_frame, text="Revocación parcial", style='White.TLabel').pack(anchor="w", padx=10, pady=(5,0))
revocacion_parcial_var = tk.StringVar()
revocacion_parcial_entry = ttk.Entry(estimacion_parcial_frame, textvariable=revocacion_parcial_var, width=40)
revocacion_parcial_entry.pack(anchor="w", padx=10, pady=(0,5))
revocacion_parcial_entry.bind("<Button-3>", lambda event: mostrar_menu(event, revocacion_parcial_entry))

# Widgets para "Desestimación con costas"
cuantia_costas_var = tk.StringVar(value="800")
cuantia_costas_label = ttk.Label(left_scrollable_frame, text="Cuantía de la condena en costas", style='White.TLabel')
cuantia_costas_label.grid(row=30, column=0, sticky="ew", padx=10, pady=(5,0))
cuantia_costas_label.grid_remove()  # Ocultar por defecto
cuantia_costas_entry = ttk.Entry(left_scrollable_frame, textvariable=cuantia_costas_var, width=40)
cuantia_costas_entry.grid(row=31, column=0, sticky="ew", padx=10, pady=(0,5))
cuantia_costas_entry.grid_remove()  # Ocultar por defecto
cuantia_costas_entry.bind("<Button-3>", lambda event: mostrar_menu(event, cuantia_costas_entry))

# Subframe y widgets para "Desestimación sin costas"
razon_var = tk.StringVar(value="beneficiario de justicia gratuita, ex artículo 2.d) Ley 1/1996")
razones_frame = ttk.LabelFrame(left_scrollable_frame, text="Razones para su no imposición", borderwidth=2, relief="groove", style='White.TLabelframe')
razones_frame.grid(row=32, column=0, sticky="ew", padx=10, pady=5)
razones_frame.grid_remove()  # Ocultar por defecto
# Diccionario que mapea las opciones con los textos específicos
razones_texts = {
    "Beneficiario de justicia gratuita": "beneficiario de justicia gratuita, ex artículo 2.d) Ley 1/1996",
    "Sindicato": "un sindicato, ex artículo 20.4 LRJS",
    "Funcionario público": "un funcionario público, ex artículo 21.5 LRJS",
    "Personal laboral": "personal laboral, ex artículo 2.d) Ley 1/1996",
    "Entidad gestora": "una Entidad Gestora de la Seguridad Social, ex artículo 2.b) Ley 1/1996"
}


# Crear Radiobuttons usando el diccionario
for razon, texto in razones_texts.items():
    rb = ttk.Radiobutton(razones_frame, text=razon, variable=razon_var, value=texto, style='White.TRadiobutton')
    rb.pack(anchor="w", pady=2)


def toggle_costas_frames():
    # Diccionario que mapea los valores de costas_var a los widgets a mostrar
    frames_to_toggle = {
        "Estimación": [estimacion_frame],
        "Estimación parcial": [estimacion_parcial_frame],
        "Desestimación con costas": [cuantia_costas_label, cuantia_costas_entry],
        "Desestimación sin costas": [razones_frame]
    }

    # Primero, ocultar todos los frames/subframes
    for widget_list in frames_to_toggle.values():
        for widget in widget_list:
            widget.grid_remove()

    # Luego, mostrar los widgets correspondientes a la selección actual
    selected_option = costas_var.get()
    for widget in frames_to_toggle.get(selected_option, []):
        widget.grid()

# Vincular la función toggle_costas_frames al cambio de la variable costas_var
costas_var.trace("w", lambda *_: toggle_costas_frames())

# Iniciar la función una vez para mostrar el frame/subframe inicial
toggle_costas_frames()

# Asegurarse de que update_text se llama cuando cambian los valores de los nuevos checkboxes
cuestiones_var.trace("w", update_text)
inadmisibilidad_var.trace("w", update_text)
infraccion_var.trace("w", update_text)
sintesis_var.trace("w", update_text)
revision_var.trace("w", update_text)
sustantivas_var.trace("w", update_text)
articulos_var.trace("w", update_text)
jurisprudencia_var.trace("w", update_text)
trabajador_empresa_var.trace("w", update_text)
revocacion_var.trace("w", update_text)
revocacion_parcial_var.trace("w", update_text)
cuantia_costas_var.trace("w", update_text)
razon_var.trace("w", update_text)


for var in tipos_vars:
    var.trace("w", update_text)


# Buttons
button_frame = ttk.Frame(left_scrollable_frame, style='White.TFrame')
button_frame.grid(row=100, column=0, sticky="ew", padx=10, pady=2)

# Configuramos el grid del button_frame para alinear mejor los widgets
button_frame.grid_columnconfigure(0, weight=1)
button_frame.grid_columnconfigure(1, weight=1)


# Botón "Guardar"
save_button = ttk.Button(button_frame, text="Guardar", command=save_to_doc)
# Colocamos el botón "Guardar" en la primera columna
save_button.grid(row=0, column=0, padx=2, pady=2, sticky="ew")

# Crear atajo de teclado para crear texto

root.bind('<Control-Return>', lambda event: save_button.invoke())

# Botón "Limpiar"
clear_button = ttk.Button(button_frame, text="Limpiar", command= lambda: ejecutar_en_hilo(clear_all))
# Colocamos el botón "Limpiar" en la segunda columna
clear_button.grid(row=0, column=1, padx=2, pady=2, sticky="ew")

# Crear atajo de teclado para limpiar

root.bind('<Control-BackSpace>', lambda event: clear_button.invoke())

# Barra de progreso
progress_bar = Progressbar(button_frame, length=100, mode='indeterminate')
# Colocamos la barra de progreso en la tercera columna pero inicialmente no la muestres
progress_bar.grid(row=0, column=2, padx=3, pady=2, sticky="ew")
progress_bar.grid_remove()  # Inicia con la barra de progreso oculta

def start_progress_bar():
    """Función para comenzar la animación de la barra de progreso."""
    progress_bar.pack()
    progress_bar.start()

def stop_progress_bar():
    """Función para detener y ocultar la barra de progreso."""
    progress_bar.stop()
    progress_bar.pack_forget()

# Right text widget
right_text = tk.Text(right_frame, wrap="word", padx=10, pady=10, font=("SegoeUI", 9), relief='groove', bd=2)
right_text.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
right_text.bind("<Button-3>", lambda event: mostrar_menu_derecha(event, right_text))
right_frame.grid_rowconfigure(0, weight=1)
right_frame.grid_columnconfigure(0, weight=1)

# Crear un menú contextual
menu_contextual = tk.Menu(root, tearoff=0)
menu_contextual.add_command(label="Resolver revisión fáctica", command=analizar_texto)

# Enlazar la liberación del botón izquierdo para mostrar el menú contextual
right_text.bind("<ButtonRelease-1>", menu_contextual_resolver_revision_factica)


# Trigger updates
partido_var.trace("w", update_text)
juzgado_var.trace("w", update_text)
autos_var.trace("w", update_text)
autos_year_var.trace("w", update_text)
recurrente_var.trace("w", update_text)
revisiones_var.trace("w", update_text)
censuras_var.trace("w", update_text)
impugnante_var.trace("w", update_text)

for widget in left_scrollable_frame.winfo_children():
    bind_to_mousewheel(widget)

# Ajustamos el embalaje del Canvas y el Scrollbar para reducir el espacio entre los widgets y el scrollbar.
left_scrollbar.pack(side="right", fill="y", expand=False)
left_canvas.pack(side="left", fill="both", expand=True, padx=(10, 0), pady=0)  # Agregar un padding de 10 píxeles a la izquierda

root.grid_columnconfigure(1, weight=3)
root.grid_rowconfigure(0, weight=1)

# Centra la ventana en la pantalla
center_window(root)

# Imagen de inicio # Cargar imagen
image_path = "login.png"  # Asegúrate de que este sea el nombre de tu archivo y que esté en el mismo directorio que tu script
image = PhotoImage(file=image_path)


# Crear un label y mostrar la imagen
image_label = tk.Label(right_frame, image=image, bg='white')
image_label.grid(row=0, column=0, sticky="nsew")  # Usa grid para colocar el label en la ventana principal
right_frame.grid_rowconfigure(0, weight=1)
right_frame.grid_columnconfigure(0, weight=1)

# Función para ocultar la imagen
def hide_image():
    image_label.grid_forget()  # Eliminar el label de la ventana con grid_forget()

# Establecer un temporizador para llamar a hide_image después de 5 segundos (5000 milisegundos)
root.after(4000, hide_image)

# Inicia la verificación del portapapeles
verificar_portapapeles()

root.mainloop()
